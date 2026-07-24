import os
import re
import io
import gc
import zipfile
import tempfile
import json
import random
import datetime
from urllib.request import urlopen, Request as _UrlRequest
from urllib.error import HTTPError as _UrlHTTPError

from fastapi import FastAPI, HTTPException, Depends, BackgroundTasks, UploadFile, File, Form, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.responses import FileResponse, Response, HTMLResponse, StreamingResponse
from pydantic import BaseModel
from typing import List, Optional
from sqlalchemy.orm import Session
from sqlalchemy import text as sa_text

from auth import hash_password, verify_password, create_access_token, verify_token
from config import ADMIN_EMAILS
from database import SessionLocal, engine, get_db
import concept_engine
from models import (
    Base, User as DBUser, MockTest as DBMockTest,
    Question as DBQuestion, TestAttempt as DBTestAttempt, Answer as DBAnswer,
    AdminEmail as DBAdminEmail, StudentProfile as DBStudentProfile,
    ReviewItem as DBReviewItem, QuestionFlag as DBQuestionFlag,
    ResolvedMistake as DBResolvedMistake, Bookmark as DBBookmark,
    KnowledgeSource as DBKnowledgeSource, KnowledgeChunk as DBKnowledgeChunk,
    SyllabusProgress as DBSyllabusProgress, GuidedProgress as DBGuidedProgress,
    DailyMissionDone as DBDailyMissionDone, MainsAnswer as DBMainsAnswer,
    ChatMessage as DBChatMessage,
    NcertPdf as DBNcertPdf,
    ReaderNote as DBReaderNote,
    NcertReading as DBNcertReading,
    ConceptJob as DBConceptJob,
    ConceptInventory as DBConceptInventory,
    InterviewPrep as DBInterviewPrep,
    MockScore as DBMockScore, ExamGoal as DBExamGoal,
    CsatPyqPaper as DBCsatPyqPaper,
    MentorTopic as DBMentorTopic,
    ConceptAttempt as DBConceptAttempt,
    ConceptMastery as DBConceptMastery,
    SkillMastery as DBSkillMastery,
    TeachingEvent as DBTeachingEvent,
)
from gemini_service import (
    generate_questions, generate_and_parse_questions,
    explain_concept, analyze_performance, chat_with_mentor,
    generate_mentor_report, generate_ncert_mcqs, generate_verified_questions,
    generate_verified_questions_stream,
    diagnose_mistake,
    generate_study_notes, generate_flashcards, generate_mnemonics,
    generate_mindmap, current_affairs_analysis, extract_mcqs_from_text,
    evaluate_mains_answer, weekly_mentor_narrative,
)
import syllabus_data
import geo_data
import syllabus_tracker_data
import study_planner
import program as guided_program
import prepos
import teaching
import diagnostic

Base.metadata.create_all(bind=engine)

# Set True once pgvector is confirmed available on the (Postgres) database. While
# False, all semantic search transparently falls back to keyword (ILIKE) search.
VECTOR_OK = False

def _ensure_schema():
    """Lightweight migration: add columns newer code expects to pre-existing DBs.
    Works on both SQLite (local) and PostgreSQL (production)."""
    try:
        dialect = engine.dialect.name
        with engine.connect() as conn:
            # New per-question metadata columns for book/subject/topic-wise generation.
            new_q_cols = ["book", "chapter", "topic", "difficulty", "question_type"]
            if dialect == "sqlite":
                qcols = [r[1] for r in conn.execute(sa_text("PRAGMA table_info(questions)"))]
                if "subject" not in qcols:
                    conn.execute(sa_text("ALTER TABLE questions ADD COLUMN subject VARCHAR")); conn.commit()
                for col in new_q_cols:
                    if col not in qcols:
                        conn.execute(sa_text(f"ALTER TABLE questions ADD COLUMN {col} VARCHAR")); conn.commit()
                ucols = [r[1] for r in conn.execute(sa_text("PRAGMA table_info(users)"))]
                if "created_at" not in ucols:
                    conn.execute(sa_text("ALTER TABLE users ADD COLUMN created_at DATETIME")); conn.commit()
                if "phone" not in ucols:
                    conn.execute(sa_text("ALTER TABLE users ADD COLUMN phone VARCHAR")); conn.commit()
                # New per-answer learning-loop columns.
                acols = [r[1] for r in conn.execute(sa_text("PRAGMA table_info(answers)"))]
                if "time_taken" not in acols:
                    conn.execute(sa_text("ALTER TABLE answers ADD COLUMN time_taken INTEGER")); conn.commit()
                for col in ("confidence", "wrong_reason"):
                    if col not in acols:
                        conn.execute(sa_text(f"ALTER TABLE answers ADD COLUMN {col} VARCHAR")); conn.commit()
                # New DAF-style student_profiles columns.
                try:
                    pcols = [r[1] for r in conn.execute(sa_text("PRAGMA table_info(student_profiles)"))]
                    for col in ("category", "district", "prep_location", "prep_city",
                                "coaching_status", "coaching_method",
                                "prep_level", "knowledge_level", "comprehension_skill",
                                "reading_speed", "study_time_windows", "study_place",
                                "prep_intensity", "failure_stage", "failure_reason", "materials_owned",
                                "full_name", "parent_name", "dob", "gender",
                                "marital_status", "mains_language",
                                "phone", "email", "address",
                                "graduation_stream", "schooling_medium",
                                "degree_percentage", "additional_qualification",
                                "work_experience"):
                        if col not in pcols:
                            conn.execute(sa_text(f"ALTER TABLE student_profiles ADD COLUMN {col} VARCHAR")); conn.commit()
                    for col in ("diagnostic_gs", "diagnostic_csat", "age"):
                        if col not in pcols:
                            conn.execute(sa_text(f"ALTER TABLE student_profiles ADD COLUMN {col} INTEGER")); conn.commit()
                    if "progress_history" not in pcols:
                        conn.execute(sa_text("ALTER TABLE student_profiles ADD COLUMN progress_history TEXT")); conn.commit()
                except Exception:
                    pass
                try:
                    kcols = [r[1] for r in conn.execute(sa_text("PRAGMA table_info(knowledge_sources)"))]
                    for col in ("description", "file_type", "taxonomy", "raw_b64", "proc_mode"):
                        if col not in kcols:
                            conn.execute(sa_text(f"ALTER TABLE knowledge_sources ADD COLUMN {col} TEXT")); conn.commit()
                except Exception:
                    pass
                try:
                    ncols = [r[1] for r in conn.execute(sa_text("PRAGMA table_info(ncert_pdfs)"))]
                    if "src_url" not in ncols:
                        conn.execute(sa_text("ALTER TABLE ncert_pdfs ADD COLUMN src_url TEXT")); conn.commit()
                except Exception:
                    pass
                # Reader notes: colour label + spaced-repetition revision tracking.
                try:
                    rcols = [r[1] for r in conn.execute(sa_text("PRAGMA table_info(reader_notes)"))]
                    if "label" not in rcols:
                        conn.execute(sa_text("ALTER TABLE reader_notes ADD COLUMN label VARCHAR")); conn.commit()
                    if "revise_stage" not in rcols:
                        conn.execute(sa_text("ALTER TABLE reader_notes ADD COLUMN revise_stage INTEGER DEFAULT 0")); conn.commit()
                    if "last_revised" not in rcols:
                        conn.execute(sa_text("ALTER TABLE reader_notes ADD COLUMN last_revised DATETIME")); conn.commit()
                    if "next_review" not in rcols:
                        conn.execute(sa_text("ALTER TABLE reader_notes ADD COLUMN next_review DATETIME")); conn.commit()
                except Exception:
                    pass
                # Concept jobs: progress-tracking columns for the extraction UI.
                try:
                    cjcols = [r[1] for r in conn.execute(sa_text("PRAGMA table_info(concept_jobs)"))]
                    if "stage" not in cjcols:
                        conn.execute(sa_text("ALTER TABLE concept_jobs ADD COLUMN stage VARCHAR")); conn.commit()
                    if "progress" not in cjcols:
                        conn.execute(sa_text("ALTER TABLE concept_jobs ADD COLUMN progress INTEGER DEFAULT 0")); conn.commit()
                    if "started_at" not in cjcols:
                        conn.execute(sa_text("ALTER TABLE concept_jobs ADD COLUMN started_at DATETIME")); conn.commit()
                except Exception:
                    pass
                # Teaching events: time-to-mastery column.
                try:
                    tecols = [r[1] for r in conn.execute(sa_text("PRAGMA table_info(teaching_events)"))]
                    if "seconds" not in tecols:
                        conn.execute(sa_text("ALTER TABLE teaching_events ADD COLUMN seconds INTEGER DEFAULT 0")); conn.commit()
                except Exception:
                    pass
            else:
                conn.execute(sa_text("ALTER TABLE questions ADD COLUMN IF NOT EXISTS subject VARCHAR")); conn.commit()
                for col in new_q_cols:
                    conn.execute(sa_text(f"ALTER TABLE questions ADD COLUMN IF NOT EXISTS {col} VARCHAR")); conn.commit()
                # AML (Mastery Loop) reservoir metadata on questions.
                for col in ("concept_key", "pattern"):
                    conn.execute(sa_text(f"ALTER TABLE questions ADD COLUMN IF NOT EXISTS {col} VARCHAR")); conn.commit()
                conn.execute(sa_text("ALTER TABLE questions ADD COLUMN IF NOT EXISTS material_ref TEXT")); conn.commit()
                conn.execute(sa_text("ALTER TABLE users ADD COLUMN IF NOT EXISTS created_at TIMESTAMP")); conn.commit()
                conn.execute(sa_text("ALTER TABLE users ADD COLUMN IF NOT EXISTS phone VARCHAR")); conn.commit()
                # New per-answer learning-loop columns.
                conn.execute(sa_text("ALTER TABLE answers ADD COLUMN IF NOT EXISTS time_taken INTEGER")); conn.commit()
                conn.execute(sa_text("ALTER TABLE answers ADD COLUMN IF NOT EXISTS confidence VARCHAR")); conn.commit()
                conn.execute(sa_text("ALTER TABLE answers ADD COLUMN IF NOT EXISTS wrong_reason VARCHAR")); conn.commit()
                # New DAF-style student_profiles columns.
                for col in ("category", "district", "prep_location", "prep_city",
                            "coaching_status", "coaching_method",
                            "prep_level", "knowledge_level", "comprehension_skill",
                            "reading_speed", "study_time_windows", "study_place",
                            "prep_intensity", "failure_stage", "failure_reason", "materials_owned",
                            "full_name", "parent_name", "dob", "gender",
                            "marital_status", "mains_language",
                            "phone", "email", "address",
                            "graduation_stream", "schooling_medium",
                            "degree_percentage", "additional_qualification",
                            "work_experience"):
                    conn.execute(sa_text(f"ALTER TABLE student_profiles ADD COLUMN IF NOT EXISTS {col} VARCHAR")); conn.commit()
                for col in ("diagnostic_gs", "diagnostic_csat", "age"):
                    conn.execute(sa_text(f"ALTER TABLE student_profiles ADD COLUMN IF NOT EXISTS {col} INTEGER")); conn.commit()
                # AML (Mastery Loop): learner mode + exam timeline.
                conn.execute(sa_text("ALTER TABLE student_profiles ADD COLUMN IF NOT EXISTS prep_mode VARCHAR")); conn.commit()
                conn.execute(sa_text("ALTER TABLE student_profiles ADD COLUMN IF NOT EXISTS exam_date DATE")); conn.commit()
                conn.execute(sa_text("ALTER TABLE student_profiles ADD COLUMN IF NOT EXISTS progress_history TEXT")); conn.commit()
                conn.execute(sa_text("ALTER TABLE teaching_events ADD COLUMN IF NOT EXISTS seconds INTEGER DEFAULT 0")); conn.commit()
                # New knowledge_sources columns (description / file type / taxonomy / resume).
                for col in ("description", "file_type", "taxonomy", "raw_b64", "proc_mode"):
                    conn.execute(sa_text(f"ALTER TABLE knowledge_sources ADD COLUMN IF NOT EXISTS {col} TEXT")); conn.commit()
                # NCERT PDFs: source URL so chapters stream from source (DB stays tiny).
                conn.execute(sa_text("ALTER TABLE ncert_pdfs ADD COLUMN IF NOT EXISTS src_url TEXT")); conn.commit()
                # Reader notes: colour label + spaced-repetition revision tracking.
                conn.execute(sa_text("ALTER TABLE reader_notes ADD COLUMN IF NOT EXISTS label VARCHAR")); conn.commit()
                conn.execute(sa_text("ALTER TABLE reader_notes ADD COLUMN IF NOT EXISTS revise_stage INTEGER DEFAULT 0")); conn.commit()
                conn.execute(sa_text("ALTER TABLE reader_notes ADD COLUMN IF NOT EXISTS last_revised TIMESTAMP")); conn.commit()
                conn.execute(sa_text("ALTER TABLE reader_notes ADD COLUMN IF NOT EXISTS next_review TIMESTAMP")); conn.commit()
                # Concept jobs: progress-tracking columns for the extraction UI.
                try:
                    conn.execute(sa_text("ALTER TABLE concept_jobs ADD COLUMN IF NOT EXISTS stage VARCHAR")); conn.commit()
                    conn.execute(sa_text("ALTER TABLE concept_jobs ADD COLUMN IF NOT EXISTS progress INTEGER DEFAULT 0")); conn.commit()
                    conn.execute(sa_text("ALTER TABLE concept_jobs ADD COLUMN IF NOT EXISTS started_at TIMESTAMP")); conn.commit()
                except Exception:
                    pass
                # pgvector: semantic-search embeddings for the knowledge base (RAG).
                try:
                    conn.execute(sa_text("CREATE EXTENSION IF NOT EXISTS vector")); conn.commit()
                    conn.execute(sa_text(
                        "ALTER TABLE knowledge_chunks ADD COLUMN IF NOT EXISTS embedding vector(768)")); conn.commit()
                    try:
                        conn.execute(sa_text(
                            "CREATE INDEX IF NOT EXISTS idx_kchunks_embedding ON knowledge_chunks "
                            "USING hnsw (embedding vector_cosine_ops)")); conn.commit()
                    except Exception:
                        pass    # index is an optimisation; search still works without it
                    globals()["VECTOR_OK"] = True
                except Exception:
                    pass        # extension unavailable / no permission → stay on keyword search
    except Exception:
        pass

_ensure_schema()

# ── Verified previous-year question bank ──────────────────────────────────────
SYSTEM_USER_EMAIL = "official@iasmentor.system"
YEARWISE_DURATION_MIN = 120  # full year-wise PYQ papers default to 120 minutes
PYQ_BANK_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "pyq_bank.json")

def _load_pyq_bank():
    """Read the verified previous-year question bank from disk once at startup."""
    if not os.path.exists(PYQ_BANK_FILE):
        return {"papers": []}
    try:
        with open(PYQ_BANK_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {"papers": []}

PYQ_BANK = _load_pyq_bank()

# Flatten every verified question and tag it with its year + paper. The same bank
# then serves two views: YEAR-WISE (full official papers) and SUBJECT-WISE
# (questions grouped/filtered by subject across all years). No AI is involved —
# these are real exam questions only.
PYQ_QUESTIONS = [
    {**q, "year": paper.get("year"), "paper": paper.get("paper", "UPSC Prelims GS Paper I")}
    for paper in PYQ_BANK.get("papers", [])
    for q in paper.get("questions", [])
]
PYQ_YEARS = sorted({q["year"] for q in PYQ_QUESTIONS if q.get("year")}, reverse=True)
PYQ_SUBJECTS = sorted({q["subject"] for q in PYQ_QUESTIONS if q.get("subject")})

def _load_json_file(fname, default):
    path = os.path.join(os.path.dirname(os.path.abspath(__file__)), fname)
    if not os.path.exists(path):
        return default
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return default

# Descriptive UPSC Mains previous-year questions (browse-only) and original
# AIVORA CSAT practice sets.
MAINS_BANK = _load_json_file("mains_pyq.json", {"papers": []})
MAINS_PAPERS = MAINS_BANK.get("papers", [])
MAINS_QUESTIONS = [
    {**q, "year": p.get("year"), "paper": p.get("paper"), "paper_code": p.get("paper_code")}
    for p in MAINS_PAPERS for q in p.get("questions", [])
]
MAINS_YEARS = sorted({p["year"] for p in MAINS_PAPERS if p.get("year")}, reverse=True)
_MAINS_PAPER_ORDER = ["GS1", "GS2", "GS3", "GS4", "ESSAY"]
MAINS_PAPER_DEFS = []
for code in _MAINS_PAPER_ORDER:
    nm = next((p.get("paper") for p in MAINS_PAPERS if p.get("paper_code") == code), None)
    if nm:
        MAINS_PAPER_DEFS.append({"code": code, "name": nm})
MAINS_SUBJECTS = sorted({q.get("subject") for q in MAINS_QUESTIONS if q.get("subject")})

CSAT_BANK = _load_json_file("csat_practice.json", {"areas": []})
CSAT_AREAS = CSAT_BANK.get("areas", [])

# Year-wise official CSAT (Paper II) previous-year papers. Each paper holds its
# reading-comprehension passages plus 80 questions with the official answer key,
# so the app can serve full past CSAT papers year by year (like GS Paper I PYQs).
CSAT_PYQ_BANK = _load_json_file("csat_pyq.json", {"papers": []})
CSAT_PYQ_PAPERS = CSAT_PYQ_BANK.get("papers", [])

def seed_previous_year_papers():
    """Seed the full verified papers into the DB as ready-to-take mock tests owned
    by a system user (the YEAR-WISE view). Idempotent: each year's paper is created once."""
    data = PYQ_BANK
    if not data.get("papers"):
        return
    db = SessionLocal()
    try:
        sys_user = db.query(DBUser).filter(DBUser.email == SYSTEM_USER_EMAIL).first()
        if not sys_user:
            sys_user = DBUser(name="UPSC Official", email=SYSTEM_USER_EMAIL, hashed_password="!disabled")
            db.add(sys_user); db.commit(); db.refresh(sys_user)
        for paper in data.get("papers", []):
            paper_name = paper.get("paper", "UPSC Prelims GS Paper I")
            year = paper.get("year")
            title = f"{paper_name} — {year} (Official PYQ)"
            questions = paper.get("questions", [])
            if not questions:
                continue
            existing = db.query(DBMockTest).filter(DBMockTest.title == title).first()
            if existing:
                try:
                    q_rows = (db.query(DBQuestion)
                              .filter(DBQuestion.mock_test_id == existing.id)
                              .order_by(DBQuestion.id).all())
                    if q_rows and len(q_rows) == len(questions):
                        # Questions already seeded → refresh fields IN PLACE. We must NOT
                        # delete + re-insert: users' answers (and other tables) reference
                        # these question rows via foreign keys, so a delete raises a
                        # ForeignKeyViolation and crashes startup. An in-place update
                        # swaps the (previously third-party) explanations for our original
                        # ones — and keeps text/options/answer/subject in sync — without
                        # disturbing any of those references.
                        changed = 0
                        for row, q in zip(q_rows, questions):
                            for attr, val in (("text", q.get("text")),
                                              ("option_a", q.get("option_a")),
                                              ("option_b", q.get("option_b")),
                                              ("option_c", q.get("option_c")),
                                              ("option_d", q.get("option_d")),
                                              ("correct_answer", q.get("correct_answer")),
                                              ("explanation", q.get("explanation") or ""),
                                              ("subject", q.get("subject"))):
                                if (getattr(row, attr) or "") != (val or ""):
                                    setattr(row, attr, val); changed += 1
                        if (existing.description or "") != (paper.get("source_note") or ""):
                            existing.description = paper.get("source_note", ""); changed += 1
                        if existing.duration_minutes != YEARWISE_DURATION_MIN:
                            existing.duration_minutes = YEARWISE_DURATION_MIN; changed += 1
                        if changed:
                            db.commit()
                        continue  # in-place refresh done (or already up to date)
                    # Question COUNT changed → a genuine structural rebuild. Remove the
                    # dependent answers first to satisfy the FK, then the questions. Wrapped
                    # so a constraint issue can never fail the whole startup.
                    qids = [r.id for r in q_rows]
                    if qids:
                        db.query(DBAnswer).filter(
                            DBAnswer.question_id.in_(qids)).delete(synchronize_session=False)
                        db.query(DBQuestion).filter(
                            DBQuestion.mock_test_id == existing.id).delete(synchronize_session=False)
                    existing.description = paper.get("source_note", "")
                    existing.total_questions = len(questions)
                    existing.duration_minutes = YEARWISE_DURATION_MIN
                    db.commit()
                    db_test = existing
                except Exception as _seed_err:
                    db.rollback()
                    print(f"[seed] skipped refresh for {title}: "
                          f"{type(_seed_err).__name__}: {str(_seed_err)[:160]}")
                    continue
            else:
                db_test = DBMockTest(
                    title=title,
                    description=paper.get("source_note", ""),
                    subject="UPSC Prelims (Mixed)",
                    total_questions=len(questions),
                    duration_minutes=YEARWISE_DURATION_MIN,
                    user_id=sys_user.id,
                )
                db.add(db_test); db.commit(); db.refresh(db_test)
            for q in questions:
                db.add(DBQuestion(
                    text=q["text"], option_a=q["option_a"], option_b=q["option_b"],
                    option_c=q["option_c"], option_d=q["option_d"],
                    correct_answer=q["correct_answer"], explanation=q.get("explanation", ""),
                    subject=q.get("subject"),
                    mock_test_id=db_test.id,
                ))
            db.commit()
    finally:
        db.close()

seed_previous_year_papers()

app = FastAPI(title="IAS Mentor AI", description="AI-powered UPSC exam preparation platform", version="2.0.0")

app.add_middleware(
    CORSMiddleware, allow_origins=["*"],
    allow_credentials=True, allow_methods=["*"], allow_headers=["*"],
)

security = HTTPBearer()
security_optional = HTTPBearer(auto_error=False)

def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security), db: Session = Depends(get_db)) -> DBUser:
    payload = verify_token(credentials.credentials)
    if not payload:
        raise HTTPException(status_code=401, detail="Invalid or expired token")
    email = payload.get("sub")
    db_user = db.query(DBUser).filter(DBUser.email == email).first()
    if not db_user:
        raise HTTPException(status_code=404, detail="User not found")
    return db_user

def get_optional_user(credentials: Optional[HTTPAuthorizationCredentials] = Depends(security_optional),
                      db: Session = Depends(get_db)) -> Optional[DBUser]:
    """Like get_current_user but never raises — returns None for logged-out visitors.
    Lets public browse endpoints serve content without a login, while still
    recognising a signed-in user when a valid token is present."""
    if not credentials:
        return None
    payload = verify_token(credentials.credentials)
    if not payload:
        return None
    email = payload.get("sub")
    return db.query(DBUser).filter(DBUser.email == email).first()

def _admin_email_set(db: Session) -> set:
    """Built-in (config) admins plus any added via the admin UI (stored in DB)."""
    emails = set(ADMIN_EMAILS)
    try:
        emails |= {(a.email or "").lower() for a in db.query(DBAdminEmail).all()}
    except Exception:
        pass
    return emails

def is_admin(user: DBUser, db: Session) -> bool:
    return bool(user and (user.email or "").lower() in _admin_email_set(db))

def require_admin(current_user: DBUser = Depends(get_current_user), db: Session = Depends(get_db)) -> DBUser:
    if not is_admin(current_user, db):
        raise HTTPException(status_code=403, detail="Admin access required")
    return current_user

# ── Schemas ──────────────────────────────────────────────────────────────────
class UserRegister(BaseModel):
    name: Optional[str] = None
    email: str
    password: str
    phone: Optional[str] = None
    target_year: Optional[str] = None
    home_state: Optional[str] = None

class UserLogin(BaseModel):
    email: Optional[str] = None
    identifier: Optional[str] = None   # email OR phone number
    password: str

class ForgotPasswordRequest(BaseModel):
    email: str
    phone: str
    new_password: str

class ProfileIn(BaseModel):
    # Personal details
    full_name: Optional[str] = None
    parent_name: Optional[str] = None
    dob: Optional[str] = None
    age: Optional[int] = None
    gender: Optional[str] = None
    marital_status: Optional[str] = None
    phone: Optional[str] = None
    email: Optional[str] = None
    address: Optional[str] = None
    home_state: Optional[str] = None
    district: Optional[str] = None
    prep_location: Optional[str] = None
    prep_city: Optional[str] = None
    # Educational qualification
    education: Optional[str] = None
    graduation_stream: Optional[str] = None
    schooling_medium: Optional[str] = None
    degree_percentage: Optional[str] = None
    additional_qualification: Optional[str] = None
    # Mains exam details
    optional_subject: Optional[str] = None
    mains_language: Optional[str] = None
    medium: Optional[str] = None
    # Category
    category: Optional[str] = None
    # Working experience
    working_professional: Optional[bool] = False
    prep_intensity: Optional[str] = None
    work_experience: Optional[str] = None
    # Previous attempts
    attempts: Optional[str] = None
    failure_stage: Optional[str] = None
    coaching_status: Optional[str] = None
    failure_reason: Optional[str] = None
    # Other details to gauge level
    target_year: Optional[str] = None
    prep_level: Optional[str] = None
    knowledge_level: Optional[str] = None
    comprehension_skill: Optional[str] = None
    reading_speed: Optional[str] = None
    learning_style: Optional[str] = None
    study_hours: Optional[str] = None
    study_time_windows: Optional[str] = None
    study_place: Optional[str] = None
    strong_subjects: Optional[str] = None
    weak_subjects: Optional[str] = None

class DiagnosticStep(BaseModel):
    answered: Optional[List[dict]] = None   # [{id, selected}] accumulated by the client

class QuestionCreate(BaseModel):
    text: str
    option_a: str
    option_b: str
    option_c: str
    option_d: str
    correct_answer: str
    explanation: Optional[str] = None

class MockTestCreate(BaseModel):
    title: str
    description: Optional[str] = None
    subject: str
    topic: Optional[str] = None
    total_questions: int
    duration_minutes: int
    auto_generate: bool = False  # NEW: auto-generate questions with AI

class AnswerSubmit(BaseModel):
    question_id: int
    selected_option: str
    confidence: Optional[str] = None       # sure | unsure | guess
    time_taken: Optional[int] = None       # seconds spent on this question

class RevisionAnswer(BaseModel):
    question_id: int
    selected_option: str
    confidence: Optional[str] = None

class WrongReasonIn(BaseModel):
    reason: str                            # conceptual | factual | careless | misread | guess

class AIGenerateRequest(BaseModel):
    subject: str
    topic: str
    num_questions: int = 5

class AIExplainRequest(BaseModel):
    topic: str
    context: Optional[str] = None

class AIChatRequest(BaseModel):
    message: str

class PreviousYearRequest(BaseModel):
    subject: Optional[str] = None     # None / "" / "All" = all subjects (subject-wise filter)
    year: Optional[str] = None        # e.g. "2023"; blank/None = all years (year-wise filter)
    num_questions: Optional[int] = None  # None / 0 = ALL matching questions
    duration_minutes: Optional[int] = None  # defaults to ~1 min/question if omitted

class NcertGenerateRequest(BaseModel):
    book_key: str                       # key from /ncert/books
    chapter: str                        # chapter name within that book
    num_questions: Optional[int] = 5
    difficulty: Optional[str] = "medium"
    question_type: Optional[str] = "all"   # factual | analytical | all

class VerifiedGenerateRequest(BaseModel):
    subject: Optional[str] = None
    topic: Optional[str] = None
    book: Optional[str] = None
    num_questions: Optional[int] = 5
    difficulty: Optional[str] = "medium"
    question_type: Optional[str] = "all"   # factual | analytical | all
    reuse: Optional[bool] = False          # reuse cached verified questions for this concept

# ── Frontend + Health ─────────────────────────────────────────────────────────
FRONTEND_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "frontend", "index.html")
ADMIN_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "frontend", "admin.html")
BLUEPRINT_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "frontend", "blueprint.html")
HERO_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "frontend", "hero.jpg")

_NOCACHE = {"Cache-Control": "no-store, no-cache, must-revalidate, max-age=0",
            "Pragma": "no-cache", "Expires": "0"}

@app.get("/", include_in_schema=False)
def serve_frontend():
    """Serve the single-page web app so one deployment hosts both the UI and API."""
    if os.path.exists(FRONTEND_FILE):
        return FileResponse(FRONTEND_FILE, headers=_NOCACHE)
    return {"message": "IAS Mentor AI API is running. Frontend file not found."}

@app.get("/hero.jpg", include_in_schema=False)
def serve_hero():
    """Serve the landing hero image (cached; versioned via ?v=N query on the client)."""
    if os.path.exists(HERO_FILE):
        return FileResponse(HERO_FILE, media_type="image/jpeg",
                            headers={"Cache-Control": "public, max-age=2592000, immutable"})
    return Response(status_code=404)

@app.get("/admin", include_in_schema=False)
def serve_admin():
    """Serve the separate admin dashboard at /admin (its own login)."""
    if os.path.exists(ADMIN_FILE):
        return FileResponse(ADMIN_FILE, headers=_NOCACHE)
    return {"message": "Admin page not found."}

@app.get("/admin/blueprint", include_in_schema=False)
def serve_blueprint():
    """Serve the content-requirements dashboard (embedded in the admin as a tab).
    Static, aggregate content only — no secrets and no DB access."""
    if os.path.exists(BLUEPRINT_FILE):
        return FileResponse(BLUEPRINT_FILE, headers=_NOCACHE)
    return HTMLResponse("<h2 style='font-family:sans-serif'>Blueprint not deployed yet.</h2>",
                        status_code=404)

@app.get("/admin/blueprint/data", include_in_schema=False)
def blueprint_data():
    """Live counts for the Content Blueprint dashboard — aggregate library sizes
    only (no per-record data, no secrets). The dashboard fetches this so its
    'have' figures update themselves as the library grows."""
    import glob as _glob
    from collections import Counter as _Counter
    base = os.path.dirname(os.path.abspath(__file__))

    def _load(path, fn, default=None):
        try:
            with open(path, "r", encoding="utf-8") as f:
                return fn(json.load(f))
        except Exception:
            return default

    weight = _Counter(q.get("subject") for q in PYQ_QUESTIONS if q.get("subject"))
    weightage = [{"s": s, "v": v} for s, v in weight.most_common()]

    csat = sum(len(a.get("questions", [])) for a in CSAT_AREAS)
    diagnostic = _load(os.path.join(base, "diagnostic_bank.json"),
                       lambda d: len(d.get("items", [])), 0)
    # Concepts come LIVE from the permanent inventory the extraction pipeline
    # fills; the legacy concept_graph.json file is only a fallback.
    concepts, concepts_24h, concepts_7d, concept_subjects = 0, 0, 0, []
    try:
        from sqlalchemy import func as _func
        from datetime import datetime as _dt, timedelta as _td
        _db = SessionLocal()
        concepts = _db.query(DBConceptInventory).count()
        if concepts:
            _now = _dt.utcnow()
            concepts_24h = (_db.query(DBConceptInventory)
                            .filter(DBConceptInventory.created_at >= _now - _td(hours=24)).count())
            concepts_7d = (_db.query(DBConceptInventory)
                           .filter(DBConceptInventory.created_at >= _now - _td(days=7)).count())
            _rows = (_db.query(DBConceptInventory.subject, _func.count(DBConceptInventory.id))
                     .group_by(DBConceptInventory.subject).all())
            concept_subjects = sorted([{"s": s or "Unclassified", "v": int(v)} for s, v in _rows],
                                      key=lambda x: -x["v"])[:12]
        _db.close()
    except Exception:
        pass
    if not concepts:
        concepts = _load(os.path.join(base, "concept_graph.json"),
                         lambda d: len(d.get("concepts", [])), 0)
    omcq = _load(os.path.join(base, "content", "original_mcqs.json"),
                 lambda d: sum(len(s.get("questions", [])) for s in d.get("subjects", [])), 0)
    omains = _load(os.path.join(base, "content", "original_mains.json"),
                   lambda d: sum(len(s.get("questions", [])) for s in d.get("subjects", [])), 0)
    notes = len(_glob.glob(os.path.join(base, "content", "notes_*.md")))
    mocks = len(PYQ_BANK.get("papers", []))

    total = (len(PYQ_QUESTIONS) + len(MAINS_QUESTIONS) + csat + diagnostic
             + omcq + omains + concepts)

    return {
        "prelims_pyq": len(PYQ_QUESTIONS),
        "mains_pyq": len(MAINS_QUESTIONS),
        "csat": csat,
        "diagnostic": diagnostic,
        "concepts": concepts,
        "concepts_24h": concepts_24h,
        "concepts_7d": concepts_7d,
        "concept_subjects": concept_subjects,
        "original_mcqs": omcq,
        "original_mains": omains,
        "original_total": omcq + omains,
        "notes_subjects": notes,
        "mock_tests": mocks,
        "weightage": weightage,
        "total_items": total,
    }

@app.get("/health", tags=["Health"])
def health():
    return {"message": "IAS Mentor AI API v2.0 is running ✅"}

# ── Auth ──────────────────────────────────────────────────────────────────────
@app.post("/register", tags=["Auth"])
def register(user: UserRegister, db: Session = Depends(get_db)):
    email = (user.email or "").strip().lower()
    if not email or "@" not in email:
        raise HTTPException(status_code=400, detail="A valid email is required")
    if not (user.password or "") or len(user.password) < 4:
        raise HTTPException(status_code=400, detail="Password must be at least 4 characters")
    phone_digits = "".join(c for c in (user.phone or "") if c.isdigit())
    if len(phone_digits) < 10:
        raise HTTPException(status_code=400, detail="A valid phone number (at least 10 digits) is required")
    if db.query(DBUser).filter(DBUser.email == email).first():
        raise HTTPException(status_code=400, detail="Email already registered")
    new_user = DBUser(name=(user.name or "").strip() or None, email=email,
                      phone=phone_digits, hashed_password=hash_password(user.password))
    db.add(new_user)
    db.commit()
    db.refresh(new_user)
    # Seed a starter profile if onboarding details were provided at sign-up.
    ty = (user.target_year or "").strip()
    st = (user.home_state or "").strip()
    if ty or st:
        try:
            db.add(DBStudentProfile(user_id=new_user.id, target_year=ty or None, home_state=st or None))
            db.commit()
        except Exception:
            db.rollback()
    return {"status": "success", "message": "User registered successfully", "data": {"id": new_user.id, "email": new_user.email}}

@app.post("/login", tags=["Auth"])
def login(user: UserLogin, db: Session = Depends(get_db)):
    ident = (user.identifier or user.email or "").strip()
    db_user = None
    if ident:
        if "@" in ident:
            db_user = db.query(DBUser).filter(DBUser.email == ident.lower()).first()
        else:
            digits = "".join(c for c in ident if c.isdigit())
            if digits:
                # match on stored phone (stored as digits-only at registration)
                db_user = db.query(DBUser).filter(DBUser.phone == digits).first()
            if not db_user:
                # fall back to treating the value as an email/login string
                db_user = db.query(DBUser).filter(DBUser.email == ident.lower()).first()
    if not db_user or not verify_password(user.password, db_user.hashed_password):
        raise HTTPException(status_code=401, detail="Invalid email/phone or password")
    token = create_access_token({"sub": db_user.email})
    return {"status": "success", "message": "Login successful", "data": {"access_token": token, "token_type": "bearer", "name": db_user.name or db_user.email, "is_admin": is_admin(db_user, db)}}

@app.post("/forgot-password", tags=["Auth"])
def forgot_password(req: ForgotPasswordRequest, db: Session = Depends(get_db)):
    """Self-service reset: verify identity with the registered email + phone, then
    set a new password. (No email/SMS infra needed; phone is the verification factor.)"""
    email = (req.email or "").strip().lower()
    db_user = db.query(DBUser).filter(DBUser.email == email).first()
    if not db_user:
        raise HTTPException(status_code=404, detail="No account found with that email")
    if not db_user.phone:
        raise HTTPException(status_code=400, detail="This account has no phone on record. Please contact support to reset your password.")
    given = "".join(c for c in (req.phone or "") if c.isdigit())
    stored = "".join(c for c in (db_user.phone or "") if c.isdigit())
    # Match on the last 10 digits (handles country-code differences).
    if not given or given[-10:] != stored[-10:]:
        raise HTTPException(status_code=400, detail="Phone number does not match our records for this email")
    if not (req.new_password or "") or len(req.new_password) < 4:
        raise HTTPException(status_code=400, detail="New password must be at least 4 characters")
    db_user.hashed_password = hash_password(req.new_password)
    db.commit()
    return {"status": "success", "message": "Password reset successfully. You can now log in."}

@app.post("/__migrate_import", tags=["Admin"])
def __migrate_import(key: str = ""):
    """ONE-TIME data import: copy every table from OLD_DATABASE_URL into the current
    database. Key-gated; disable by clearing the MIGRATION_KEY env var when done."""
    expected = os.getenv("MIGRATION_KEY", "")
    if not expected or key != expected:
        raise HTTPException(status_code=403, detail="Forbidden")
    old_url = os.getenv("OLD_DATABASE_URL", "")
    if not old_url:
        raise HTTPException(status_code=400, detail="OLD_DATABASE_URL is not set")
    if old_url.startswith("postgres://"):
        old_url = old_url.replace("postgres://", "postgresql://", 1)
    from sqlalchemy import create_engine as _ce
    old_engine = _ce(old_url)
    copied, skipped = {}, {}
    tables = list(Base.metadata.sorted_tables)
    try:
        with old_engine.connect() as oconn:
            with engine.begin() as nconn:
                for t in reversed(tables):
                    try:
                        nconn.execute(sa_text(f'TRUNCATE TABLE "{t.name}" RESTART IDENTITY CASCADE'))
                    except Exception:
                        pass
                for t in tables:
                    try:
                        rows = [dict(r._mapping) for r in oconn.execute(t.select())]
                    except Exception as e:
                        skipped[t.name] = str(e)[:140]
                        continue
                    for i in range(0, len(rows), 500):
                        if rows[i:i + 500]:
                            nconn.execute(t.insert(), rows[i:i + 500])
                    copied[t.name] = len(rows)
                for t in tables:
                    if "id" in t.c:
                        try:
                            nconn.execute(sa_text(
                                f"SELECT setval(pg_get_serial_sequence('\"{t.name}\"','id'), "
                                f"GREATEST((SELECT COALESCE(MAX(id),1) FROM \"{t.name}\"),1))"))
                        except Exception:
                            pass
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Import failed: {str(e)[:300]}")
    finally:
        old_engine.dispose()
    return {"status": "success", "copied": copied, "skipped": skipped,
            "total_rows": sum(copied.values())}

@app.get("/me", tags=["Auth"])
def whoami(current_user: DBUser = Depends(get_current_user), db: Session = Depends(get_db)):
    return {"status": "success", "name": current_user.name or current_user.email,
            "email": current_user.email, "is_admin": is_admin(current_user, db)}

# ── Student Profile ─────────────────────────────────────────────────────────────
_PROFILE_FIELDS = [
    # Personal details
    "full_name", "parent_name", "dob", "age", "gender", "marital_status",
    "home_state", "district", "phone", "email", "address",
    # Educational qualification
    "education", "graduation_stream", "schooling_medium", "degree_percentage",
    "additional_qualification",
    # Mains exam details
    "optional_subject", "mains_language", "medium",
    # Category
    "category",
    # Working experience
    "working_professional", "prep_intensity", "work_experience",
    # Preparation setup
    "prep_location", "prep_city",
    # Previous attempts
    "attempts", "failure_stage", "coaching_status", "failure_reason",
    # Other details to gauge level
    "target_year", "prep_level", "knowledge_level", "comprehension_skill",
    "reading_speed", "learning_style", "study_hours", "study_time_windows",
    "study_place", "strong_subjects", "weak_subjects",
]

# Fields not counted toward the completion percentage (optional/conditional,
# or removed from the profile form so they can never be filled).
_PROFILE_OPTIONAL = {"working_professional", "prep_intensity", "work_experience",
                     "prep_city", "parent_name", "age", "failure_reason", "study_place",
                     "additional_qualification", "degree_percentage",
                     # removed from the UI — must not drag completion below 100%
                     "dob", "marital_status", "district", "address",
                     "prep_location", "coaching_status", "mains_language"}

def _profile_dict(p):
    if not p:
        return {f: ("" if f != "working_professional" else False) for f in _PROFILE_FIELDS}
    d = {f: getattr(p, f) for f in _PROFILE_FIELDS}
    for f in _PROFILE_FIELDS:
        if f != "working_professional" and d[f] is None:
            d[f] = ""
    d["working_professional"] = bool(d.get("working_professional"))
    return d

def _profile_completion(d):
    keys = [k for k in _PROFILE_FIELDS if k not in _PROFILE_OPTIONAL]
    filled = sum(1 for k in keys if str(d.get(k) or "").strip())
    return round(100 * filled / len(keys))

# The personalisation signals the daily-mission engine reads off the profile.
_MISSION_SIGNALS = ("study_hours", "prep_intensity", "working_professional",
                    "weak_subjects", "learning_style", "reading_speed",
                    "knowledge_level", "comprehension_skill", "failure_stage",
                    "failure_reason", "diagnostic_gs", "diagnostic_csat")

def _mission_profile(p):
    return {f: getattr(p, f, None) for f in _MISSION_SIGNALS} if p else {}

@app.get("/geo/profile-options", tags=["Auth"])
def geo_profile_options(current_user: Optional[DBUser] = Depends(get_optional_user)):
    """Categories, states, state->districts, optionals etc. for the DAF profile dropdowns.
    Public so a logged-out visitor can fill their profile before registering."""
    return geo_data.profile_options()

@app.get("/me/profile", tags=["Auth"])
def get_my_profile(current_user: DBUser = Depends(get_current_user), db: Session = Depends(get_db)):
    p = db.query(DBStudentProfile).filter(DBStudentProfile.user_id == current_user.id).first()
    d = _profile_dict(p)
    return {"status": "success", "profile": d, "completion": _profile_completion(d),
            "name": current_user.name or "", "email": current_user.email,
            "phone": current_user.phone or ""}

@app.put("/me/profile", tags=["Auth"])
def update_my_profile(payload: ProfileIn, current_user: DBUser = Depends(get_current_user), db: Session = Depends(get_db)):
    p = db.query(DBStudentProfile).filter(DBStudentProfile.user_id == current_user.id).first()
    if not p:
        p = DBStudentProfile(user_id=current_user.id)
        db.add(p)
    data = payload.dict()
    for f in _PROFILE_FIELDS:
        if f in data and data[f] is not None:
            setattr(p, f, data[f])
    db.commit()
    db.refresh(p)
    d = _profile_dict(p)
    return {"status": "success", "profile": d, "completion": _profile_completion(d)}


# ── Adaptive Diagnostic (objective knowledge + comprehension baseline) ──────────
def _know_label(s):
    return "Strong" if s >= 75 else ("Good" if s >= 55 else ("Moderate" if s >= 35 else "Low"))

def _csat_label(s):
    return "Strong" if s >= 70 else ("Average" if s >= 45 else "Needs improvement")

@app.post("/me/diagnostic/step", tags=["Diagnostic"])
def diagnostic_step(payload: DiagnosticStep, current_user: DBUser = Depends(get_current_user),
                    db: Session = Depends(get_db)):
    """Drive the adaptive diagnostic one step at a time. Send the answers given so
    far; get the next adapted question, or — when complete — the scored result,
    which is stored on the profile to calibrate the candidate's tailored plan."""
    answered = payload.answered or []
    q = diagnostic.next_question(answered)
    if q:
        return {"status": "success", "done": False, "question": q, "answered": len(answered)}
    res = diagnostic.score(answered)
    p = db.query(DBStudentProfile).filter(DBStudentProfile.user_id == current_user.id).first()
    if not p:
        p = DBStudentProfile(user_id=current_user.id); db.add(p)
    p.diagnostic_gs = res["gs"]; p.diagnostic_csat = res["csat"]
    # Seed the self-rating fields from the objective result if the user left them blank.
    if not (p.knowledge_level or "").strip():
        p.knowledge_level = _know_label(res["gs"])
    if not (p.comprehension_skill or "").strip():
        p.comprehension_skill = _csat_label(res["csat"])
    db.commit()
    res["knowledge_label"] = _know_label(res["gs"])
    res["comprehension_label"] = _csat_label(res["csat"])
    return {"status": "success", "done": True, "result": res}

@app.get("/me/diagnostic/result", tags=["Diagnostic"])
def diagnostic_result(current_user: DBUser = Depends(get_current_user), db: Session = Depends(get_db)):
    p = db.query(DBStudentProfile).filter(DBStudentProfile.user_id == current_user.id).first()
    gs = p.diagnostic_gs if p else None
    csat = p.diagnostic_csat if p else None
    return {"status": "success", "taken": gs is not None,
            "gs": gs, "csat": csat,
            "knowledge_label": (_know_label(gs) if gs is not None else None),
            "comprehension_label": (_csat_label(csat) if csat is not None else None)}

# ── Syllabus Tracker ──────────────────────────────────────────────────────────
class SyllabusToggle(BaseModel):
    topic_id: str
    completed: bool = True

@app.get("/me/syllabus", tags=["Syllabus"])
def get_my_syllabus(current_user: Optional[DBUser] = Depends(get_optional_user), db: Session = Depends(get_db)):
    """Full UPSC syllabus tree + this user's completed topic IDs and progress.
    Public browse: logged-out visitors get the full tree with no completed topics;
    ticking a topic (POST /me/syllabus/toggle) still requires an account."""
    valid = syllabus_tracker_data.all_topic_ids()
    if current_user is None:
        return {"status": "success", "tree": syllabus_tracker_data.tree_with_ids(),
                "completed": [], "total": len(valid), "completed_count": 0}
    done = {r.topic_id for r in db.query(DBSyllabusProgress)
            .filter(DBSyllabusProgress.user_id == current_user.id).all()}
    done = {t for t in done if t in valid}
    return {"status": "success", "tree": syllabus_tracker_data.tree_with_ids(),
            "completed": sorted(done), "total": len(valid), "completed_count": len(done)}

@app.post("/me/syllabus/toggle", tags=["Syllabus"])
def toggle_syllabus_topic(payload: SyllabusToggle,
                          current_user: DBUser = Depends(get_current_user),
                          db: Session = Depends(get_db)):
    tid = (payload.topic_id or "").strip()
    if tid not in syllabus_tracker_data.all_topic_ids():
        raise HTTPException(status_code=400, detail="Unknown topic")
    row = db.query(DBSyllabusProgress).filter(
        DBSyllabusProgress.user_id == current_user.id,
        DBSyllabusProgress.topic_id == tid).first()
    if payload.completed and not row:
        db.add(DBSyllabusProgress(user_id=current_user.id, topic_id=tid))
    elif not payload.completed and row:
        db.delete(row)
    db.commit()
    total = db.query(DBSyllabusProgress).filter(DBSyllabusProgress.user_id == current_user.id).count()
    return {"status": "success", "completed": payload.completed,
            "completed_count": total, "total": syllabus_tracker_data.total_topics()}

@app.post("/me/syllabus/reset", tags=["Syllabus"])
def reset_syllabus(current_user: DBUser = Depends(get_current_user), db: Session = Depends(get_db)):
    db.query(DBSyllabusProgress).filter(DBSyllabusProgress.user_id == current_user.id).delete()
    db.commit()
    return {"status": "success", "completed_count": 0, "total": syllabus_tracker_data.total_topics()}

# ── Study Planner ─────────────────────────────────────────────────────────────
@app.get("/me/study-plan", tags=["Planner"])
def my_study_plan(target_year: Optional[str] = None, hours: Optional[str] = None,
                  current_user: DBUser = Depends(get_current_user), db: Session = Depends(get_db)):
    """Generate a phase-wise AIVORA study schedule for the time left until Prelims.
    Pulls defaults from the student profile; query params override them."""
    p = db.query(DBStudentProfile).filter(DBStudentProfile.user_id == current_user.id).first()
    ty = target_year or (p.target_year if p else None) or ""
    hb = hours or (p.study_hours if p else None) or "2-4"
    weak = (p.weak_subjects if p else None) or ""
    plan = study_planner.generate_plan(ty, hb, weak)
    plan["from_profile"] = {"target_year": (p.target_year if p else None),
                            "study_hours": (p.study_hours if p else None),
                            "weak_subjects": weak}
    return {"status": "success", "plan": plan}

# ── Guided Program (follow-along curriculum) ──────────────────────────────────
def _program_inputs(p):
    ty = (p.target_year if p else None) or ""
    hb = (p.study_hours if p else None) or "2-4"
    weak = (p.weak_subjects if p else None) or ""
    return ty, hb, weak

class GuidedToggle(BaseModel):
    task_id: str
    completed: bool = True

@app.get("/me/program", tags=["Planner"])
def my_program(target_year: Optional[str] = None, hours: Optional[str] = None,
               current_user: DBUser = Depends(get_current_user), db: Session = Depends(get_db)):
    """The full guided task curriculum + the user's completed task IDs and progress."""
    p = db.query(DBStudentProfile).filter(DBStudentProfile.user_id == current_user.id).first()
    ty, hb, weak = _program_inputs(p)
    if target_year:
        ty = target_year
    if hours:
        hb = hours
    prog = guided_program.generate_program(ty, hb, weak)
    done = {r.task_id for r in db.query(DBGuidedProgress)
            .filter(DBGuidedProgress.user_id == current_user.id).all()}
    valid = {t["id"] for d in prog["days"] for t in d["tasks"]}
    done = {t for t in done if t in valid}
    prog["completed"] = sorted(done)
    prog["completed_count"] = len(done)
    return {"status": "success", "program": prog}

@app.post("/me/program/toggle", tags=["Planner"])
def toggle_program_task(payload: GuidedToggle,
                        current_user: DBUser = Depends(get_current_user),
                        db: Session = Depends(get_db)):
    p = db.query(DBStudentProfile).filter(DBStudentProfile.user_id == current_user.id).first()
    ty, hb, weak = _program_inputs(p)
    tid = (payload.task_id or "").strip()
    if tid not in guided_program.all_task_ids(ty, hb, weak):
        raise HTTPException(status_code=400, detail="Unknown task")
    row = db.query(DBGuidedProgress).filter(
        DBGuidedProgress.user_id == current_user.id,
        DBGuidedProgress.task_id == tid).first()
    if payload.completed and not row:
        db.add(DBGuidedProgress(user_id=current_user.id, task_id=tid))
    elif not payload.completed and row:
        db.delete(row)
    db.commit()
    total = db.query(DBGuidedProgress).filter(DBGuidedProgress.user_id == current_user.id).count()
    return {"status": "success", "completed": payload.completed, "completed_count": total}

# ── PrepOS "Today" — the daily decision engine ────────────────────────────────
def _gather_answers(db, user_id):
    rows = (db.query(DBAnswer.is_correct, DBAnswer.time_taken, DBQuestion.subject,
                     DBQuestion.topic, DBQuestion.difficulty, DBTestAttempt.completed_at,
                     DBAnswer.confidence)
            .join(DBTestAttempt, DBAnswer.test_attempt_id == DBTestAttempt.id)
            .join(DBQuestion, DBAnswer.question_id == DBQuestion.id)
            .filter(DBTestAttempt.user_id == user_id)
            .order_by(DBTestAttempt.completed_at.asc()).all())
    return [{"is_correct": bool(r[0]), "time_taken": r[1], "subject": r[2],
             "topic": r[3], "difficulty": r[4], "completed_at": r[5],
             "confidence": r[6]} for r in rows]

def _next_guided_task(db, user, profile):
    ty = (profile.target_year if profile else None) or ""
    hb = (profile.study_hours if profile else None) or "2-4"
    weak = (profile.weak_subjects if profile else None) or ""
    prog = guided_program.generate_program(ty, hb, weak)
    done = {r.task_id for r in db.query(DBGuidedProgress)
            .filter(DBGuidedProgress.user_id == user.id).all()}
    for d in prog["days"]:
        for t in d["tasks"]:
            if t["id"] not in done and t["kind"] in ("read", "ncert_test", "subject_test", "pyq"):
                return {"id": t["id"], "kind": t["kind"], "title": t["title"], "params": t["params"]}
    return None

@app.get("/me/today", tags=["Planner"])
def my_today(current_user: DBUser = Depends(get_current_user), db: Session = Depends(get_db)):
    """The full PrepOS state: knowledge map, scorecard, forecast, today's mission, interventions."""
    today = datetime.date.today()
    profile = db.query(DBStudentProfile).filter(DBStudentProfile.user_id == current_user.id).first()
    answers = _gather_answers(db, current_user.id)
    review_items = [{"times_seen": r.times_seen, "times_correct": r.times_correct,
                     "mastered": r.mastered, "next_review": r.next_review}
                    for r in db.query(DBReviewItem).filter(DBReviewItem.user_id == current_user.id).all()]
    # Exclude Teaching-Engine "mission check" attempts from the mocks metric — their
    # answers still count toward mastery (via _gather_answers), but a 3-Q check isn't
    # a full mock. outerjoin keeps attempts whose mock test is missing/null.
    attempts = [{"completed_at": a.completed_at, "score": a.score}
                for a in (db.query(DBTestAttempt)
                          .outerjoin(DBMockTest, DBTestAttempt.mock_test_id == DBMockTest.id)
                          .filter(DBTestAttempt.user_id == current_user.id)
                          .filter((DBMockTest.title.is_(None)) | (~DBMockTest.title.like("🔎 Mission check%")))
                          .all())]
    now = datetime.datetime.utcnow()
    review_due = db.query(DBReviewItem).filter(
        DBReviewItem.user_id == current_user.id,
        DBReviewItem.next_review <= now,
        DBReviewItem.mastered == False).count()  # noqa: E712
    syl_done = db.query(DBSyllabusProgress).filter(DBSyllabusProgress.user_id == current_user.id).count()
    coverage_pct = round(100 * syl_done / max(1, syllabus_tracker_data.total_topics()))

    km = prepos.build_knowledge_map(answers, today)
    mains_evals = db.query(DBMainsAnswer.overall_pct).filter(
        DBMainsAnswer.user_id == current_user.id).order_by(DBMainsAnswer.id.desc()).limit(10).all()
    writing_avg = (sum(m[0] for m in mains_evals) / len(mains_evals)) if mains_evals else None
    scores = prepos.compute_scores(answers, review_items, attempts, coverage_pct, today, writing_avg)
    ty_label = "UPSC Prelims"; days_left = None
    try:
        _plan = study_planner.generate_plan(
            (profile.target_year if profile else None) or "",
            (profile.study_hours if profile else None) or "2-4",
            (profile.weak_subjects if profile else None) or "")
        ty_label = _plan["exam_label"]; days_left = _plan.get("days_left")
    except Exception:
        pass
    fcast = prepos.forecast(scores, ty_label)
    nxt = _next_guided_task(db, current_user, profile)
    hb = (profile.study_hours if profile else None) or "2-4"
    mp = _mission_profile(profile)
    mission = prepos.daily_mission(km, review_due, hb, nxt, today, mp)
    lever = prepos.growth_lever(scores, mp)
    inter = prepos.interventions(answers, km, attempts, today)
    checkins = prepos.checkins(mp, scores, km, attempts, answers, review_due, today)
    report = prepos.measurement_report(scores, km, lever, mp, inter)
    pred = prepos.prediction(scores, fcast, lever, days_left)
    # Record one progress snapshot per day so movement/deltas can be shown over time.
    try:
        _hist = json.loads(profile.progress_history) if (profile and profile.progress_history) else []
    except Exception:
        _hist = []
    if not any(h.get("d") == today.isoformat() for h in _hist):
        _hist.append(prepos.snapshot_row(scores, pred, today))
        _hist = _hist[-40:]
        if profile:
            profile.progress_history = json.dumps(_hist)
            db.commit()
    prog = prepos.progress(scores, pred, _hist, review_items, attempts, answers,
                           syl_done, coverage_pct, days_left, ty_label,
                           (current_user.name or current_user.email.split("@")[0]), today)
    decision = prepos.decide(scores, lever, review_due, km, mp, pred)
    brief = prepos.briefing(current_user.name or current_user.email.split("@")[0], scores, mission, datetime.datetime.now())

    done_keys = {r.task_key for r in db.query(DBDailyMissionDone).filter(
        DBDailyMissionDone.user_id == current_user.id,
        DBDailyMissionDone.day == today.isoformat()).all()}
    for t in mission["tasks"]:
        t["done"] = t["key"] in done_keys

    return {"status": "success", "today": {
        "briefing": brief, "mission": mission, "scores": scores,
        "forecast": fcast, "knowledge_map": km, "interventions": inter,
        "exam_label": ty_label, "coverage": coverage_pct, "checkins": checkins,
        "growth_lever": lever, "report": report, "prediction": pred, "decision": decision,
        "placement": {"gs": (profile.diagnostic_gs if profile else None),
                      "csat": (profile.diagnostic_csat if profile else None)},
        "progress": prog,
    }}

class TodayDone(BaseModel):
    task_key: str
    completed: bool = True

@app.post("/me/today/done", tags=["Planner"])
def today_done(payload: TodayDone, current_user: DBUser = Depends(get_current_user),
               db: Session = Depends(get_db)):
    day = datetime.date.today().isoformat()
    row = db.query(DBDailyMissionDone).filter(
        DBDailyMissionDone.user_id == current_user.id,
        DBDailyMissionDone.day == day,
        DBDailyMissionDone.task_key == payload.task_key).first()
    if payload.completed and not row:
        db.add(DBDailyMissionDone(user_id=current_user.id, day=day, task_key=payload.task_key))
    elif not payload.completed and row:
        db.delete(row)
    db.commit()
    return {"status": "success", "completed": payload.completed}

# ── Mains answer writing + AI evaluation (feeds writing quality) ──────────────
class MainsEvalRequest(BaseModel):
    question: str
    answer: str
    marks: Optional[int] = 10
    paper: Optional[str] = None

@app.post("/me/mains/evaluate", tags=["Mains PYQ"])
def mains_evaluate(req: MainsEvalRequest, current_user: DBUser = Depends(get_current_user),
                   db: Session = Depends(get_db)):
    q = (req.question or "").strip()
    a = (req.answer or "").strip()
    if not q or len(a) < 30:
        raise HTTPException(status_code=400, detail="Please write at least a few sentences for evaluation.")
    marks = req.marks or 10
    ev = evaluate_mains_answer(q, a, marks)
    row = DBMainsAnswer(user_id=current_user.id, question=q[:2000], answer=a[:8000],
                        paper=(req.paper or "")[:60], overall_pct=ev.get("overall_pct", 0),
                        overall_marks=str(ev.get("overall_marks", "")), marks=marks,
                        eval_json=json.dumps(ev))
    db.add(row); db.commit(); db.refresh(row)
    return {"status": "success", "id": row.id, "evaluation": ev}

@app.get("/me/mains/history", tags=["Mains PYQ"])
def mains_history(current_user: DBUser = Depends(get_current_user), db: Session = Depends(get_db)):
    rows = (db.query(DBMainsAnswer).filter(DBMainsAnswer.user_id == current_user.id)
            .order_by(DBMainsAnswer.id.desc()).limit(30).all())
    out = []
    for r in rows:
        try:
            ev = json.loads(r.eval_json) if r.eval_json else {}
        except Exception:
            ev = {}
        out.append({"id": r.id, "question": r.question, "paper": r.paper,
                    "overall_pct": r.overall_pct, "overall_marks": r.overall_marks,
                    "marks": r.marks, "evaluation": ev,
                    "created_at": r.created_at.isoformat() if r.created_at else None})
    avg = round(sum(r.overall_pct for r in rows) / len(rows)) if rows else None
    return {"status": "success", "answers": out, "count": len(out), "writing_quality": avg}

# ── Weekly mentor report + trends ─────────────────────────────────────────────
@app.get("/me/report", tags=["Planner"])
def my_report(current_user: DBUser = Depends(get_current_user), db: Session = Depends(get_db)):
    today = datetime.date.today()
    answers = _gather_answers(db, current_user.id)
    attempts = [{"completed_at": a.completed_at, "score": a.score}
                for a in db.query(DBTestAttempt).filter(DBTestAttempt.user_id == current_user.id).all()]
    report = prepos.weekly_report(answers, attempts, today)
    name = current_user.name or (current_user.email or "Aspirant").split("@")[0]
    narrative = weekly_mentor_narrative(report["summary"], name) if report["has_data"] else (
        "Welcome to AIVORA. Take a few tests this week and your first weekly report — with trends and a "
        "personal note from your AI mentor — will appear right here.")
    report["narrative"] = narrative
    return {"status": "success", "report": report}

# ── Mock Tests ────────────────────────────────────────────────────────────────
@app.post("/mock-tests/", tags=["Mock Tests"])
def create_mock_test(mock_test: MockTestCreate, db: Session = Depends(get_db), current_user: DBUser = Depends(get_current_user)):
    db_test = DBMockTest(
        title=mock_test.title, description=mock_test.description,
        subject=mock_test.subject, total_questions=mock_test.total_questions,
        duration_minutes=mock_test.duration_minutes, user_id=current_user.id,
    )
    db.add(db_test)
    db.commit()
    db.refresh(db_test)

    questions_added = 0
    if mock_test.auto_generate:
        topic = mock_test.topic or mock_test.subject
        try:
            parsed = generate_and_parse_questions(mock_test.subject, topic, mock_test.total_questions)
            for q in parsed:
                db_q = DBQuestion(
                    text=q['text'], option_a=q['option_a'], option_b=q['option_b'],
                    option_c=q['option_c'], option_d=q['option_d'],
                    correct_answer=q['correct_answer'], explanation=q['explanation'],
                    mock_test_id=db_test.id,
                )
                db.add(db_q)
            db.commit()
            questions_added = len(parsed)
        except Exception as e:
            pass  # Test created, questions will need to be added manually

    return {
        "status": "success",
        "message": f"Mock test created{'with ' + str(questions_added) + ' AI questions' if questions_added else ''}",
        "mock_test_id": db_test.id,
        "questions_added": questions_added,
    }

@app.get("/previous-year-papers/", tags=["Mock Tests"])
def list_previous_year_papers(db: Session = Depends(get_db), current_user: Optional[DBUser] = Depends(get_optional_user)):
    """List the real, verified previous-year papers seeded from the question bank."""
    sys_user = db.query(DBUser).filter(DBUser.email == SYSTEM_USER_EMAIL).first()
    if not sys_user:
        return {"status": "success", "papers": []}
    tests = db.query(DBMockTest).filter(DBMockTest.user_id == sys_user.id).order_by(DBMockTest.title.desc()).all()
    return {
        "status": "success",
        "papers": [
            {
                "id": t.id, "title": t.title, "description": t.description,
                "total_questions": t.total_questions, "duration_minutes": t.duration_minutes,
                "questions_added": db.query(DBQuestion).filter(DBQuestion.mock_test_id == t.id).count(),
            }
            for t in tests
        ],
    }

@app.get("/previous-year-papers/subjects/", tags=["Mock Tests"])
def previous_year_subjects(current_user: Optional[DBUser] = Depends(get_optional_user)):
    """SUBJECT-WISE index of the verified bank: how many real questions exist per
    subject (and per year), so the UI can offer subject-wise practice sets."""
    by_subject = {}
    for q in PYQ_QUESTIONS:
        s = q.get("subject")
        if not s:
            continue
        entry = by_subject.setdefault(s, {"subject": s, "count": 0, "years": {}})
        entry["count"] += 1
        yr = q.get("year")
        if yr:
            entry["years"][yr] = entry["years"].get(yr, 0) + 1
    subjects = [
        {"subject": e["subject"], "count": e["count"],
         "years": sorted(e["years"].keys(), reverse=True)}
        for e in sorted(by_subject.values(), key=lambda x: x["subject"])
    ]
    return {
        "status": "success",
        "subjects": subjects,
        "years": PYQ_YEARS,
        "total_questions": len(PYQ_QUESTIONS),
    }

@app.post("/mock-tests/previous-year/", tags=["Mock Tests"])
def create_previous_year_test(request: PreviousYearRequest, db: Session = Depends(get_db), current_user: DBUser = Depends(get_current_user)):
    """Build a ready-to-take mock test from the VERIFIED previous-year bank,
    filtered by subject and/or year. Real exam questions only — no AI generation."""
    subject = (request.subject or "").strip()
    year = (request.year or "").strip()
    all_subjects = (not subject) or subject.lower() in ("all", "all subjects")

    # Filter the verified bank by the requested subject and/or year.
    pool = PYQ_QUESTIONS
    if not all_subjects:
        pool = [q for q in pool if q.get("subject") == subject]
    if year:
        pool = [q for q in pool if str(q.get("year")) == year]

    if not pool:
        raise HTTPException(
            status_code=404,
            detail="No verified previous-year questions match that subject/year yet.",
        )

    # No count (subject-wise practice) => use ALL matching questions; otherwise cap.
    if not request.num_questions or request.num_questions <= 0:
        selected = random.sample(pool, len(pool))
    else:
        num = max(1, min(request.num_questions, 250))
        selected = random.sample(pool, min(num, len(pool)))

    label_subject = "All Subjects" if all_subjects else subject
    duration = request.duration_minutes if request.duration_minutes else max(10, len(selected))
    title = f"PYQ • {label_subject}" + (f" • {year}" if year else "")

    db_test = DBMockTest(
        title=title,
        description=(
            f"Verified previous-year UPSC questions — {label_subject}"
            + (f" ({year})" if year else " (all years)")
        ),
        subject=label_subject,
        total_questions=len(selected),
        duration_minutes=duration,
        user_id=current_user.id,
    )
    db.add(db_test)
    db.commit()
    db.refresh(db_test)

    for q in selected:
        db.add(DBQuestion(
            text=q["text"], option_a=q["option_a"], option_b=q["option_b"],
            option_c=q["option_c"], option_d=q["option_d"],
            correct_answer=q["correct_answer"], explanation=q.get("explanation", ""),
            subject=q.get("subject"),
            mock_test_id=db_test.id,
        ))
    db.commit()

    return {
        "status": "success",
        "message": f"Verified previous-year test created with {len(selected)} questions",
        "mock_test_id": db_test.id,
        "questions_added": len(selected),
        "available": len(pool),
        "duration_minutes": duration,
    }

# ── Mains PYQs (descriptive, browse-only) ─────────────────────────────────────
@app.get("/pyq/mains/index", tags=["Mains PYQ"])
def mains_index(current_user: Optional[DBUser] = Depends(get_optional_user)):
    """Available years, papers and subject themes for the Mains PYQ browser."""
    return {"status": "success", "years": MAINS_YEARS, "papers": MAINS_PAPER_DEFS,
            "subjects": MAINS_SUBJECTS, "total_questions": len(MAINS_QUESTIONS)}

@app.get("/pyq/mains/questions", tags=["Mains PYQ"])
def mains_questions(year: Optional[str] = None, paper_code: Optional[str] = None,
                    subject: Optional[str] = None,
                    current_user: Optional[DBUser] = Depends(get_optional_user)):
    """Filtered list of real descriptive UPSC Mains questions (read/practice only)."""
    pool = MAINS_QUESTIONS
    if year:
        pool = [q for q in pool if str(q.get("year")) == str(year)]
    if paper_code:
        pool = [q for q in pool if (q.get("paper_code") or "").upper() == paper_code.upper()]
    if subject and subject.lower() not in ("all", "all subjects"):
        pool = [q for q in pool if q.get("subject") == subject]
    # Sort newest year first, then by paper order
    order = {c: i for i, c in enumerate(_MAINS_PAPER_ORDER)}
    items = sorted(pool, key=lambda q: (-(q.get("year") or 0), order.get(q.get("paper_code"), 9)))
    out = [{"q": q.get("q"), "subject": q.get("subject"), "marks": q.get("marks"),
            "words": q.get("words"), "year": q.get("year"), "paper": q.get("paper"),
            "paper_code": q.get("paper_code")} for q in items]
    return {"status": "success", "count": len(out), "questions": out}

# ── CSAT practice (original AIVORA sets, attemptable as MCQs) ──────────────────
@app.get("/pyq/csat/areas", tags=["CSAT"])
def csat_areas(current_user: Optional[DBUser] = Depends(get_optional_user)):
    return {"status": "success",
            "areas": [{"code": a["code"], "name": a["name"],
                       "count": len(a.get("questions", []))} for a in CSAT_AREAS]}

class CSATStart(BaseModel):
    area: Optional[str] = None

@app.post("/pyq/csat/start", tags=["CSAT"])
def csat_start(req: CSATStart, db: Session = Depends(get_db),
               current_user: DBUser = Depends(get_current_user)):
    """Build an attemptable mock test from an AIVORA CSAT practice area."""
    area = next((a for a in CSAT_AREAS if a["code"] == (req.area or "")), None)
    if not area:
        raise HTTPException(status_code=404, detail="Unknown CSAT area")
    qs = area.get("questions", [])
    if not qs:
        raise HTTPException(status_code=404, detail="No questions in this area yet")
    title = f"CSAT Practice • {area['name']}"
    db_test = DBMockTest(title=title, description="AIVORA CSAT practice set (original, exam-pattern)",
                         subject="CSAT", total_questions=len(qs),
                         duration_minutes=max(10, len(qs) * 2), user_id=current_user.id)
    db.add(db_test); db.commit(); db.refresh(db_test)
    for q in qs:
        db.add(DBQuestion(
            text=q["text"], option_a=q["option_a"], option_b=q["option_b"],
            option_c=q["option_c"], option_d=q["option_d"],
            correct_answer=q["correct_answer"], explanation=q.get("explanation", ""),
            subject="CSAT", topic=area["name"], difficulty="medium",
            question_type="csat", mock_test_id=db_test.id))
    db.commit()
    return {"status": "success", "mock_test_id": db_test.id,
            "questions_added": len(qs), "duration_minutes": db_test.duration_minutes,
            "title": title}


def _csat_db_papers(db):
    """Admin-authored CSAT papers from the DB → {year: paper_dict}."""
    out = {}
    for row in db.query(DBCsatPyqPaper).all():
        try:
            d = json.loads(row.data or "{}")
        except Exception:
            d = {}
        out[int(row.year)] = {
            "year": int(row.year),
            "paper": row.title or "UPSC Prelims CSAT (GS Paper II)",
            "passages": d.get("passages", []),
            "questions": d.get("questions", []),
        }
    return out


@app.get("/pyq/csat-years", tags=["CSAT"])
def csat_pyq_years(db: Session = Depends(get_db), current_user: Optional[DBUser] = Depends(get_optional_user)):
    """List CSAT (Paper II) papers available, newest first. An admin-authored paper
    (full paper, with passages) takes precedence over the built-in aptitude bank."""
    by_year = {}
    for p in CSAT_PYQ_PAPERS:
        if p.get("year"):
            by_year[int(p["year"])] = {"year": int(p["year"]),
                                       "count": len(p.get("questions", [])), "kind": "aptitude"}
    for y, p in _csat_db_papers(db).items():
        by_year[y] = {"year": y, "count": len(p.get("questions", [])), "kind": "full"}
    out = sorted(by_year.values(), key=lambda x: x["year"], reverse=True)
    return {"status": "success", "years": out}


class CSATYearStart(BaseModel):
    year: int


@app.post("/pyq/csat-year/start", tags=["CSAT"])
def csat_pyq_year_start(req: CSATYearStart, db: Session = Depends(get_db),
                        current_user: DBUser = Depends(get_current_user)):
    """Build an attemptable CSAT paper for a year. An admin-authored paper (with
    reading passages) is used when present; otherwise the built-in aptitude bank.
    Passages are shown once, above the first question of each passage group."""
    yr = int(req.year)
    paper = _csat_db_papers(db).get(yr) or next(
        (p for p in CSAT_PYQ_PAPERS if int(p.get("year", 0)) == yr), None)
    if not paper:
        raise HTTPException(status_code=404, detail="That CSAT year isn't available yet.")
    qs = paper.get("questions", [])
    if not qs:
        raise HTTPException(status_code=404, detail="No questions for that year.")
    passages = {p.get("id"): p.get("text", "") for p in paper.get("passages", [])}
    title = f"CSAT PYQ • {yr}"
    db_test = DBMockTest(
        title=title, description=f"UPSC CSAT Paper II — {yr} (PYQ)",
        subject="CSAT", total_questions=len(qs),
        duration_minutes=int(paper.get("duration_minutes") or 120), user_id=current_user.id,
    )
    db.add(db_test); db.commit(); db.refresh(db_test)
    prev_pid = None
    for q in qs:
        pid = q.get("passage_id")
        text = q.get("text", "")
        # Prepend the passage only on the FIRST question of each passage group.
        if pid and pid != prev_pid and passages.get(pid):
            text = ("Directions: Read the following passage and answer the item(s) that "
                    "follow. Your answer should be based only on the passage.\n\nPASSAGE:\n"
                    + passages[pid] + "\n\n" + text)
        prev_pid = pid
        db.add(DBQuestion(
            text=text, option_a=q.get("option_a"), option_b=q.get("option_b"),
            option_c=q.get("option_c"), option_d=q.get("option_d"),
            correct_answer=(q.get("correct_answer") or "").upper(),
            explanation=q.get("explanation", ""), subject="CSAT",
            topic=q.get("type") or "CSAT", difficulty="medium",
            question_type="csat", mock_test_id=db_test.id,
        ))
    db.commit()
    return {"status": "success", "mock_test_id": db_test.id, "questions_added": len(qs),
            "duration_minutes": db_test.duration_minutes, "title": title}


# ── Admin: CSAT Paper Builder (site owner authors full papers into the DB) ──────
class CSATPassageIn(BaseModel):
    id: str
    text: str

class CSATQuestionIn(BaseModel):
    q_no: int
    passage_id: Optional[str] = None
    text: str
    option_a: str
    option_b: str
    option_c: str
    option_d: str
    correct_answer: str
    type: Optional[str] = "general"

class CSATPaperIn(BaseModel):
    year: int
    title: Optional[str] = None
    passages: List[CSATPassageIn] = []
    questions: List[CSATQuestionIn] = []

@app.post("/admin/csat-pyq/save", tags=["CSAT"])
def admin_csat_pyq_save(req: CSATPaperIn, admin: DBUser = Depends(require_admin),
                        db: Session = Depends(get_db)):
    """Create or replace the CSAT paper for a year with admin-authored content."""
    if not req.questions:
        raise HTTPException(status_code=400, detail="Add at least one question.")
    data = json.dumps({
        "passages": [{"id": p.id, "text": p.text} for p in req.passages],
        "questions": [{"q_no": q.q_no, "passage_id": q.passage_id, "text": q.text,
                       "option_a": q.option_a, "option_b": q.option_b,
                       "option_c": q.option_c, "option_d": q.option_d,
                       "correct_answer": (q.correct_answer or "").upper(),
                       "type": q.type or "general"} for q in req.questions],
    }, ensure_ascii=False)
    row = db.query(DBCsatPyqPaper).filter(DBCsatPyqPaper.year == int(req.year)).first()
    if row:
        if req.title:
            row.title = req.title
        row.data = data
    else:
        db.add(DBCsatPyqPaper(year=int(req.year),
                              title=req.title or "UPSC Prelims CSAT (GS Paper II)", data=data))
    db.commit()
    return {"status": "success", "year": int(req.year),
            "questions": len(req.questions), "passages": len(req.passages)}

@app.get("/admin/csat-pyq", tags=["CSAT"])
def admin_csat_pyq_list(admin: DBUser = Depends(require_admin), db: Session = Depends(get_db)):
    rows = db.query(DBCsatPyqPaper).order_by(DBCsatPyqPaper.year.desc()).all()
    out = []
    for r in rows:
        try:
            d = json.loads(r.data or "{}")
        except Exception:
            d = {}
        out.append({"year": r.year, "title": r.title,
                    "questions": len(d.get("questions", [])),
                    "passages": len(d.get("passages", []))})
    return {"status": "success", "papers": out}

@app.get("/admin/csat-pyq/{year}", tags=["CSAT"])
def admin_csat_pyq_get(year: int, admin: DBUser = Depends(require_admin),
                       db: Session = Depends(get_db)):
    """Fetch an admin-authored CSAT paper for editing (empty shell if none yet)."""
    row = db.query(DBCsatPyqPaper).filter(DBCsatPyqPaper.year == int(year)).first()
    if not row:
        apt = next((p for p in CSAT_PYQ_PAPERS if int(p.get("year", 0)) == int(year)), None)
        q_apt = []
        if apt:
            for q in apt.get("questions", []):
                q_apt.append({"q_no": q.get("q_no"), "passage_id": None,
                              "text": q.get("text", ""),
                              "option_a": q.get("option_a", ""), "option_b": q.get("option_b", ""),
                              "option_c": q.get("option_c", ""), "option_d": q.get("option_d", ""),
                              "correct_answer": (q.get("correct_answer") or "").upper(),
                              "type": q.get("type", "general")})
        return {"status": "success", "year": int(year), "title": None,
                "passages": [], "questions": q_apt}
    try:
        d = json.loads(row.data or "{}")
    except Exception:
        d = {}
    return {"status": "success", "year": int(year), "title": row.title,
            "passages": d.get("passages", []), "questions": d.get("questions", [])}

@app.get("/mock-tests/", tags=["Mock Tests"])
def get_all_mock_tests(db: Session = Depends(get_db), current_user: DBUser = Depends(get_current_user)):
    tests = db.query(DBMockTest).all()
    return {
        "status": "success",
        "mock_tests": [
            {
                "id": t.id, "title": t.title, "subject": t.subject,
                "total_questions": t.total_questions, "duration_minutes": t.duration_minutes,
                "questions_added": db.query(DBQuestion).filter(DBQuestion.mock_test_id == t.id).count(),
                "attempts_count": db.query(DBTestAttempt).filter(DBTestAttempt.mock_test_id == t.id).count(),
            }
            for t in tests
        ],
    }

@app.get("/mock-tests/{mock_test_id}/", tags=["Mock Tests"])
def get_mock_test(mock_test_id: int, db: Session = Depends(get_db), current_user: DBUser = Depends(get_current_user)):
    db_test = db.query(DBMockTest).filter(DBMockTest.id == mock_test_id).first()
    if not db_test:
        raise HTTPException(status_code=404, detail="Mock test not found")
    question_count = db.query(DBQuestion).filter(DBQuestion.mock_test_id == mock_test_id).count()
    return {
        "status": "success",
        "mock_test": {
            "id": db_test.id, "title": db_test.title, "description": db_test.description,
            "subject": db_test.subject, "total_questions": db_test.total_questions,
            "duration_minutes": db_test.duration_minutes, "questions_added": question_count,
        },
    }

@app.post("/mock-tests/{mock_test_id}/questions/", tags=["Mock Tests"])
def add_question(mock_test_id: int, question: QuestionCreate, db: Session = Depends(get_db), current_user: DBUser = Depends(get_current_user)):
    if not db.query(DBMockTest).filter(DBMockTest.id == mock_test_id).first():
        raise HTTPException(status_code=404, detail="Mock test not found")
    if question.correct_answer.upper() not in ("A", "B", "C", "D"):
        raise HTTPException(status_code=422, detail="correct_answer must be A, B, C, or D")
    db_q = DBQuestion(
        text=question.text, option_a=question.option_a, option_b=question.option_b,
        option_c=question.option_c, option_d=question.option_d,
        correct_answer=question.correct_answer.upper(), explanation=question.explanation,
        mock_test_id=mock_test_id,
    )
    db.add(db_q)
    db.commit()
    db.refresh(db_q)
    return {"status": "success", "message": "Question added", "question_id": db_q.id}

@app.post("/mock-tests/{mock_test_id}/ai-populate/", tags=["Mock Tests"])
def ai_populate_test(mock_test_id: int, topic: str, db: Session = Depends(get_db), current_user: DBUser = Depends(get_current_user)):
    """Generate and save AI questions directly into an existing test."""
    db_test = db.query(DBMockTest).filter(DBMockTest.id == mock_test_id).first()
    if not db_test:
        raise HTTPException(status_code=404, detail="Mock test not found")
    try:
        parsed = generate_and_parse_questions(db_test.subject, topic, db_test.total_questions)
        for q in parsed:
            db_q = DBQuestion(
                text=q['text'], option_a=q['option_a'], option_b=q['option_b'],
                option_c=q['option_c'], option_d=q['option_d'],
                correct_answer=q['correct_answer'], explanation=q['explanation'],
                mock_test_id=mock_test_id,
            )
            db.add(db_q)
        db.commit()
        return {"status": "success", "questions_added": len(parsed)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI Error: {str(e)}")

@app.get("/mock-tests/{mock_test_id}/questions/", tags=["Mock Tests"])
def get_mock_test_questions(mock_test_id: int, db: Session = Depends(get_db), current_user: DBUser = Depends(get_current_user)):
    if not db.query(DBMockTest).filter(DBMockTest.id == mock_test_id).first():
        raise HTTPException(status_code=404, detail="Mock test not found")
    questions = db.query(DBQuestion).filter(DBQuestion.mock_test_id == mock_test_id).all()
    return {
        "status": "success",
        "questions": [{"id": q.id, "text": q.text, "option_a": q.option_a, "option_b": q.option_b, "option_c": q.option_c, "option_d": q.option_d, "subject": q.subject, "topic": q.topic} for q in questions],
    }

# Spaced-repetition forgetting curve (days). Index by `repetitions`.
SR_INTERVALS = [1, 7, 30, 90]

def _schedule_review_on_wrong(db, user_id, question_id):
    """A missed question enters (or resets in) the spaced-repetition queue."""
    item = db.query(DBReviewItem).filter(
        DBReviewItem.user_id == user_id, DBReviewItem.question_id == question_id).first()
    now = datetime.datetime.utcnow()
    if not item:
        item = DBReviewItem(user_id=user_id, question_id=question_id)
        db.add(item)
    item.repetitions = 0
    item.interval_days = SR_INTERVALS[0]
    item.next_review = now + datetime.timedelta(days=SR_INTERVALS[0])
    item.mastered = False

def _advance_review(db, user_id, question_id, correct):
    """Update an item's schedule after a revision attempt (SM-2-lite)."""
    item = db.query(DBReviewItem).filter(
        DBReviewItem.user_id == user_id, DBReviewItem.question_id == question_id).first()
    if not item:
        return
    now = datetime.datetime.utcnow()
    item.last_reviewed = now
    item.times_seen = (item.times_seen or 0) + 1
    if correct:
        item.times_correct = (item.times_correct or 0) + 1
        item.repetitions = (item.repetitions or 0) + 1
        if item.repetitions >= len(SR_INTERVALS):
            item.mastered = True
            item.interval_days = SR_INTERVALS[-1]
            item.next_review = now + datetime.timedelta(days=365)
        else:
            item.interval_days = SR_INTERVALS[item.repetitions]
            item.next_review = now + datetime.timedelta(days=item.interval_days)
    else:
        item.repetitions = 0
        item.interval_days = SR_INTERVALS[0]
        item.next_review = now + datetime.timedelta(days=SR_INTERVALS[0])
        item.mastered = False


@app.post("/mock-tests/{mock_test_id}/submit/", tags=["Mock Tests"])
def submit_mock_test(mock_test_id: int, answers: List[AnswerSubmit], db: Session = Depends(get_db), current_user: DBUser = Depends(get_current_user)):
    db_test = db.query(DBMockTest).filter(DBMockTest.id == mock_test_id).first()
    if not db_test:
        raise HTTPException(status_code=404, detail="Mock test not found")
    score = 0
    answer_objs = []
    total_time = 0
    _aml_pending = []
    for ans in answers:
        db_q = db.query(DBQuestion).filter(DBQuestion.id == ans.question_id, DBQuestion.mock_test_id == mock_test_id).first()
        if not db_q:
            continue
        is_correct = db_q.correct_answer == ans.selected_option.upper()
        if is_correct:
            score += 1
        tt = ans.time_taken if (ans.time_taken is not None and ans.time_taken >= 0) else None
        if tt:
            total_time += tt
        conf = (ans.confidence or "").lower().strip() or None
        if conf not in (None, "sure", "unsure", "guess"):
            conf = None
        answer_objs.append(DBAnswer(
            selected_option=ans.selected_option.upper(), is_correct=is_correct,
            question_id=ans.question_id, time_taken=tt, confidence=conf,
        ))
        # Wrong answers feed the spaced-repetition queue.
        if not is_correct:
            _schedule_review_on_wrong(db, current_user.id, db_q.id)
        _aml_pending.append((db_q, is_correct, ans.selected_option.upper(),
                             (tt * 1000 if tt else None), conf))
    db_attempt = DBTestAttempt(score=score, time_taken_seconds=total_time, user_id=current_user.id, mock_test_id=mock_test_id, answers=answer_objs)
    db.add(db_attempt)
    db.commit()
    db.refresh(db_attempt)
    # ── AML mastery capture — the test is already committed above, so this
    #    isolated block can never affect the student's submission. ──
    try:
        for (_q, _correct, _sel, _rms, _conf) in _aml_pending:
            _aml_record(db, current_user.id, _q, _correct, _sel, _rms, _conf, "mock")
        db.commit()
    except Exception:
        db.rollback()
    total = db_test.total_questions
    return {
        "status": "success", "message": "Test submitted successfully",
        "score": score, "total_questions": total,
        "percentage": round((score / total) * 100, 1) if total else 0,
        "attempt_id": db_attempt.id,
    }

# ── NCERT MCQs (book-wise + chapter-wise, AI-verified) ────────────────────────
@app.get("/ncert/books", tags=["NCERT"])
def ncert_books():
    """List NCERT books available for chapter-wise practice (with their chapters)."""
    books = []
    for b in syllabus_data.NCERT_BOOKS:
        books.append({
            "key": b["key"], "book": b["book"], "subject": b["subject"],
            "grade": b["grade"], "chapters": b["chapters"],
            "read_url": b.get("read_url", ""),
        })
    return {"status": "success", "books": books}


# ── Admin-uploaded NCERT book PDFs (read fully in-app) ──────────────────────────
# NOTE: everything here is written to be MEMORY-LEAN — the free tier has little
# RAM, so we never hold a whole zip (or all chapters) in memory at once. Zips are
# streamed to a temp file on disk, and chapters are stored one at a time.
def _ncert_chapter_names(zf):
    """Names of the chapter PDFs inside the zip, in reading order. Front-matter /
    answer files (…ps.pdf, …an.pdf, cover) are skipped; chapter files end in a number."""
    items = []
    for name in zf.namelist():
        low = name.lower()
        if low.endswith("/") or not low.endswith(".pdf"):
            continue
        base = low.rsplit("/", 1)[-1]
        m = re.search(r'(\d+)\.pdf$', base)
        if not m:
            continue
        items.append((int(m.group(1)), base, name))
    items.sort(key=lambda x: (x[0], x[1]))
    return items


def _store_book_from_zip_path(db, book_key, zip_path):
    """Store each chapter PDF from a zip on disk, ONE at a time (commit + free
    between chapters), so peak memory stays at a single chapter."""
    stored = 0
    with zipfile.ZipFile(zip_path) as zf:
        items = _ncert_chapter_names(zf)
        if not items:
            return 0
        db.query(DBNcertPdf).filter(DBNcertPdf.book_key == book_key).delete()
        db.commit()
        for _num, base, fullname in items:
            try:
                data = zf.read(fullname)
            except Exception:
                continue
            if data[:4] != b"%PDF":
                del data
                continue
            obj = DBNcertPdf(book_key=book_key, chapter_index=stored,
                             filename=base[:200], data=data, size=len(data))
            db.add(obj)
            db.commit()
            db.expunge(obj)          # drop the big bytes from the session identity map
            stored += 1
            del data, obj
            gc.collect()
    return stored


def _download_ncert_to_tempfile(url, timeout=120):
    """Stream an NCERT zip to a temp file on disk (never fully in memory).
    Render's servers have open egress to ncert.nic.in. Returns the temp path."""
    import ssl
    ctx = ssl.create_default_context()
    ctx.check_hostname = False
    ctx.verify_mode = ssl.CERT_NONE
    headers = {
        "User-Agent": ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
                       "(KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36"),
        "Accept": "*/*",
        "Referer": "https://ncert.nic.in/textbook.php",
    }
    last = ""
    for _u in (url, url.replace("https://", "http://")):
        try:
            with urlopen(_UrlRequest(_u, headers=headers), timeout=timeout, context=ctx) as r:
                tf = tempfile.NamedTemporaryFile(delete=False, suffix=".zip")
                try:
                    while True:
                        chunk = r.read(1 << 20)   # 1 MB at a time
                        if not chunk:
                            break
                        tf.write(chunk)
                finally:
                    tf.close()
                if os.path.getsize(tf.name) > 0:
                    return tf.name
                os.remove(tf.name)
        except Exception as e:  # noqa: BLE001
            last = f"{type(e).__name__} {getattr(e, 'code', '')}".strip()
    raise HTTPException(status_code=502, detail=f"Could not download from NCERT ({last or 'unknown'}).")


@app.post("/admin/ncert/upload-zip", tags=["NCERT"])
async def admin_ncert_upload_zip(file: UploadFile = File(...), book_key: str = Form(...),
                                 admin: DBUser = Depends(require_admin), db: Session = Depends(get_db)):
    """Admin: upload the NCERT book zip (as downloaded from ncert.nic.in). We
    stream it to disk, extract each chapter PDF and store it — memory-lean."""
    book = syllabus_data.get_ncert_book(book_key)
    if not book:
        raise HTTPException(status_code=400, detail="Unknown NCERT book key.")
    tf = tempfile.NamedTemporaryFile(delete=False, suffix=".zip")
    try:
        while True:
            chunk = await file.read(1 << 20)
            if not chunk:
                break
            tf.write(chunk)
    finally:
        tf.close()
    tmp = tf.name
    try:
        if os.path.getsize(tmp) == 0:
            raise HTTPException(status_code=400, detail="The uploaded file is empty.")
        try:
            stored = _store_book_from_zip_path(db, book_key, tmp)
        except zipfile.BadZipFile:
            raise HTTPException(status_code=400, detail="That doesn't look like a valid .zip file. Upload the zip downloaded from NCERT.")
        if not stored:
            raise HTTPException(status_code=400, detail="No chapter PDFs found inside the zip.")
    finally:
        try:
            os.remove(tmp)
        except OSError:
            pass
        gc.collect()
    return {"status": "success", "book_key": book_key, "chapters_stored": stored,
            "message": f"Stored {stored} chapter PDF(s) for {book['book']}."}


def _http_get(url, timeout=30, attempts=2):
    """Small server-side GET into memory (archive metadata + single chapter PDFs).
    Uses the same browser-like headers as the NCERT proxy (NCERT wants a Referer),
    with an http fallback and a bounded retry so a flaky moment recovers fast."""
    import ssl
    import time as _time
    ctx = ssl.create_default_context()
    ctx.check_hostname = False
    ctx.verify_mode = ssl.CERT_NONE
    headers = {
        "User-Agent": ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
                       "(KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36"),
        "Accept": "application/pdf,application/octet-stream,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.9",
        "Referer": "https://ncert.nic.in/textbook.php",
    }
    variants = [url]
    if url.startswith("https://"):
        variants.append(url.replace("https://", "http://", 1))
    last = None
    for i in range(max(1, attempts)):
        for u in variants:
            try:
                with urlopen(_UrlRequest(u, headers=headers), timeout=timeout, context=ctx) as r:
                    data = r.read()
                if data:
                    return data
            except Exception as e:  # noqa: BLE001
                last = e
        _time.sleep(0.7)
    if last:
        raise last
    raise RuntimeError("empty response")


def _record_chapter_urls(db, book_key, url_list):
    """Store just the chapter SOURCE URLs (no bytes) so the DB stays tiny; the
    reader streams each chapter from its source on demand."""
    db.query(DBNcertPdf).filter(DBNcertPdf.book_key == book_key).delete()
    db.commit()
    for i, (fn, src) in enumerate(url_list):
        db.add(DBNcertPdf(book_key=book_key, chapter_index=i, filename=(fn or "")[:200],
                          src_url=src, data=None, size=0))
    db.commit()
    return len(url_list)


def _import_from_archive(db, book_key, item_id):
    """Record a book's chapter URLs from an archive.org item (freely-distributed
    NCERT textbooks NCERT no longer hosts). No bytes are stored."""
    meta = json.loads(_http_get("https://archive.org/metadata/" + item_id).decode("utf-8", "ignore"))
    files = meta.get("files") or []
    chaps = []
    for f in files:
        nm = f.get("name", "")
        low = nm.lower()
        if not low.endswith(".pdf") or low.endswith("ps.pdf") or low.endswith("an.pdf"):
            continue
        m = re.search(r'(\d+)\.pdf$', low)   # chapter files end in a number
        if not m:
            continue
        chaps.append((int(m.group(1)), nm))
    chaps.sort(key=lambda x: (x[0], x[1]))
    if not chaps:
        return 0
    url_list = [(nm, "https://archive.org/download/" + item_id + "/" + nm) for _n, nm in chaps]
    return _record_chapter_urls(db, book_key, url_list)


def _ncert_zip_chapter_urls(zip_path):
    """From an NCERT book zip, list (filename, public chapter-PDF URL) in reading
    order — the chapter files also live at ncert.nic.in/textbook/pdf/<name>."""
    out = []
    with zipfile.ZipFile(zip_path) as zf:
        for _num, base, _full in _ncert_chapter_names(zf):
            out.append((base, "https://ncert.nic.in/textbook/pdf/" + base))
    return out


@app.post("/admin/ncert/import/{book_key}", tags=["NCERT"])
def admin_ncert_import(book_key: str, admin: DBUser = Depends(require_admin), db: Session = Depends(get_db)):
    """Admin: pull a book's chapters server-side and store them — no manual work.
    Source is the NCERT zip, or an archive.org mirror (for books NCERT dropped)."""
    book = syllabus_data.get_ncert_book(book_key)
    if not book:
        raise HTTPException(status_code=400, detail="Unknown NCERT book key.")
    url = (book.get("read_url") or "").strip()
    if not url:
        raise HTTPException(status_code=400, detail="No download link on file for this book — upload its zip manually.")
    if url.startswith("archive:"):
        try:
            stored = _import_from_archive(db, book_key, url[len("archive:"):])
        except Exception as e:  # noqa: BLE001
            try:
                db.rollback()
            except Exception:
                pass
            raise HTTPException(status_code=502, detail=f"Archive import failed: {type(e).__name__}: {str(e)[:180]}")
        if not stored:
            raise HTTPException(status_code=502, detail="No chapter PDFs found at the archive source.")
        return {"status": "success", "book_key": book_key, "chapters_stored": stored,
                "message": f"Imported {stored} chapter(s) for {book['book']}."}
    # NCERT-hosted book: fetch the zip once just to read the chapter list, then
    # store only the public chapter URLs (no bytes).
    tmp = _download_ncert_to_tempfile(url)
    try:
        try:
            url_list = _ncert_zip_chapter_urls(tmp)
        except zipfile.BadZipFile:
            raise HTTPException(status_code=502, detail="NCERT returned something that isn't a valid zip. Try again, or upload manually.")
        if not url_list:
            raise HTTPException(status_code=502, detail="No chapter PDFs found in the downloaded zip.")
        stored = _record_chapter_urls(db, book_key, url_list)
    finally:
        try:
            os.remove(tmp)
        except OSError:
            pass
        gc.collect()
    return {"status": "success", "book_key": book_key, "chapters_stored": stored,
            "message": f"Imported {stored} chapter(s) for {book['book']} from NCERT."}


@app.post("/admin/db/free-ncert-space", tags=["NCERT"])
def admin_free_ncert_space(admin: DBUser = Depends(require_admin), db: Session = Depends(get_db)):
    """Recovery: clear the stored NCERT PDF bytes to free the database (which went
    read-only after it filled up). Tries to force the session read-write first."""
    out = {}
    for stmt in ("SET default_transaction_read_only = off",
                 "SET SESSION CHARACTERISTICS AS TRANSACTION READ WRITE"):
        try:
            db.execute(sa_text(stmt)); out[stmt] = "ok"
        except Exception as e:  # noqa: BLE001
            out[stmt] = type(e).__name__
    try:
        out["rows_before"] = db.execute(sa_text("SELECT count(*) FROM ncert_pdfs")).scalar()
    except Exception as e:  # noqa: BLE001
        out["count_err"] = str(e)[:120]
    try:
        db.execute(sa_text("TRUNCATE TABLE ncert_pdfs"))
        db.commit()
        out["truncated"] = True
    except Exception as e:  # noqa: BLE001
        try:
            db.rollback()
        except Exception:
            pass
        out["truncated"] = False
        out["error"] = f"{type(e).__name__}: {str(e)[:200]}"
    return out


@app.post("/admin/maintenance/housekeeping", tags=["Admin"])
def admin_maintenance_housekeeping(admin: DBUser = Depends(require_admin), db: Session = Depends(get_db)):
    """One-time cleanup, admin-only and idempotent (safe to re-run):
      (1) scrub the dummy 'general' concept rows written to the owner account during
          AML live-verification, and
      (2) tag CSAT aptitude concepts — rows with subject='Other' whose subtopic matches
          the reasoning/quant pattern — as exam_track='csat', splitting them out of the
          prelims GS track.
    Returns a before/after report."""
    out = {"scrub": {}, "csat": {}}
    OWNER_EMAIL = "ltr2mm@gmail.com"
    CSAT_RX = (
        r"reasoning|series|coding|decoding|syllogism|analog|seating|arrangement|"
        r"permutation|combination|probability|calendar|clock|direction sense|"
        r"blood relation|data sufficiency|data interpretation|comprehension|"
        r"number system|number theory|\yratio\y|percentage|profit|time and work|"
        r"mensuration|verbal|non-verbal|assumption|conclusion|\ypuzzle|lcm|hcf|"
        r"arithmetic|averages?|mixture|cube|dice|venn|statement and|inequalit|"
        r"ranking|order and|figure|paper (folding|cutting)|counting|quadratic|"
        r"algebra|\ygeometry\y|trigonometr|logical"
    )
    # --- HK1: scrub owner's 'general' test rows ---
    try:
        oid = db.execute(sa_text("SELECT id FROM users WHERE lower(email)=lower(:e)"),
                         {"e": OWNER_EMAIL}).scalar()
        out["scrub"]["owner_id"] = oid
        if oid is not None:
            a = db.execute(sa_text(
                "DELETE FROM concept_attempts WHERE user_id=:u AND concept_key='general'"),
                {"u": oid}).rowcount
            m = db.execute(sa_text(
                "DELETE FROM concept_mastery WHERE user_id=:u AND concept_key='general'"),
                {"u": oid}).rowcount
            db.commit()
            out["scrub"]["attempts_deleted"] = a
            out["scrub"]["mastery_deleted"] = m
    except Exception as e:  # noqa: BLE001
        try:
            db.rollback()
        except Exception:
            pass
        out["scrub"]["error"] = f"{type(e).__name__}: {str(e)[:200]}"
    # --- HK2: CSAT track split ---
    try:
        has_et = db.execute(sa_text(
            "SELECT 1 FROM information_schema.columns "
            "WHERE table_name='concept_inventory' AND column_name='exam_track'")).scalar()
        out["csat"]["exam_track_column"] = bool(has_et)
        if has_et:
            before = db.execute(sa_text(
                "SELECT exam_track, count(*) FROM concept_inventory "
                "GROUP BY exam_track ORDER BY 2 DESC")).fetchall()
            out["csat"]["before"] = {str(k): v for k, v in before}
            matched = db.execute(sa_text(
                "SELECT count(*) FROM concept_inventory "
                "WHERE subject='Other' AND subtopic ~* :rx"), {"rx": CSAT_RX}).scalar()
            out["csat"]["matched"] = matched
            upd = db.execute(sa_text(
                "UPDATE concept_inventory SET exam_track='csat' "
                "WHERE subject='Other' AND subtopic ~* :rx"), {"rx": CSAT_RX}).rowcount
            db.commit()
            out["csat"]["updated"] = upd
            after = db.execute(sa_text(
                "SELECT exam_track, count(*) FROM concept_inventory "
                "GROUP BY exam_track ORDER BY 2 DESC")).fetchall()
            out["csat"]["after"] = {str(k): v for k, v in after}
    except Exception as e:  # noqa: BLE001
        try:
            db.rollback()
        except Exception:
            pass
        out["csat"]["error"] = f"{type(e).__name__}: {str(e)[:200]}"
    return out


@app.get("/admin/ncert/importable", tags=["NCERT"])
def admin_ncert_importable(admin: DBUser = Depends(require_admin), db: Session = Depends(get_db)):
    """Admin: every NCERT book, whether it has a download link, and whether it's
    already stored — drives the one-click 'Import from NCERT' list."""
    have = {}
    for (bk,) in db.query(DBNcertPdf.book_key).all():
        have[bk] = have.get(bk, 0) + 1
    out = []
    for b in syllabus_data.NCERT_BOOKS:
        out.append({"book_key": b["key"], "book": b["book"], "subject": b["subject"],
                    "grade": b.get("grade", ""), "has_link": bool(b.get("read_url")),
                    "stored_chapters": have.get(b["key"], 0)})
    return {"status": "success", "books": out}


@app.get("/admin/ncert/uploaded", tags=["NCERT"])
def admin_ncert_uploaded(admin: DBUser = Depends(require_admin), db: Session = Depends(get_db)):
    """Admin: which NCERT books have uploaded PDFs, with chapter counts & size."""
    rows = db.query(DBNcertPdf.book_key, DBNcertPdf.chapter_index, DBNcertPdf.size).all()
    agg = {}
    for bk, ci, sz in rows:
        a = agg.setdefault(bk, {"chapters": 0, "bytes": 0})
        a["chapters"] += 1
        a["bytes"] += (sz or 0)
    out = []
    for b in syllabus_data.NCERT_BOOKS:
        a = agg.get(b["key"])
        if a:
            out.append({"book_key": b["key"], "book": b["book"], "subject": b["subject"],
                        "chapters": a["chapters"], "mb": round(a["bytes"] / 1048576, 1)})
    return {"status": "success", "uploaded": out}


@app.delete("/admin/ncert/{book_key}", tags=["NCERT"])
def admin_ncert_delete(book_key: str, admin: DBUser = Depends(require_admin), db: Session = Depends(get_db)):
    n = db.query(DBNcertPdf).filter(DBNcertPdf.book_key == book_key).delete()
    db.commit()
    return {"status": "success", "removed": n}


@app.get("/ncert/uploaded-books", tags=["NCERT"])
def ncert_uploaded_books(current_user: DBUser = Depends(get_current_user), db: Session = Depends(get_db)):
    """Which NCERT books have an in-app readable upload (book_key -> chapter count)."""
    rows = db.query(DBNcertPdf.book_key).all()
    counts = {}
    for (bk,) in rows:
        counts[bk] = counts.get(bk, 0) + 1
    return {"status": "success", "counts": counts}


# In-memory LRU cache of fetched NCERT chapter PDFs. The DB stores only the
# source link, so without this every "Open chapter" re-fetches the remote PDF
# from ncert.nic.in — the main cause of slow chapter loads. Bounded so Cloud
# Run memory stays safe; first open is a remote fetch, repeats are instant.
from collections import OrderedDict as _OrderedDict
_NCERT_PDF_CACHE = _OrderedDict()
_NCERT_CACHE_MAX = 60 * 1024 * 1024   # ~60 MB across all cached chapters

def _ncert_cache_get(url):
    b = _NCERT_PDF_CACHE.get(url)
    if b is not None:
        _NCERT_PDF_CACHE.move_to_end(url)
    return b

def _ncert_cache_put(url, data):
    if not data:
        return
    _NCERT_PDF_CACHE[url] = data
    _NCERT_PDF_CACHE.move_to_end(url)
    total = sum(len(v) for v in _NCERT_PDF_CACHE.values())
    while total > _NCERT_CACHE_MAX and len(_NCERT_PDF_CACHE) > 1:
        _, v = _NCERT_PDF_CACHE.popitem(last=False)
        total -= len(v)

@app.get("/ncert/book-pdf/{book_key}/{ch}", tags=["NCERT"])
def ncert_book_pdf(book_key: str, ch: int, db: Session = Depends(get_db)):
    """Serve one NCERT chapter PDF, same-origin, so it renders in-app. Streams from
    the chapter's source URL (DB holds only the link); manual uploads serve bytes."""
    row = (db.query(DBNcertPdf)
           .filter(DBNcertPdf.book_key == book_key, DBNcertPdf.chapter_index == ch)
           .first())
    if not row:
        raise HTTPException(status_code=404, detail="That chapter isn't available yet.")
    hdrs = {"Content-Disposition": "inline; filename=ncert-chapter.pdf",
            "Cache-Control": "public, max-age=86400"}
    if row.data:                       # manual upload kept raw bytes
        return Response(content=row.data, media_type="application/pdf", headers=hdrs)
    if row.src_url:                    # stream from the source (NCERT / archive)
        data = _ncert_cache_get(row.src_url)
        if data is None:
            try:
                data = _http_get(row.src_url)
            except Exception:
                raise HTTPException(status_code=502, detail="Couldn't fetch this chapter from its source right now.")
            if not data or data[:4] != b"%PDF":
                raise HTTPException(status_code=502, detail="The source didn't return a valid PDF.")
            _ncert_cache_put(row.src_url, data)
        return Response(content=data, media_type="application/pdf", headers=hdrs)
    raise HTTPException(status_code=404, detail="That chapter isn't available yet.")


_BOOKPDF_PATH_RE = re.compile(r'^/ncert/book-pdf/[A-Za-z0-9_\-]+/\d+$')

@app.get("/read", response_class=HTMLResponse, tags=["NCERT"])
def read_pdf_viewer(u: str, t: str = "NCERT"):
    """A tiny standalone PDF viewer page. 'Open in new tab' points here so the
    chapter always DISPLAYS in the browser (rendered with PDF.js) instead of
    downloading — regardless of the browser's PDF-handling setting."""
    if not _BOOKPDF_PATH_RE.match(u or ""):
        raise HTTPException(status_code=400, detail="Invalid document reference.")
    safe_t = (t or "NCERT").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")
    safe_u = u.replace("&", "&amp;").replace('"', "&quot;")
    html = """<!doctype html><html lang="en"><head><meta charset="utf-8">
<title>__T__</title><meta name="viewport" content="width=device-width,initial-scale=1">
<style>body{margin:0;background:#525659;font-family:-apple-system,Segoe UI,Roboto,sans-serif}
#bar{position:sticky;top:0;z-index:5;background:#323639;color:#fff;padding:9px 14px;display:flex;gap:14px;align-items:center;font-size:14px}
#bar b{flex:1;white-space:nowrap;overflow:hidden;text-overflow:ellipsis}
#bar a{color:#8ab4f8;text-decoration:none} #wrap{padding:14px 0;text-align:center}
canvas{display:block;margin:0 auto 12px;max-width:100%;box-shadow:0 2px 12px rgba(0,0,0,.5);background:#fff}
#msg{color:#fff;padding:2.5rem;font-size:15px}</style></head>
<body><div id="bar"><b>__T__</b><a href="__U__" download>Download</a></div>
<div id="wrap"><div id="msg">Loading…</div></div>
<script src="https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.min.js"></script>
<script>
(function(){
 try{pdfjsLib.GlobalWorkerOptions.workerSrc='https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.worker.min.js';}catch(e){}
 (async function(){
  var wrap=document.getElementById('wrap');
  try{
   var pdf=await pdfjsLib.getDocument('__U__').promise;
   wrap.innerHTML='';
   var ww=Math.min(950, window.innerWidth-20);
   for(var p=1;p<=pdf.numPages;p++){
    var page=await pdf.getPage(p);
    var base=page.getViewport({scale:1});
    var scale=Math.max(0.6, Math.min(2.2, ww/base.width));
    var vp=page.getViewport({scale:scale});
    var c=document.createElement('canvas'); c.width=vp.width; c.height=vp.height;
    wrap.appendChild(c);
    await page.render({canvasContext:c.getContext('2d'), viewport:vp}).promise;
   }
  }catch(e){ wrap.innerHTML='<div id="msg">Couldn\\'t display this file. <a href="__U__" style="color:#8ab4f8">Open the raw PDF</a></div>'; }
 })();
})();
</script></body></html>"""
    html = html.replace("__T__", safe_t).replace("__U__", safe_u)
    return HTMLResponse(content=html)


# ── Reader highlights & notes (saved for revision) ──────────────────────────────
class ReaderNoteIn(BaseModel):
    book_key: str
    chapter_index: Optional[int] = 0
    kind: Optional[str] = "note"          # 'note' | 'highlight'
    page: Optional[int] = None
    text: Optional[str] = None
    rects: Optional[List[dict]] = None    # normalized [{x,y,w,h}] for highlights
    color: Optional[str] = None
    label: Optional[str] = None           # 'Important' | 'Revise' | 'Doubt' | 'Fact'

def _note_out(n, with_book=False):
    _now = datetime.datetime.utcnow()
    nr = getattr(n, "next_review", None)
    due = (nr is None) or (nr <= _now)
    d = {"id": n.id, "book_key": n.book_key, "chapter_index": n.chapter_index,
         "kind": n.kind, "page": n.page, "text": n.text or "",
         "color": n.color or "", "label": getattr(n, "label", None) or "",
         "revise_stage": getattr(n, "revise_stage", 0) or 0,
         "last_revised": (n.last_revised.isoformat() if getattr(n, "last_revised", None) else None),
         "next_review": (nr.isoformat() if nr else None),
         "due": due,
         "created_at": n.created_at.isoformat() if n.created_at else None,
         "rects": (json.loads(n.rects) if n.rects else None)}
    if with_book:
        b = syllabus_data.get_ncert_book(n.book_key)
        d["book"] = b["book"] if b else n.book_key
        chs = (b or {}).get("chapters") or []
        d["chapter"] = chs[n.chapter_index] if (0 <= (n.chapter_index or 0) < len(chs)) else f"Chapter {(n.chapter_index or 0)+1}"
    return d

@app.post("/me/reader/note", tags=["Reader"])
def reader_note_create(payload: ReaderNoteIn, current_user: DBUser = Depends(get_current_user),
                       db: Session = Depends(get_db)):
    if not (payload.book_key or "").strip():
        raise HTTPException(status_code=400, detail="book_key is required.")
    if payload.kind == "note" and not (payload.text or "").strip():
        raise HTTPException(status_code=400, detail="Note text is empty.")
    n = DBReaderNote(
        user_id=current_user.id, book_key=payload.book_key.strip(),
        chapter_index=payload.chapter_index or 0,
        kind=(payload.kind or "note"), page=payload.page,
        text=(payload.text or "")[:8000],
        rects=(json.dumps(payload.rects) if payload.rects else None),
        color=(payload.color or None), label=(payload.label or None),
        revise_stage=0, next_review=None)
    db.add(n); db.commit(); db.refresh(n)
    return {"status": "success", "note": _note_out(n)}

@app.get("/me/reader/notes", tags=["Reader"])
def reader_notes_for_chapter(book_key: str, chapter_index: int = 0,
                             current_user: DBUser = Depends(get_current_user),
                             db: Session = Depends(get_db)):
    rows = (db.query(DBReaderNote)
            .filter(DBReaderNote.user_id == current_user.id,
                    DBReaderNote.book_key == book_key,
                    DBReaderNote.chapter_index == chapter_index)
            .order_by(DBReaderNote.id.asc()).all())
    return {"status": "success", "notes": [_note_out(n) for n in rows]}

@app.get("/me/reader/all-notes", tags=["Reader"])
def reader_all_notes(current_user: DBUser = Depends(get_current_user),
                     db: Session = Depends(get_db)):
    rows = (db.query(DBReaderNote)
            .filter(DBReaderNote.user_id == current_user.id)
            .order_by(DBReaderNote.book_key.asc(), DBReaderNote.chapter_index.asc(),
                      DBReaderNote.id.asc()).all())
    return {"status": "success", "notes": [_note_out(n, with_book=True) for n in rows],
            "count": len(rows)}

class ReaderNoteEdit(BaseModel):
    text: Optional[str] = None
    color: Optional[str] = None
    label: Optional[str] = None

# Colour label → highlight colour, so a label and its colour stay in sync.
_LABEL_COLORS = {
    "Important": "#f87171",   # red
    "Revise":    "#f5d032",   # yellow
    "Doubt":     "#a78bfa",   # purple
    "Fact":      "#5eead4",   # teal
}

@app.put("/me/reader/note/{note_id}", tags=["Reader"])
def reader_note_edit(note_id: int, payload: ReaderNoteEdit,
                     current_user: DBUser = Depends(get_current_user),
                     db: Session = Depends(get_db)):
    n = (db.query(DBReaderNote)
         .filter(DBReaderNote.id == note_id, DBReaderNote.user_id == current_user.id).first())
    if not n:
        raise HTTPException(status_code=404, detail="Note not found.")
    if payload.text is not None:
        n.text = (payload.text or "")[:8000]
    if payload.label is not None:
        lbl = (payload.label or "").strip() or None
        n.label = lbl
        # keep the on-page highlight colour in sync with the chosen label
        if lbl in _LABEL_COLORS:
            n.color = _LABEL_COLORS[lbl]
    if payload.color is not None:
        n.color = (payload.color or None)
    db.commit(); db.refresh(n)
    return {"status": "success", "note": _note_out(n, with_book=True)}

@app.post("/me/reader/note/{note_id}/revised", tags=["Reader"])
def reader_note_revised(note_id: int, current_user: DBUser = Depends(get_current_user),
                        db: Session = Depends(get_db)):
    """Mark an item revised — advance it along the spaced-repetition curve
    (1 → 7 → 30 → 90 days) and schedule the next review."""
    n = (db.query(DBReaderNote)
         .filter(DBReaderNote.id == note_id, DBReaderNote.user_id == current_user.id).first())
    if not n:
        raise HTTPException(status_code=404, detail="Note not found.")
    now = datetime.datetime.utcnow()
    stage = (n.revise_stage or 0)
    interval = SR_INTERVALS[stage] if stage < len(SR_INTERVALS) else SR_INTERVALS[-1]
    n.revise_stage = min(stage + 1, len(SR_INTERVALS))
    n.last_revised = now
    n.next_review = now + datetime.timedelta(days=interval)
    db.commit(); db.refresh(n)
    return {"status": "success", "note": _note_out(n, with_book=True)}

@app.get("/me/reader/revise-today", tags=["Reader"])
def reader_revise_today(limit: int = 5, current_user: DBUser = Depends(get_current_user),
                        db: Session = Depends(get_db)):
    """Items due for revision today: never-revised ones first (oldest first),
    then anything whose next_review date has passed."""
    now = datetime.datetime.utcnow()
    rows = (db.query(DBReaderNote)
            .filter(DBReaderNote.user_id == current_user.id)
            .filter((DBReaderNote.next_review == None) | (DBReaderNote.next_review <= now))  # noqa: E711
            .all())
    # due soonest / oldest unseen first; unrevised (next_review None) treated as most due
    def _key(n):
        nr = n.next_review
        return (0, n.created_at or now) if nr is None else (1, nr)
    rows.sort(key=_key)
    total_due = len(rows)
    rows = rows[:max(1, min(limit, 50))]
    return {"status": "success",
            "notes": [_note_out(n, with_book=True) for n in rows],
            "due_count": total_due}

# ── AI over a candidate's own saved notes (Tier 3) ────────────────────────────
class ReaderAIIn(BaseModel):
    book_key: Optional[str] = None
    chapter_index: Optional[int] = None     # None → whole book (or all notes if no book_key)
    n: Optional[int] = None

def _collect_note_text(db, user_id, book_key=None, chapter_index=None, kinds=None):
    """Gather a user's saved highlights/notes into a single study-text blob.
    Returns (text, count, subject, title)."""
    q = db.query(DBReaderNote).filter(DBReaderNote.user_id == user_id)
    if book_key:
        q = q.filter(DBReaderNote.book_key == book_key)
    if chapter_index is not None:
        q = q.filter(DBReaderNote.chapter_index == chapter_index)
    if kinds:
        q = q.filter(DBReaderNote.kind.in_(list(kinds)))
    rows = q.order_by(DBReaderNote.chapter_index.asc(), DBReaderNote.id.asc()).all()
    parts, subject, title = [], "", ""
    b = syllabus_data.get_ncert_book(book_key) if book_key else None
    if b:
        subject = b.get("subject") or ""
        title = b.get("book") or book_key
        chs = b.get("chapters") or []
        if chapter_index is not None and 0 <= chapter_index < len(chs):
            title = f"{title} — {chs[chapter_index]}"
    for n in rows:
        t = (n.text or "").strip()
        if t:
            tag = "[note] " if n.kind == "note" else ""
            parts.append("• " + tag + t)
    return "\n".join(parts), len(parts), subject, title

@app.post("/me/reader/ai/mcqs", tags=["Reader"])
def reader_ai_mcqs(payload: ReaderAIIn, current_user: DBUser = Depends(get_current_user),
                   db: Session = Depends(get_db)):
    text, cnt, subject, title = _collect_note_text(db, current_user.id, payload.book_key, payload.chapter_index)
    if cnt < 1 or len(text) < 30:
        raise HTTPException(status_code=400, detail="Not enough saved notes/highlights here yet to build questions. Highlight a few more passages first.")
    import gemini_service
    try:
        qs = gemini_service.mcqs_from_notes(text, subject=subject, n=(payload.n or 5))
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"AI could not generate questions right now: {e}")
    if not qs:
        raise HTTPException(status_code=502, detail="Couldn't generate questions from these notes. Please try again.")
    return {"status": "success", "questions": qs, "source_count": cnt, "title": title}

@app.post("/me/reader/ai/flashcards", tags=["Reader"])
def reader_ai_flashcards(payload: ReaderAIIn, current_user: DBUser = Depends(get_current_user),
                         db: Session = Depends(get_db)):
    text, cnt, subject, title = _collect_note_text(db, current_user.id, payload.book_key, payload.chapter_index)
    if cnt < 1 or len(text) < 30:
        raise HTTPException(status_code=400, detail="Not enough saved notes/highlights here yet to build flashcards. Highlight a few more passages first.")
    import gemini_service
    try:
        cards = gemini_service.flashcards_from_notes(text, n=(payload.n or 10))
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"AI could not generate flashcards right now: {e}")
    if not cards:
        raise HTTPException(status_code=502, detail="Couldn't generate flashcards from these notes. Please try again.")
    return {"status": "success", "cards": cards, "source_count": cnt, "title": title}

@app.post("/me/reader/ai/summary", tags=["Reader"])
def reader_ai_summary(payload: ReaderAIIn, current_user: DBUser = Depends(get_current_user),
                      db: Session = Depends(get_db)):
    text, cnt, subject, title = _collect_note_text(db, current_user.id, payload.book_key, payload.chapter_index)
    if cnt < 1 or len(text) < 30:
        raise HTTPException(status_code=400, detail="Not enough saved highlights here yet to summarise. Highlight a few passages first.")
    import gemini_service
    try:
        summary = gemini_service.summarize_highlights(text, title=title)
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"AI could not summarise right now: {e}")
    return {"status": "success", "summary": summary or "", "source_count": cnt, "title": title}

@app.delete("/me/reader/note/{note_id}", tags=["Reader"])
def reader_note_delete(note_id: int, current_user: DBUser = Depends(get_current_user),
                       db: Session = Depends(get_db)):
    n = (db.query(DBReaderNote)
         .filter(DBReaderNote.id == note_id, DBReaderNote.user_id == current_user.id).first())
    if not n:
        raise HTTPException(status_code=404, detail="Note not found.")
    db.delete(n); db.commit()
    return {"status": "success"}


_NCERT_PDF_RE = re.compile(r'^https://ncert\.nic\.in/textbook/pdf/[a-z0-9]+\.pdf$', re.I)

@app.get("/ncert/pdf", tags=["NCERT"])
def ncert_pdf_proxy(u: str):
    """Proxy an official NCERT chapter PDF so it can be read INSIDE the app.
    NCERT only offers whole-book .zip download links and blocks third-party PDF
    viewers, so we fetch the chapter PDF server-side and stream it same-origin.
    Public content only (no auth needed — an <iframe> can't send a bearer token),
    and locked to official ncert.nic.in chapter-PDF URLs."""
    if not _NCERT_PDF_RE.match(u or ""):
        raise HTTPException(status_code=400, detail="Only official NCERT chapter PDFs can be opened here.")
    import ssl
    _ctx = ssl.create_default_context()
    _ctx.check_hostname = False
    _ctx.verify_mode = ssl.CERT_NONE
    _headers = {
        "User-Agent": ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
                       "(KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36"),
        "Accept": "application/pdf,application/octet-stream,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.9",
        "Referer": "https://ncert.nic.in/textbook.php",
    }
    resp = None
    last_err = ""
    for _url in (u, u.replace("https://", "http://")):   # some NCERT edges only serve http
        try:
            req = _UrlRequest(_url, headers=_headers)
            resp = urlopen(req, timeout=20, context=_ctx)   # connection only; body streamed below
            break
        except Exception as e:  # noqa: BLE001
            last_err = f"{type(e).__name__} {getattr(e, 'code', '')}".strip()
            resp = None
    if resp is None:
        raise HTTPException(status_code=502,
                            detail=f"Could not fetch the NCERT PDF ({last_err or 'unknown'}). Use 'Open in new tab'.")

    def _stream(r):
        # Stream in 64 KB chunks so a large chapter PDF is never fully held in RAM
        # (reading whole PDFs into memory was OOM-crashing the container -> 502s).
        try:
            while True:
                chunk = r.read(65536)
                if not chunk:
                    break
                yield chunk
        finally:
            try:
                r.close()
            except Exception:
                pass

    return StreamingResponse(_stream(resp), media_type="application/pdf",
                             headers={"Content-Disposition": "inline; filename=ncert-chapter.pdf",
                                      "Cache-Control": "public, max-age=86400"})


# ── Full-book PDFs we host in our own Google Cloud Storage bucket ───────────────
# Served SAME-ORIGIN (through this app) so the browser renders them inline in the
# reader instead of showing a cross-origin download placeholder. GCS is Google
# infrastructure in the same project/region, so there is no IP block and the fetch
# is fast. Range requests are forwarded so large books (100 MB+) scroll smoothly.
_NCERT_GCS_BUCKET = "aivora-production-ncert-books"
_NCERT_KEY_RE = re.compile(r'^[a-z0-9_]{1,64}$', re.I)
# Cloud Run caps a single HTTP/1 response at ~32 MiB, so we never emit more than
# this per response. Books larger than this are served as byte ranges (the PDF.js
# reader requests small ranges and learns the real total from Content-Range).
_GCS_MAX_CHUNK = 30 * 1024 * 1024

def _gcs_req_headers(extra=None):
    h = {
        "User-Agent": ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
                       "(KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36"),
        "Accept": "application/pdf,application/octet-stream,*/*;q=0.8",
    }
    if extra:
        h.update(extra)
    return h

@app.get("/ncert/gcs/{key}", tags=["NCERT"])
def ncert_gcs_proxy(key: str, request: Request):
    """Stream a hosted full-book PDF from our GCS bucket, same-origin, so it opens
    inside the in-app reader. Public content only; key is locked to [a-z0-9_].
    Every response is kept under Cloud Run's ~32 MiB limit by clamping the byte
    window; the true total size is always reported via Content-Range so the reader
    can page through large books with range requests."""
    if not _NCERT_KEY_RE.match(key or ""):
        raise HTTPException(status_code=400, detail="Invalid book id.")
    url = f"https://storage.googleapis.com/{_NCERT_GCS_BUCKET}/{key}.pdf"
    # Resolve the byte WINDOW to fetch from GCS. For normal bounded ranges and
    # no-range requests we do NOT need the total size up front — we request a
    # capped window and read the real total back from GCS's Content-Range. Only a
    # suffix range ("bytes=-N") needs the total first, so that alone costs a probe.
    rng = request.headers.get("range")
    had_range = bool(rng)
    gstart = 0
    gend = _GCS_MAX_CHUNK - 1            # default window (no-range / open-ended)
    suffix_n = None
    if rng:
        m = re.match(r"bytes=(\d*)-(\d*)", rng.strip())
        if m:
            g1, g2 = m.group(1), m.group(2)
            if g1 == "" and g2 != "":                       # suffix: last N bytes
                suffix_n = int(g2)
            else:
                gstart = int(g1) if g1 else 0
                if g2 != "":
                    gend = int(g2)
                    if gend - gstart + 1 > _GCS_MAX_CHUNK:
                        gend = gstart + _GCS_MAX_CHUNK - 1
                else:                                        # open-ended
                    gend = gstart + _GCS_MAX_CHUNK - 1
    if suffix_n is not None:
        try:
            probe = urlopen(_UrlRequest(url, headers=_gcs_req_headers({"Range": "bytes=0-0"})), timeout=15)
            cr0 = probe.getheader("Content-Range") or ""
            try:
                probe.read(1)
            except Exception:
                pass
            probe.close()
            total0 = int(cr0.split("/")[-1])
        except Exception as e:  # noqa: BLE001
            raise HTTPException(status_code=502,
                                detail=f"Could not load the book right now ({type(e).__name__}). Use 'Open in new tab'.")
        gstart = max(0, total0 - suffix_n)
        gend = total0 - 1
        if gend - gstart + 1 > _GCS_MAX_CHUNK:
            gstart = gend - _GCS_MAX_CHUNK + 1
    if gstart < 0:
        gstart = 0
    if gend < gstart:
        gend = gstart
    # Fetch exactly that window (single GCS request).
    try:
        resp = urlopen(_UrlRequest(url, headers=_gcs_req_headers({"Range": f"bytes={gstart}-{gend}"})), timeout=30)
    except Exception as e:  # noqa: BLE001
        raise HTTPException(status_code=502,
                            detail=f"Could not load the book right now ({type(e).__name__}). Use 'Open in new tab'.")
    # Read the true total + actual end from GCS's Content-Range ("bytes s-e/total").
    cr = resp.getheader("Content-Range") or ""
    total = None
    actual_end = gend
    try:
        seg, tot = cr.split(" ", 1)[1].split("/")
        total = int(tot)
        actual_end = int(seg.split("-")[1])
    except Exception:
        cl_probe = resp.getheader("Content-Length")
        if cl_probe:
            actual_end = gstart + int(cl_probe) - 1
            total = actual_end + 1
    start, end = gstart, actual_end
    length = end - start + 1
    is_full = (not had_range) and start == 0 and (total is not None) and end == total - 1
    out_headers = {
        "Accept-Ranges": "bytes",
        "Content-Length": str(length),
        "Content-Type": "application/pdf",
        "Content-Disposition": f'inline; filename="{key}.pdf"',
        "Cache-Control": "public, max-age=86400",
    }
    if not is_full and total is not None:
        out_headers["Content-Range"] = f"bytes {start}-{end}/{total}"

    def _stream(r):
        try:
            while True:
                chunk = r.read(65536)
                if not chunk:
                    break
                yield chunk
        finally:
            try:
                r.close()
            except Exception:
                pass

    return StreamingResponse(_stream(resp), status_code=(200 if is_full else 206),
                             media_type="application/pdf", headers=out_headers)


# ── NCERT reading progress (dashboard + continue-reading) ──────────────────────
class NcertProgressUpdate(BaseModel):
    book_key: str
    status: Optional[str] = None          # 'reading' | 'completed'
    last_page: Optional[int] = None       # 0-based furthest page reached
    pages_total: Optional[int] = None

def _ncert_key_maps():
    key_title, key_subject, key_grade, subject_order = {}, {}, {}, []
    for b in syllabus_data.NCERT_BOOKS:
        key_title[b["key"]] = b["book"]
        key_subject[b["key"]] = b["subject"]
        key_grade[b["key"]] = b.get("grade", "")
        if b["subject"] not in subject_order:
            subject_order.append(b["subject"])
    return key_title, key_subject, key_grade, subject_order

@app.post("/ncert/progress", tags=["NCERT"])
def ncert_progress_update(payload: NcertProgressUpdate, db: Session = Depends(get_db),
                          current_user: DBUser = Depends(get_current_user)):
    """Record reading progress: open a book ('reading'), update furthest page, or
    mark it 'completed'. One row per (user, book)."""
    if not syllabus_data.get_ncert_book(payload.book_key):
        raise HTTPException(status_code=404, detail="NCERT book not found")
    row = (db.query(DBNcertReading)
             .filter(DBNcertReading.user_id == current_user.id,
                     DBNcertReading.book_key == payload.book_key).first())
    if not row:
        row = DBNcertReading(user_id=current_user.id, book_key=payload.book_key, status="reading")
        db.add(row)
    if payload.status in ("reading", "completed"):
        row.status = payload.status
    if payload.pages_total and payload.pages_total > 0:
        row.pages_total = payload.pages_total
    if payload.last_page is not None and payload.last_page >= 0:
        row.last_page = max(int(row.last_page or 0), int(payload.last_page))
    db.commit()
    return {"status": "success", "book_status": row.status}

@app.get("/ncert/progress", tags=["NCERT"])
def ncert_progress(db: Session = Depends(get_db),
                   current_user: DBUser = Depends(get_current_user)):
    """Everything the NCERT dashboard needs: completion %, questions attended,
    reading-coverage + MCQ-accuracy by subject, and completed/pending/in-progress lists."""
    key_title, key_subject, key_grade, subject_order = _ncert_key_maps()
    all_keys = list(key_title.keys())
    total_books = len(all_keys)

    rows = db.query(DBNcertReading).filter(DBNcertReading.user_id == current_user.id).all()
    by_key = {r.book_key: r for r in rows}
    completed_keys = [k for k in all_keys if k in by_key and by_key[k].status == "completed"]
    completed_set = set(completed_keys)

    # reading coverage by subject: completed vs total
    reading_by_subject = []
    for s in subject_order:
        keys_in = [k for k in all_keys if key_subject[k] == s]
        done = sum(1 for k in keys_in if k in completed_set)
        reading_by_subject.append({"subject": s, "done": done, "total": len(keys_in)})

    # MCQ accuracy by subject — NCERT-only (questions whose book is an NCERT title)
    ncert_titles = set(key_title.values())
    q = (db.query(DBQuestion.subject, DBAnswer.is_correct)
           .join(DBAnswer, DBAnswer.question_id == DBQuestion.id)
           .join(DBTestAttempt, DBAnswer.test_attempt_id == DBTestAttempt.id)
           .filter(DBTestAttempt.user_id == current_user.id,
                   DBQuestion.book.in_(ncert_titles)))
    acc = {}
    questions_attended = 0
    for subj, ok in q.all():
        questions_attended += 1
        a = acc.setdefault(subj or "Other", [0, 0])
        a[1] += 1
        if ok:
            a[0] += 1
    mcq_by_subject = []
    for s in subject_order:
        c, t = acc.get(s, [0, 0])
        if t > 0:
            mcq_by_subject.append({"subject": s, "correct": c, "total": t,
                                   "pct": round(100 * c / t)})

    def _info(k):
        return {"key": k, "title": key_title[k], "subject": key_subject[k], "grade": key_grade[k]}

    completed = [_info(k) for k in completed_keys]
    pending = [_info(k) for k in all_keys if k not in completed_set]
    in_progress = []
    for k in all_keys:
        r = by_key.get(k)
        if r and r.status != "completed" and (r.last_page or 0) > 0:
            pt = r.pages_total or 0
            in_progress.append({**_info(k), "last_page": r.last_page or 0,
                                "pages_total": pt,
                                "pct": (round(100 * (r.last_page + 1) / pt) if pt else 0)})

    return {
        "status": "success",
        "total_books": total_books,
        "completed_count": len(completed_keys),
        "pct": round(100 * len(completed_keys) / total_books) if total_books else 0,
        "questions_attended": questions_attended,
        "reading_by_subject": reading_by_subject,
        "mcq_by_subject": mcq_by_subject,
        "completed": completed,
        "pending": pending,
        "in_progress": in_progress,
    }


@app.get("/reference-books", tags=["NCERT"])
def reference_books():
    """Standard reference books -> subject -> high-yield topics (for grounded generation)."""
    return {"status": "success", "reference_books": syllabus_data.REFERENCE_BOOKS}


@app.get("/subjects", tags=["NCERT"])
def subjects():
    """Subjects with their high-yield topics (for the Subjectwise practice tab)."""
    return {"status": "success", "subjects": syllabus_data.SUBJECT_TOPICS}


@app.post("/ncert/generate", tags=["NCERT"])
def ncert_generate(request: NcertGenerateRequest, db: Session = Depends(get_db),
                   current_user: DBUser = Depends(get_current_user)):
    """Generate a ready-to-take, AI-VERIFIED MCQ set from a specific NCERT book + chapter."""
    book = syllabus_data.get_ncert_book(request.book_key)
    if not book:
        raise HTTPException(status_code=404, detail="NCERT book not found")

    # "All chapters" => generate across the whole book (chapter left blank).
    chapter_raw = (request.chapter or "").strip()
    all_chapters = (not chapter_raw) or chapter_raw.lower() in ("all", "all chapters", "__all__")
    if not all_chapters and chapter_raw not in book["chapters"]:
        raise HTTPException(status_code=400, detail="Chapter not found in this book")
    gen_chapter = "" if all_chapters else chapter_raw
    label = "All Chapters" if all_chapters else chapter_raw

    num = max(1, min(int(request.num_questions or 5), 100))
    difficulty = (request.difficulty or "medium").lower()
    qtype = (request.question_type or "all").lower()
    try:
        questions = generate_ncert_mcqs(
            book=book["book"], chapter=gen_chapter,
            subject=book["subject"], num_questions=num, difficulty=difficulty,
            question_type=qtype,
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Generation error: {str(e)}")
    if not questions:
        raise HTTPException(status_code=502, detail="Could not generate verified questions. Please try again.")

    title = f"NCERT • {book['book']} • {label}"
    duration = max(5, len(questions))
    db_test = DBMockTest(
        title=title,
        description=f"AI-verified NCERT MCQs from {book['book']} — {label} ({difficulty})",
        subject=book["subject"],
        total_questions=len(questions),
        duration_minutes=duration,
        user_id=current_user.id,
    )
    db.add(db_test); db.commit(); db.refresh(db_test)
    for q in questions:
        db.add(DBQuestion(
            text=q["text"], option_a=q["option_a"], option_b=q["option_b"],
            option_c=q["option_c"], option_d=q["option_d"],
            correct_answer=q["correct_answer"], explanation=q.get("explanation", ""),
            subject=book["subject"],
            book=book["book"], chapter=(None if all_chapters else chapter_raw),
            topic=q.get("topic") or label,
            difficulty=difficulty, question_type=q.get("question_type", "direct"),
            mock_test_id=db_test.id,
        ))
    db.commit()
    return {
        "status": "success",
        "message": f"Generated {len(questions)} verified NCERT questions",
        "mock_test_id": db_test.id,
        "questions_added": len(questions),
        "duration_minutes": duration,
        "title": title,
    }


def _reuse_pool(db, user_id, subject, topic, difficulty, limit):
    """Pull previously-generated verified questions for this concept that the student
    hasn't answered yet — caching by concept to cut generation cost (roadmap §7)."""
    if limit <= 0:
        return []
    answered = {a.question_id for a in (db.query(DBAnswer.question_id)
                .join(DBTestAttempt, DBAnswer.test_attempt_id == DBTestAttempt.id)
                .filter(DBTestAttempt.user_id == user_id).all())}
    q = db.query(DBQuestion).filter(DBQuestion.subject == subject)
    if topic:
        q = q.filter(DBQuestion.topic == topic)
    if difficulty:
        q = q.filter(DBQuestion.difficulty == difficulty)
    cands = q.order_by(DBQuestion.id.desc()).limit(400).all()
    seen_text, out = set(), []
    for dq in cands:
        if dq.id in answered or dq.text in seen_text or not dq.correct_answer:
            continue
        seen_text.add(dq.text)
        out.append({"text": dq.text, "option_a": dq.option_a, "option_b": dq.option_b,
                    "option_c": dq.option_c, "option_d": dq.option_d,
                    "correct_answer": dq.correct_answer, "explanation": dq.explanation or "",
                    "topic": dq.topic, "question_type": dq.question_type or "direct"})
        if len(out) >= limit:
            break
    return out

@app.post("/mock-tests/generate-verified/", tags=["Tests"])
def generate_verified_test(request: VerifiedGenerateRequest, db: Session = Depends(get_db),
                           current_user: DBUser = Depends(get_current_user)):
    """Generate an AI-VERIFIED test from a subject/topic (optionally grounded in a book)."""
    num = max(1, min(int(request.num_questions or 5), 100))
    difficulty = (request.difficulty or "medium").lower()
    qtype = (request.question_type or "all").lower()
    subject = request.subject or "General Studies"
    topic = (request.topic or "").strip()
    book = (request.book or "").strip()
    reused = []
    if getattr(request, "reuse", False) and not book:
        try:
            reused = _reuse_pool(db, current_user.id, subject, topic, difficulty, num // 2)
        except Exception:
            reused = []
    gen_needed = num - len(reused)
    source_context = ""
    try:
        if db.query(DBKnowledgeChunk).filter(DBKnowledgeChunk.subject == subject).first():
            source_context = _retrieve_context(db, subject, topic)
    except Exception:
        source_context = ""
    if not source_context and not book:
        try:
            source_context = _concept_context(db, subject, topic)
        except Exception:
            source_context = ""
    plan = []
    try:
        if not book:
            plan = concept_engine.plan_set(subject, topic, gen_needed)
    except Exception:
        plan = []
    questions = []
    if gen_needed > 0 and plan:
        for _item in plan:
            try:
                _ctx = _concept_context(db, subject, _item["topic"]) or source_context
            except Exception:
                _ctx = source_context
            try:
                _qs = generate_verified_questions(subject=subject, topic=_item["topic"],
                    num_questions=_item["count"], difficulty=difficulty,
                    question_type=_item["question_type"], source_context=_ctx)
            except Exception:
                _qs = []
            for _q in (_qs or []):
                _q.setdefault("topic", _item["topic"])
            questions.extend(_qs or [])
    if gen_needed > 0 and not questions:
        try:
            questions = generate_verified_questions(
                subject=subject, topic=topic, num_questions=gen_needed, difficulty=difficulty, book=book,
                question_type=qtype, source_context=source_context,
            )
        except Exception as e:
            if not reused:
                raise HTTPException(status_code=500, detail=f"Generation error: {str(e)}")
    questions = reused + (questions or [])
    if not questions:
        raise HTTPException(status_code=502, detail="Could not generate verified questions. Please try again.")

    label = topic or book or subject
    title = f"Verified • {subject}" + (f" • {label}" if label and label != subject else "")
    duration = max(5, len(questions))
    db_test = DBMockTest(
        title=title, description=f"AI-verified MCQs — {label} ({difficulty})",
        subject=subject, total_questions=len(questions),
        duration_minutes=duration, user_id=current_user.id,
    )
    db.add(db_test); db.commit(); db.refresh(db_test)
    for q in questions:
        db.add(DBQuestion(
            text=q["text"], option_a=q["option_a"], option_b=q["option_b"],
            option_c=q["option_c"], option_d=q["option_d"],
            correct_answer=q["correct_answer"], explanation=q.get("explanation", ""),
            subject=subject, book=book or None, topic=q.get("topic") or topic,
            difficulty=difficulty, question_type=q.get("question_type", "direct"),
            mock_test_id=db_test.id,
        ))
    db.commit()
    return {
        "status": "success", "message": f"Generated {len(questions)} verified questions",
        "mock_test_id": db_test.id, "questions_added": len(questions),
        "duration_minutes": duration, "title": title,
    }


@app.post("/mock-tests/generate-stream", tags=["Tests"])
def generate_verified_stream(request: VerifiedGenerateRequest,
                             current_user: DBUser = Depends(get_current_user)):
    """Server-Sent-Events version of verified generation: streams live progress so the
    UI shows real activity instead of a frozen spinner (and starting the response
    immediately also avoids gateway timeouts). Falls back client-side to the classic
    endpoint if unavailable."""
    from fastapi.responses import StreamingResponse
    uid = current_user.id
    num = max(1, min(int(request.num_questions or 5), 100))
    difficulty = (request.difficulty or "medium").lower()
    qtype = (request.question_type or "all").lower()
    subject = request.subject or "General Studies"
    topic = (request.topic or "").strip()
    book = (request.book or "").strip()
    reuse = bool(getattr(request, "reuse", False))

    if concept_engine.is_mapped(subject, topic) and not book:
        raise HTTPException(status_code=409, detail="concept-weighted")

    def _sse(obj):
        return "data: " + json.dumps(obj, ensure_ascii=False) + "\n\n"

    def _gen():
        _db = SessionLocal()
        try:
            yield _sse({"stage": "start", "target": num})
            reused = []
            if reuse and not book:
                try:
                    reused = _reuse_pool(_db, uid, subject, topic, difficulty, num // 2)
                except Exception:
                    reused = []
            gen_needed = num - len(reused)
            source_context = ""
            try:
                if _db.query(DBKnowledgeChunk).filter(DBKnowledgeChunk.subject == subject).first():
                    source_context = _retrieve_context(_db, subject, topic)
            except Exception:
                source_context = ""
            if not source_context and not book:
                try:
                    source_context = _concept_context(_db, subject, topic)
                except Exception:
                    source_context = ""
            gen_qs = []
            if gen_needed > 0:
                try:
                    for ev in generate_verified_questions_stream(
                        subject=subject, topic=topic, num_questions=gen_needed, difficulty=difficulty,
                        book=book, question_type=qtype, source_context=source_context):
                        if isinstance(ev, dict) and ev.get("stage") == "result":
                            gen_qs = ev.get("questions") or []
                        else:
                            yield _sse(ev)
                except Exception as e:
                    if not reused:
                        yield _sse({"stage": "error", "detail": f"Generation error: {str(e)}"})
                        return
            questions = reused + (gen_qs or [])
            if not questions:
                yield _sse({"stage": "error", "detail": "Could not generate verified questions. Please try again."})
                return
            label = topic or book or subject
            title = f"Verified \u2022 {subject}" + (f" \u2022 {label}" if label and label != subject else "")
            duration = max(5, len(questions))
            db_test = DBMockTest(
                title=title, description=f"AI-verified MCQs \u2014 {label} ({difficulty})",
                subject=subject, total_questions=len(questions),
                duration_minutes=duration, user_id=uid,
            )
            _db.add(db_test); _db.commit(); _db.refresh(db_test)
            for q in questions:
                _db.add(DBQuestion(
                    text=q["text"], option_a=q["option_a"], option_b=q["option_b"],
                    option_c=q["option_c"], option_d=q["option_d"],
                    correct_answer=q["correct_answer"], explanation=q.get("explanation", ""),
                    subject=subject, book=book or None, topic=q.get("topic") or topic,
                    difficulty=difficulty, question_type=q.get("question_type", "direct"),
                    mock_test_id=db_test.id,
                ))
            _db.commit()
            yield _sse({"stage": "done", "mock_test_id": db_test.id,
                        "questions_added": len(questions), "duration_minutes": duration, "title": title})
        except Exception as e:
            yield _sse({"stage": "error", "detail": str(e)})
        finally:
            _db.close()

    return StreamingResponse(_gen(), media_type="text/event-stream",
                             headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no",
                                      "Connection": "keep-alive"})




# ── Results ───────────────────────────────────────────────────────────────────
@app.get("/attempts/{attempt_id}/results/", tags=["Results"])
def get_attempt_results(attempt_id: int, db: Session = Depends(get_db), current_user: DBUser = Depends(get_current_user)):
    db_attempt = db.query(DBTestAttempt).filter(DBTestAttempt.id == attempt_id, DBTestAttempt.user_id == current_user.id).first()
    if not db_attempt:
        raise HTTPException(status_code=404, detail="Attempt not found")
    # Map the user's submitted answers by question, then walk EVERY question in the
    # test so the review (and explanations) covers answered AND skipped questions.
    ans_map = {a.question_id: a for a in db.query(DBAnswer).filter(DBAnswer.test_attempt_id == attempt_id).all()}
    questions = db.query(DBQuestion).filter(DBQuestion.mock_test_id == db_attempt.mock_test_id).order_by(DBQuestion.id).all()
    results = []
    for q in questions:
        ans = ans_map.get(q.id)
        results.append({
            "question_id": q.id, "question_text": q.text,
            "option_a": q.option_a, "option_b": q.option_b, "option_c": q.option_c, "option_d": q.option_d,
            "selected_option": ans.selected_option if ans else None,
            "correct_answer": q.correct_answer,
            "is_correct": bool(ans.is_correct) if ans else False,
            "answered": ans is not None,
            "explanation": q.explanation,
        })
    total = db_attempt.mock_test.total_questions
    return {
        "status": "success", "score": db_attempt.score, "total_questions": total,
        "percentage": round((db_attempt.score / total) * 100, 1) if total else 0,
        "mock_test_id": db_attempt.mock_test_id,
        "title": db_attempt.mock_test.title,
        "results": results,
    }

@app.get("/my-attempts/", tags=["Results"])
def get_my_attempts(db: Session = Depends(get_db), current_user: DBUser = Depends(get_current_user)):
    attempts = db.query(DBTestAttempt).filter(DBTestAttempt.user_id == current_user.id).order_by(DBTestAttempt.completed_at.desc()).all()
    return {
        "status": "success",
        "attempts": [
            {
                "attempt_id": a.id, "mock_test_id": a.mock_test_id,
                "mock_test_title": a.mock_test.title, "subject": a.mock_test.subject,
                "score": a.score, "total_questions": a.mock_test.total_questions,
                "percentage": round((a.score / a.mock_test.total_questions) * 100, 1) if a.mock_test.total_questions else 0,
                "completed_at": a.completed_at.isoformat() if a.completed_at else None,
            }
            for a in attempts
        ],
    }

# ── Learning Loop: Mistakes · Weak Topics · Spaced Repetition ──────────────────
_OPT_LETTERS = ["A", "B", "C", "D"]

def _opt_text(q, letter):
    if not letter:
        return None
    return getattr(q, "option_" + letter.lower(), None)

def _classify_mistake(ans, q):
    """Heuristic auto-tag for a wrong answer when the student hasn't tagged it."""
    if ans.wrong_reason:
        return ans.wrong_reason
    conf = (ans.confidence or "").lower()
    tt = ans.time_taken or 0
    if conf == "guess":
        return "guess"
    if conf == "sure":
        return "conceptual"        # confidently wrong → a real knowledge gap
    if tt and tt <= 8:
        return "careless"          # very fast + wrong → likely misread/rushed
    return "review"                # default: needs review

@app.get("/me/mistakes", tags=["Learning"])
def my_mistakes(subject: Optional[str] = None, reason: Optional[str] = None,
                confidence: Optional[str] = None, include_resolved: bool = False,
                limit: int = 200, db: Session = Depends(get_db),
                current_user: DBUser = Depends(get_current_user)):
    """Wrong answers, de-duplicated to the most recent attempt per question, with
    filters, resolved-state, community difficulty and bookmark flags."""
    resolved_ids = {r.question_id for r in db.query(DBResolvedMistake).filter(
        DBResolvedMistake.user_id == current_user.id).all()}
    bookmarked_ids = {b.question_id for b in db.query(DBBookmark).filter(
        DBBookmark.user_id == current_user.id).all()}
    rows = (db.query(DBAnswer, DBQuestion, DBTestAttempt)
            .join(DBQuestion, DBAnswer.question_id == DBQuestion.id)
            .join(DBTestAttempt, DBAnswer.test_attempt_id == DBTestAttempt.id)
            .filter(DBTestAttempt.user_id == current_user.id, DBAnswer.is_correct == False)
            .order_by(DBTestAttempt.completed_at.desc(), DBAnswer.id.desc())
            .all())
    seen = set()
    items = []
    subj_counter = {}
    resolved_count = 0
    for ans, q, att in rows:
        if q.id in seen:
            continue
        seen.add(q.id)
        is_resolved = q.id in resolved_ids
        if is_resolved:
            resolved_count += 1
        # Subject chips count all active (non-resolved) mistakes, independent of filters.
        if not is_resolved:
            s0 = q.subject or "General"
            subj_counter[s0] = subj_counter.get(s0, 0) + 1
        if is_resolved and not include_resolved:
            continue
        if subject and (q.subject or "").lower() != subject.lower():
            continue
        rsn = _classify_mistake(ans, q)
        if reason and rsn != reason.lower():
            continue
        if confidence and (ans.confidence or "").lower() != confidence.lower():
            continue
        items.append({
            "answer_id": ans.id, "question_id": q.id,
            "question_text": q.text,
            "options": {l: _opt_text(q, l) for l in _OPT_LETTERS},
            "your_answer": ans.selected_option,
            "your_answer_text": _opt_text(q, ans.selected_option),
            "correct_answer": q.correct_answer,
            "correct_answer_text": _opt_text(q, q.correct_answer),
            "explanation": q.explanation or "",
            "subject": q.subject or "General", "topic": q.topic or "", "book": q.book or "",
            "confidence": ans.confidence, "time_taken": ans.time_taken,
            "reason": rsn, "reason_is_auto": not bool(ans.wrong_reason),
            "resolved": is_resolved, "bookmarked": q.id in bookmarked_ids,
            "completed_at": att.completed_at.isoformat() if att.completed_at else None,
        })
        if len(items) >= max(1, min(limit, 500)):
            break
    ids = [it["question_id"] for it in items]
    if ids:
        comm = {}
        for qid, isc in db.query(DBAnswer.question_id, DBAnswer.is_correct).filter(
                DBAnswer.question_id.in_(ids)).all():
            d = comm.setdefault(qid, [0, 0])
            d[1] += 1
            d[0] += 1 if isc else 0
        for it in items:
            d = comm.get(it["question_id"])
            it["community_accuracy"] = round(d[0] / d[1] * 100) if d and d[1] else None
            it["community_attempts"] = d[1] if d else 0
    return {"status": "success", "count": len(items), "resolved_total": resolved_count,
            "by_subject": [{"subject": k, "count": v} for k, v in sorted(subj_counter.items(), key=lambda x: -x[1])],
            "mistakes": items}

@app.post("/me/mistakes/{question_id}/resolve", tags=["Learning"])
def resolve_mistake(question_id: int, db: Session = Depends(get_db),
                    current_user: DBUser = Depends(get_current_user)):
    """Toggle a mistake as understood/resolved (hidden from the active list)."""
    existing = db.query(DBResolvedMistake).filter(
        DBResolvedMistake.user_id == current_user.id,
        DBResolvedMistake.question_id == question_id).first()
    if existing:
        db.delete(existing); db.commit()
        return {"status": "success", "resolved": False}
    db.add(DBResolvedMistake(user_id=current_user.id, question_id=question_id)); db.commit()
    return {"status": "success", "resolved": True}

@app.post("/me/mistakes/{question_id}/revise", tags=["Learning"])
def add_mistake_to_revision(question_id: int, db: Session = Depends(get_db),
                            current_user: DBUser = Depends(get_current_user)):
    """Manually push a mistake into the spaced-repetition queue, due now."""
    if not db.query(DBQuestion).filter(DBQuestion.id == question_id).first():
        raise HTTPException(status_code=404, detail="Question not found")
    _schedule_review_on_wrong(db, current_user.id, question_id)
    db.commit()
    return {"status": "success", "scheduled": True}

@app.get("/me/revision/calendar", tags=["Learning"])
def revision_calendar(db: Session = Depends(get_db), current_user: DBUser = Depends(get_current_user)):
    """Day-by-day spaced-repetition schedule for the next 14 days."""
    now = datetime.datetime.utcnow()
    td = datetime.timedelta
    items = db.query(DBReviewItem).filter(
        DBReviewItem.user_id == current_user.id, DBReviewItem.mastered == False).all()
    overdue = 0
    perday = {}
    for it in items:
        if it.next_review is None:
            continue
        d = (it.next_review.date() - now.date()).days
        if d < 0:
            overdue += 1
        elif d <= 14:
            key = it.next_review.date().isoformat()
            perday[key] = perday.get(key, 0) + 1
    days = []
    for i in range(0, 14):
        dt = (now.date() + td(i))
        days.append({"date": dt.isoformat(), "label": dt.strftime("%a %d %b"),
                     "count": perday.get(dt.isoformat(), 0) + (overdue if i == 0 else 0),
                     "is_today": i == 0})
    return {"status": "success", "overdue": overdue, "days": days}

@app.put("/me/mistakes/{answer_id}/reason", tags=["Learning"])
def tag_mistake_reason(answer_id: int, payload: WrongReasonIn,
                       db: Session = Depends(get_db), current_user: DBUser = Depends(get_current_user)):
    ans = (db.query(DBAnswer).join(DBTestAttempt, DBAnswer.test_attempt_id == DBTestAttempt.id)
           .filter(DBAnswer.id == answer_id, DBTestAttempt.user_id == current_user.id).first())
    if not ans:
        raise HTTPException(status_code=404, detail="Answer not found")
    reason = (payload.reason or "").lower().strip()
    if reason not in ("conceptual", "factual", "careless", "misread", "guess"):
        raise HTTPException(status_code=400, detail="Invalid reason")
    ans.wrong_reason = reason
    db.commit()
    return {"status": "success", "answer_id": answer_id, "reason": reason}

@app.get("/me/weak-topics", tags=["Learning"])
def my_weak_topics(min_attempts: int = 3, db: Session = Depends(get_db),
                   current_user: DBUser = Depends(get_current_user)):
    """Per subject+topic accuracy across all of the student's answers; weakest first."""
    rows = (db.query(DBAnswer, DBQuestion)
            .join(DBQuestion, DBAnswer.question_id == DBQuestion.id)
            .join(DBTestAttempt, DBAnswer.test_attempt_id == DBTestAttempt.id)
            .filter(DBTestAttempt.user_id == current_user.id)
            .all())
    agg = {}
    subj_agg = {}
    for ans, q in rows:
        subj = q.subject or "General"
        topic = q.topic or "General"
        key = (subj, topic)
        a = agg.setdefault(key, {"attempted": 0, "correct": 0})
        a["attempted"] += 1
        a["correct"] += 1 if ans.is_correct else 0
        sa = subj_agg.setdefault(subj, {"attempted": 0, "correct": 0})
        sa["attempted"] += 1
        sa["correct"] += 1 if ans.is_correct else 0
    topics = []
    for (subj, topic), a in agg.items():
        if a["attempted"] < max(1, min_attempts):
            continue
        acc = round(a["correct"] / a["attempted"] * 100, 1)
        topics.append({"subject": subj, "topic": topic, "attempted": a["attempted"],
                       "correct": a["correct"], "accuracy": acc})
    topics.sort(key=lambda x: (x["accuracy"], -x["attempted"]))
    subjects = []
    for subj, a in subj_agg.items():
        acc = round(a["correct"] / a["attempted"] * 100, 1) if a["attempted"] else 0
        subjects.append({"subject": subj, "attempted": a["attempted"],
                         "correct": a["correct"], "accuracy": acc})
    subjects.sort(key=lambda x: x["accuracy"])
    return {"status": "success", "weak_topics": topics[:12], "subjects": subjects}

@app.get("/me/revision/due", tags=["Learning"])
def revision_due(include_upcoming: bool = False, db: Session = Depends(get_db),
                 current_user: DBUser = Depends(get_current_user)):
    """Questions due for spaced-repetition review right now."""
    now = datetime.datetime.utcnow()
    q = db.query(DBReviewItem).filter(DBReviewItem.user_id == current_user.id,
                                      DBReviewItem.mastered == False)
    due_items = q.filter(DBReviewItem.next_review <= now).order_by(DBReviewItem.next_review.asc()).all()
    total_scheduled = q.count()
    mastered_count = db.query(DBReviewItem).filter(
        DBReviewItem.user_id == current_user.id, DBReviewItem.mastered == True).count()
    questions = []
    for it in due_items:
        dq = db.query(DBQuestion).filter(DBQuestion.id == it.question_id).first()
        if not dq:
            continue
        questions.append({
            "review_id": it.id, "question_id": dq.id, "text": dq.text,
            "option_a": dq.option_a, "option_b": dq.option_b,
            "option_c": dq.option_c, "option_d": dq.option_d,
            "subject": dq.subject or "General", "topic": dq.topic or "",
            "repetitions": it.repetitions or 0,
        })
    return {"status": "success", "due_count": len(questions),
            "total_scheduled": total_scheduled, "mastered": mastered_count,
            "questions": questions}

@app.post("/me/revision/submit", tags=["Learning"])
def revision_submit(answers: List[RevisionAnswer], db: Session = Depends(get_db),
                    current_user: DBUser = Depends(get_current_user)):
    """Grade a revision round and advance each question's spaced-repetition schedule."""
    correct = 0
    results = []
    for ans in answers:
        dq = db.query(DBQuestion).filter(DBQuestion.id == ans.question_id).first()
        if not dq:
            continue
        is_correct = dq.correct_answer == (ans.selected_option or "").upper()
        if is_correct:
            correct += 1
        _advance_review(db, current_user.id, dq.id, is_correct)
        results.append({
            "question_id": dq.id, "is_correct": is_correct,
            "correct_answer": dq.correct_answer,
            "your_answer": (ans.selected_option or "").upper(),
            "explanation": dq.explanation or "",
        })
    db.commit()
    return {"status": "success", "correct": correct, "total": len(results), "results": results}

@app.get("/me/learning/summary", tags=["Learning"])
def learning_summary(db: Session = Depends(get_db), current_user: DBUser = Depends(get_current_user)):
    """Small counters for the dashboard learning-loop cards."""
    now = datetime.datetime.utcnow()
    mistakes = (db.query(DBAnswer.question_id)
                .join(DBTestAttempt, DBAnswer.test_attempt_id == DBTestAttempt.id)
                .filter(DBTestAttempt.user_id == current_user.id, DBAnswer.is_correct == False)
                .distinct().count())
    due = db.query(DBReviewItem).filter(
        DBReviewItem.user_id == current_user.id, DBReviewItem.mastered == False,
        DBReviewItem.next_review <= now).count()
    mastered = db.query(DBReviewItem).filter(
        DBReviewItem.user_id == current_user.id, DBReviewItem.mastered == True).count()
    return {"status": "success", "mistakes": mistakes, "due_revisions": due, "mastered": mastered}

@app.post("/me/mistakes/{question_id}/diagnose", tags=["Learning"])
def diagnose_my_mistake(question_id: int, db: Session = Depends(get_db),
                        current_user: DBUser = Depends(get_current_user)):
    """AI, confidence-aware post-mortem of one wrong answer the student gave."""
    ans = (db.query(DBAnswer).join(DBTestAttempt, DBAnswer.test_attempt_id == DBTestAttempt.id)
           .filter(DBTestAttempt.user_id == current_user.id,
                   DBAnswer.question_id == question_id, DBAnswer.is_correct == False)
           .order_by(DBAnswer.id.desc()).first())
    if not ans:
        raise HTTPException(status_code=404, detail="No wrong answer found for this question")
    q = db.query(DBQuestion).filter(DBQuestion.id == question_id).first()
    if not q:
        raise HTTPException(status_code=404, detail="Question not found")
    try:
        text = diagnose_mistake(
            question=q.text,
            options={l: getattr(q, "option_" + l.lower(), "") for l in ["A", "B", "C", "D"]},
            correct_letter=q.correct_answer, chosen_letter=ans.selected_option,
            confidence=ans.confidence or "", subject=q.subject or "", topic=q.topic or "",
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Diagnosis error: {str(e)}")
    return {"status": "success", "question_id": question_id, "diagnosis": text}

@app.get("/me/insights", tags=["Learning"])
def my_insights(db: Session = Depends(get_db), current_user: DBUser = Depends(get_current_user)):
    """Confidence calibration + mistake-pattern intelligence from captured data."""
    rows = (db.query(DBAnswer, DBQuestion)
            .join(DBQuestion, DBAnswer.question_id == DBQuestion.id)
            .join(DBTestAttempt, DBAnswer.test_attempt_id == DBTestAttempt.id)
            .filter(DBTestAttempt.user_id == current_user.id)
            .all())
    total = len(rows)
    # Confidence calibration: per confidence level, correct vs wrong.
    conf = {c: {"correct": 0, "wrong": 0} for c in ["sure", "unsure", "guess", "untagged"]}
    reason_counts = {}
    careless = 0          # fast (<=8s) AND wrong
    fast_total = 0
    overconfident = 0     # sure but wrong
    lucky = 0             # guess but correct
    underselling = 0      # unsure but correct
    times = []
    for ans, q in rows:
        c = (ans.confidence or "untagged").lower()
        if c not in conf:
            c = "untagged"
        if ans.is_correct:
            conf[c]["correct"] += 1
        else:
            conf[c]["wrong"] += 1
        if c == "sure" and not ans.is_correct:
            overconfident += 1
        if c == "guess" and ans.is_correct:
            lucky += 1
        if c == "unsure" and ans.is_correct:
            underselling += 1
        if ans.time_taken is not None:
            times.append(ans.time_taken)
            if ans.time_taken <= 8:
                fast_total += 1
                if not ans.is_correct:
                    careless += 1
        if not ans.is_correct:
            r = ans.wrong_reason or _classify_mistake(ans, q)
            reason_counts[r] = reason_counts.get(r, 0) + 1
    # Guess quality: accuracy on guessed questions vs the 25% random baseline.
    g = conf["guess"]
    guess_attempted = g["correct"] + g["wrong"]
    guess_acc = round(g["correct"] / guess_attempted * 100, 1) if guess_attempted else None
    def _acc(d):
        n = d["correct"] + d["wrong"]
        return round(d["correct"] / n * 100, 1) if n else None
    calibration = [{"confidence": c, "correct": conf[c]["correct"], "wrong": conf[c]["wrong"],
                    "attempted": conf[c]["correct"] + conf[c]["wrong"], "accuracy": _acc(conf[c])}
                   for c in ["sure", "unsure", "guess", "untagged"]]
    avg_time = round(sum(times) / len(times), 1) if times else None
    return {
        "status": "success", "total_answers": total,
        "calibration": calibration,
        "flags": {
            "overconfident": overconfident,   # confidently wrong → real gaps
            "lucky_guesses": lucky,           # right by luck → don't trust these
            "underselling": underselling,     # unsure but right → trust yourself more
            "careless": careless,             # fast & wrong → slow down / read fully
        },
        "guess_quality": {"attempted": guess_attempted, "accuracy": guess_acc, "baseline": 25.0},
        "reason_breakdown": [{"reason": k, "count": v} for k, v in sorted(reason_counts.items(), key=lambda x: -x[1])],
        "avg_time_per_q": avg_time,
        "fast_answers": fast_total,
    }


# ── Phase 2: Coverage · Revision Plan · Readiness ─────────────────────────────
# Rough UPSC Prelims weight per subject (1=low, 2=med, 3=high), keyword-matched.
_SUBJECT_WEIGHTS = [
    (("polity", "constitution", "governance"), 3),
    (("economy", "economic"), 3),
    (("environment", "ecology", "biodiversity", "climate"), 3),
    (("geography",), 3),
    (("history", "art", "culture", "ancient", "medieval", "modern"), 3),
    (("current affairs", "current"), 3),
    (("science", "technology"), 2),
    (("international", "relations"), 1),
]

def _subject_weight(name):
    n = (name or "").lower()
    for keys, w in _SUBJECT_WEIGHTS:
        if any(k in n for k in keys):
            return w
    return 2

def _coverage_status(attempted, accuracy, days_since):
    if not attempted:
        return "untouched"
    if days_since is not None and days_since > 45:
        return "needs_revision"
    if accuracy >= 75 and attempted >= 15:
        return "strong"
    if attempted >= 15:
        return "covered"
    return "in_progress"

def _answer_rows(db, user_id):
    return (db.query(DBAnswer, DBQuestion, DBTestAttempt)
            .join(DBQuestion, DBAnswer.question_id == DBQuestion.id)
            .join(DBTestAttempt, DBAnswer.test_attempt_id == DBTestAttempt.id)
            .filter(DBTestAttempt.user_id == user_id)
            .all())

@app.get("/me/coverage", tags=["Coverage"])
def my_coverage(db: Session = Depends(get_db), current_user: DBUser = Depends(get_current_user)):
    """How much of the syllabus the student has actually practiced, and how well."""
    now = datetime.datetime.utcnow()
    rows = _answer_rows(db, current_user.id)
    by_subject, by_book = {}, {}
    for ans, q, att in rows:
        subj = q.subject or "General"
        s = by_subject.setdefault(subj, {"attempted": 0, "correct": 0, "topics": set(), "last": None})
        s["attempted"] += 1
        s["correct"] += 1 if ans.is_correct else 0
        if q.topic:
            s["topics"].add(q.topic)
        if att.completed_at and (s["last"] is None or att.completed_at > s["last"]):
            s["last"] = att.completed_at
        if q.book:
            b = by_book.setdefault(q.book, {"subject": subj, "attempted": 0, "correct": 0, "last": None})
            b["attempted"] += 1
            b["correct"] += 1 if ans.is_correct else 0
            if att.completed_at and (b["last"] is None or att.completed_at > b["last"]):
                b["last"] = att.completed_at
    canon = {item["subject"]: len(item.get("topics", [])) for item in syllabus_data.SUBJECT_TOPICS}
    subjects, matched_canon = [], set()
    for subj, s in by_subject.items():
        acc = round(s["correct"] / s["attempted"] * 100, 1) if s["attempted"] else 0
        days = (now - s["last"]).days if s["last"] else None
        tt = None
        for cs, ct in canon.items():
            cl, sl = cs.lower(), subj.lower()
            if cl == sl or cl in sl or sl in cl:
                tt = ct
                matched_canon.add(cs)
                break
        topics_touched = len(s["topics"])
        cov_pct = round(min(topics_touched, tt) / tt * 100, 1) if tt else None
        subjects.append({
            "subject": subj, "attempted": s["attempted"], "accuracy": acc,
            "topics_touched": topics_touched, "total_topics": tt, "coverage_pct": cov_pct,
            "days_since": days, "weight": _subject_weight(subj),
            "status": _coverage_status(s["attempted"], acc, days),
        })
    subjects.sort(key=lambda x: (-x["weight"], -x["attempted"]))
    gaps = [{"subject": cs, "total_topics": ct, "weight": _subject_weight(cs)}
            for cs, ct in canon.items() if cs not in matched_canon]
    gaps.sort(key=lambda x: -x["weight"])
    ncert_names = {b["book"] for b in syllabus_data.NCERT_BOOKS}
    ref_names = {b["book"] for b in syllabus_data.REFERENCE_BOOKS}
    books = []
    for bk, b in by_book.items():
        acc = round(b["correct"] / b["attempted"] * 100, 1) if b["attempted"] else 0
        days = (now - b["last"]).days if b["last"] else None
        kind = "NCERT" if bk in ncert_names else ("Reference" if bk in ref_names else "Other")
        books.append({"book": bk, "subject": b["subject"], "attempted": b["attempted"],
                      "accuracy": acc, "days_since": days, "kind": kind,
                      "status": _coverage_status(b["attempted"], acc, days)})
    books.sort(key=lambda x: -x["attempted"])
    practiced = set(by_book.keys())
    return {
        "status": "success", "subjects": subjects, "gaps": gaps, "books": books,
        "summary": {
            "subjects_practiced": len(subjects),
            "canon_subjects": len(canon), "canon_practiced": len(matched_canon),
            "ncert_practiced": len(practiced & ncert_names), "ncert_total": len(ncert_names),
            "ref_practiced": len(practiced & ref_names), "ref_total": len(ref_names),
        },
    }

@app.get("/me/revision-plan", tags=["Coverage"])
def my_revision_plan(db: Session = Depends(get_db), current_user: DBUser = Depends(get_current_user)):
    """Spaced-repetition calendar + a priority ranking of what to revise first."""
    now = datetime.datetime.utcnow()
    items = db.query(DBReviewItem).filter(DBReviewItem.user_id == current_user.id).all()
    overdue = today = this_week = later = mastered = 0
    for it in items:
        if it.mastered:
            mastered += 1
            continue
        if it.next_review is None:
            continue
        d = (it.next_review.date() - now.date()).days
        if d < 0:
            overdue += 1
        elif d == 0:
            today += 1
        elif d <= 7:
            this_week += 1
        else:
            later += 1
    # Priority = proficiency gap + forgetting + UPSC weight, per subject.
    rows = _answer_rows(db, current_user.id)
    agg = {}
    for ans, q, att in rows:
        subj = q.subject or "General"
        a = agg.setdefault(subj, {"attempted": 0, "correct": 0, "last": None})
        a["attempted"] += 1
        a["correct"] += 1 if ans.is_correct else 0
        if att.completed_at and (a["last"] is None or att.completed_at > a["last"]):
            a["last"] = att.completed_at
    priorities = []
    for subj, a in agg.items():
        if a["attempted"] < 3:
            continue
        acc = a["correct"] / a["attempted"] * 100
        days = (now - a["last"]).days if a["last"] else None
        weight = _subject_weight(subj)
        prof_gap = (100 - acc) / 100.0
        forgetting = min((days or 0) / 45.0, 1.0) if days is not None else 0.3
        score = round((0.4 * prof_gap + 0.35 * forgetting + 0.25 * ((weight - 1) / 2.0)) * 100, 1)
        priorities.append({"subject": subj, "score": score, "accuracy": round(acc, 1),
                           "days_since": days, "weight": weight, "attempted": a["attempted"]})
    priorities.sort(key=lambda x: -x["score"])
    return {
        "status": "success",
        "calendar": {"overdue": overdue, "today": today, "this_week": this_week,
                     "later": later, "mastered": mastered, "due_now": overdue + today},
        "priorities": priorities[:8],
    }

@app.get("/me/readiness", tags=["Coverage"])
def my_readiness(db: Session = Depends(get_db), current_user: DBUser = Depends(get_current_user)):
    """An honest 'if the exam were tomorrow' snapshot — a score RANGE, never a cutoff."""
    now = datetime.datetime.utcnow()
    rows = _answer_rows(db, current_user.id)
    total = len(rows)
    correct = sum(1 for ans, q, att in rows if ans.is_correct)
    wrong = total - correct
    overall_acc = round(correct / total * 100, 1) if total else 0
    # Score range from full attempts' net %, scaled to the 200-mark paper.
    attempts = db.query(DBTestAttempt).filter(DBTestAttempt.user_id == current_user.id).all()
    net_pcts = []
    for a in attempts:
        tq = a.mock_test.total_questions if a.mock_test else 0
        if not tq:
            continue
        att_correct = a.score or 0
        attempted = db.query(DBAnswer).filter(DBAnswer.test_attempt_id == a.id).count()
        att_wrong = max(0, attempted - att_correct)
        net = att_correct * 2.0 - att_wrong * (2.0 / 3.0)
        net_pcts.append(net / (tq * 2.0) * 100)
    score_range = None
    if net_pcts:
        avg = sum(net_pcts) / len(net_pcts)
        if len(net_pcts) > 1:
            var = sum((x - avg) ** 2 for x in net_pcts) / (len(net_pcts) - 1)
            spread = max(8.0, var ** 0.5)
        else:
            spread = 15.0
        center = avg / 100.0 * 200.0
        margin = spread / 100.0 * 200.0
        score_range = {"low": max(0, round(center - margin)), "high": min(200, round(center + margin)),
                       "mid": round(center), "tests_used": len(net_pcts)}
    # Per-subject strengths / focus areas.
    by_subject = {}
    for ans, q, att in rows:
        subj = q.subject or "General"
        s = by_subject.setdefault(subj, {"attempted": 0, "correct": 0})
        s["attempted"] += 1
        s["correct"] += 1 if ans.is_correct else 0
    strengths, focus = [], []
    for subj, s in by_subject.items():
        if s["attempted"] < 5:
            continue
        acc = round(s["correct"] / s["attempted"] * 100, 1)
        entry = {"subject": subj, "accuracy": acc, "attempted": s["attempted"], "weight": _subject_weight(subj)}
        if acc >= 65:
            strengths.append(entry)
        elif acc < 50:
            focus.append(entry)
    strengths.sort(key=lambda x: -x["accuracy"])
    focus.sort(key=lambda x: (x["accuracy"], -x["weight"]))
    # Readiness band from overall net %.
    net_total = correct * 2.0 - wrong * (2.0 / 3.0)
    net_pct = round(net_total / (total * 2.0) * 100, 1) if total else 0
    if net_pct >= 55:
        band = "Strong"
    elif net_pct >= 45:
        band = "Approaching"
    elif net_pct >= 33:
        band = "Developing"
    else:
        band = "Early"
    due_now = db.query(DBReviewItem).filter(
        DBReviewItem.user_id == current_user.id, DBReviewItem.mastered == False,
        DBReviewItem.next_review <= now).count()

    # ── Attempt-strategy advisor: built from the student's OWN confidence
    # calibration, so "should I attempt 85 or 100?" gets a personal answer.
    # Model: a 100-question paper mirroring their historical sure/unsure/guess
    # mix; expected net marks = n*(acc*2 - (1-acc)*2/3) per bucket attempted.
    strategy = None
    tagged = [(ans, q) for ans, q, att in rows
              if (ans.confidence or "").lower() in ("sure", "unsure", "guess")]
    if len(tagged) < 30:
        strategy = {"locked": True,
                    "note": ("Tap the confidence button after answers — once you have 30+ "
                             "tagged answers this becomes your personal attempt strategy. "
                             f"({len(tagged)}/30 so far)")}
    else:
        buckets = {}
        for ans, q in tagged:
            b = buckets.setdefault((ans.confidence or "").lower(), {"n": 0, "c": 0})
            b["n"] += 1
            b["c"] += 1 if ans.is_correct else 0
        n_all = len(tagged)
        mix = {}
        for k in ("sure", "unsure", "guess"):
            b = buckets.get(k, {"n": 0, "c": 0})
            mix[k] = {"share": round(b["n"] / n_all * 100),
                      "accuracy": round(b["c"] / b["n"] * 100, 1) if b["n"] else None,
                      "n": b["n"]}

        def _bucket_net(k):
            b = buckets.get(k, {"n": 0, "c": 0})
            if not b["n"]:
                return 0.0, 0.0
            share = b["n"] / n_all * 100.0            # questions of this kind per 100
            acc = b["c"] / b["n"]
            return share, share * (acc * 2.0 - (1.0 - acc) * (2.0 / 3.0))

        options = []
        for name, ks, note in (
                ("Conservative", ("sure",), "Attempt only what you're sure of; skip the rest."),
                ("Balanced", ("sure", "unsure"), "Attempt sure + unsure; skip pure guesses."),
                ("Aggressive", ("sure", "unsure", "guess"), "Attempt everything, guesses included.")):
            att = net = 0.0
            for k in ks:
                s, m = _bucket_net(k)
                att += s
                net += m
            options.append({"name": name, "attempt": round(att),
                            "expected_net": round(net), "note": note})
        best = max(options, key=lambda o: o["expected_net"])
        g = mix.get("guess") or {}
        gacc = g.get("accuracy")
        if gacc is None or (g.get("n") or 0) < 10:
            guess_verdict = "Not enough tagged guesses yet to judge your guessing."
        elif gacc > 30:
            guess_verdict = (f"Your guesses land {gacc}% of the time — above the 25% break-even, "
                             "so educated guessing ADDS marks for you.")
        elif gacc >= 20:
            guess_verdict = (f"Your guesses land {gacc}% — right at break-even. Guess only after "
                             "eliminating at least one option.")
        else:
            guess_verdict = (f"Your guesses land only {gacc}% — below break-even, so blind "
                             "guessing COSTS you marks. Skip when you can't eliminate options.")
        strategy = {"locked": False, "tagged_answers": n_all, "mix": mix,
                    "options": options, "recommended": best["name"],
                    "guess_verdict": guess_verdict,
                    "note": ("Based on your own confidence tags, not a generic rule. "
                             "It sharpens as you tag more answers.")}

    return {
        "status": "success",
        "has_data": total > 0,
        "total_answered": total, "overall_accuracy": overall_acc, "net_pct": net_pct,
        "band": band, "score_range": score_range,
        "strengths": strengths[:5], "focus_areas": focus[:5],
        "due_now": due_now,
        "strategy": strategy,
        "caveat": "A rough estimate from your own practice so far — not a prediction of the official cutoff.",
    }


# ── Interview module phase 1: DAF-based question generator ───────────────────
class DAFRequest(BaseModel):
    home_state: Optional[str] = None
    district: Optional[str] = None
    education: Optional[str] = None
    optional_subject: Optional[str] = None
    work_experience: Optional[str] = None
    hobbies: Optional[str] = None
    achievements: Optional[str] = None
    languages: Optional[str] = None
    extra: Optional[str] = None            # anything else the aspirant wants covered


def _daf_from_profile(req: "DAFRequest", prof) -> dict:
    """Merge the request with StudentProfile defaults so the aspirant doesn't
    retype what the app already knows."""
    def pick(v, fallback):
        v = (v or "").strip()
        return v if v else ((fallback or "").strip() if fallback else "")
    p = prof or type("X", (), {})()
    return {
        "home_state": pick(req.home_state, getattr(p, "home_state", None)),
        "district": pick(req.district, getattr(p, "district", None)),
        "education": pick(req.education, " / ".join(x for x in [
            getattr(p, "education", None), getattr(p, "graduation_stream", None),
            getattr(p, "additional_qualification", None)] if x)),
        "optional_subject": pick(req.optional_subject, getattr(p, "optional_subject", None)),
        "work_experience": pick(req.work_experience, getattr(p, "work_experience", None)),
        "hobbies": pick(req.hobbies, None),
        "achievements": pick(req.achievements, None),
        "languages": pick(req.languages, getattr(p, "mains_language", None)),
        "extra": (req.extra or "").strip(),
    }


@app.get("/me/interview", tags=["Interview"])
def my_interview(db: Session = Depends(get_db), current_user: DBUser = Depends(get_current_user)):
    """The aspirant's saved DAF details + generated interview question bank,
    plus profile-based prefill for the form when nothing is saved yet."""
    row = db.query(DBInterviewPrep).filter(DBInterviewPrep.user_id == current_user.id).first()
    prof = db.query(DBStudentProfile).filter(DBStudentProfile.user_id == current_user.id).first()
    prefill = _daf_from_profile(DAFRequest(), prof)
    if not row or not row.questions:
        return {"status": "success", "has_data": False, "prefill": prefill}
    try:
        daf = json.loads(row.daf or "{}")
    except Exception:
        daf = {}
    try:
        questions = json.loads(row.questions or "{}")
    except Exception:
        questions = {}
    return {"status": "success", "has_data": True, "daf": daf, "questions": questions,
            "updated_at": row.updated_at.isoformat() if row.updated_at else None,
            "prefill": prefill}


@app.post("/me/interview/generate", tags=["Interview"])
def generate_interview_questions(req: DAFRequest, db: Session = Depends(get_db),
                                 current_user: DBUser = Depends(get_current_user)):
    """Generate a personalised UPSC Personality Test question bank from the
    aspirant's DAF-style details (the way boards actually probe a DAF: hometown,
    education, optional, hobbies, work, and their current-affairs links).
    Regenerates in place; the previous bank is replaced."""
    import gemini_service
    import re as _re
    prof = db.query(DBStudentProfile).filter(DBStudentProfile.user_id == current_user.id).first()
    daf = _daf_from_profile(req, prof)
    filled = {k: v for k, v in daf.items() if v}
    if len(filled) < 2:
        raise HTTPException(status_code=400,
                            detail="Fill in at least two DAF fields (e.g. home state and education).")
    daf_text = "\n".join(f"- {k.replace('_', ' ').title()}: {v}" for k, v in filled.items())
    system = (
        "You are a senior UPSC Personality Test board member who has interviewed "
        "hundreds of candidates. From a candidate's DAF details you predict the "
        "questions a real board would ask. Be specific to THIS candidate's details "
        "- never generic filler. Mix factual, opinion, situational and follow-up "
        "style questions; include current-affairs angles connected to their "
        "background. For each question add a short preparation hint (what a good "
        "answer touches, 10-25 words, no model answers). Balanced, respectful, "
        "no political bias. STRICT JSON only.")
    schema = ('{"themes":[{"title":str,"why":str,'
              '"questions":[{"q":str,"hint":str}]}]}')
    prompt = (
        f"Candidate's DAF details:\n{daf_text}\n\n"
        "Generate a personalised interview question bank: 7-10 themes (e.g. home "
        "district/state, education background, optional subject, work experience, "
        "hobbies, current affairs tied to their profile, situational/ethics, "
        "service motivation). 6-8 questions per theme, each with a hint. "
        f"Return STRICT JSON exactly in this shape: {schema}")
    # Interview question bank is high-stakes + low-volume → Gemini PRO primary,
    # DeepSeek as automatic fallback.
    raw = gemini_service.gen_text(prompt, json_mode=True, prefer="gemini",
                                  model=gemini_service.MODEL_PRO, system=system)
    m = _re.search(r"\{.*\}", raw or "", _re.DOTALL)
    if not m:
        raise HTTPException(status_code=502, detail="AI returned no usable question bank - try again.")
    try:
        bank = json.loads(m.group(0))
        themes = bank.get("themes") or []
        assert isinstance(themes, list) and themes
    except Exception:
        raise HTTPException(status_code=502, detail="AI returned malformed data - try again.")
    # sanitise: keep only expected keys, cap sizes
    clean = []
    for t in themes[:12]:
        qs = [{"q": str(q.get("q", "")).strip(), "hint": str(q.get("hint", "")).strip()}
              for q in (t.get("questions") or [])[:10] if str(q.get("q", "")).strip()]
        if qs:
            clean.append({"title": str(t.get("title", "")).strip() or "Theme",
                          "why": str(t.get("why", "")).strip(), "questions": qs})
    if not clean:
        raise HTTPException(status_code=502, detail="AI returned an empty question bank - try again.")
    payload = {"themes": clean}
    row = db.query(DBInterviewPrep).filter(DBInterviewPrep.user_id == current_user.id).first()
    if not row:
        row = DBInterviewPrep(user_id=current_user.id)
        db.add(row)
    row.daf = json.dumps(daf, ensure_ascii=False)
    row.questions = json.dumps(payload, ensure_ascii=False)
    row.updated_at = datetime.datetime.utcnow()
    db.commit()
    total_q = sum(len(t["questions"]) for t in clean)
    return {"status": "success", "questions": payload, "daf": daf,
            "themes": len(clean), "total_questions": total_q}


# ── Test Tracker: logged mock scores, trends, goals & exam countdown ──────────
class MockScoreIn(BaseModel):
    stage: Optional[str] = "prelims"
    test_name: Optional[str] = None
    series: Optional[str] = None
    taken_on: Optional[str] = None          # ISO date; defaults to today
    max_marks: Optional[float] = 200
    score: Optional[float] = 0
    total_q: Optional[int] = None
    correct: Optional[int] = None
    wrong: Optional[int] = None
    unattempted: Optional[int] = None
    accuracy: Optional[float] = None         # auto-computed from correct/wrong if omitted
    weak_areas: Optional[str] = None
    notes: Optional[str] = None


class ExamGoalIn(BaseModel):
    target_score: Optional[int] = None
    target_accuracy: Optional[int] = None
    prelims_date: Optional[str] = None
    mains_date: Optional[str] = None


def _serialize_mock(m):
    return {"id": m.id, "stage": m.stage, "test_name": m.test_name, "series": m.series,
            "taken_on": m.taken_on, "max_marks": m.max_marks, "score": m.score,
            "total_q": m.total_q, "correct": m.correct, "wrong": m.wrong,
            "unattempted": m.unattempted, "accuracy": m.accuracy,
            "weak_areas": m.weak_areas, "notes": m.notes}


def _mock_stats(rows):
    """Summary KPIs + an ordered trend series for one stage's logged tests."""
    if not rows:
        return {"tests": 0, "avg_score": None, "best_score": None,
                "avg_accuracy": None, "last_mistakes": None, "series": []}
    scores = [r.score or 0 for r in rows]
    accs = [r.accuracy for r in rows if r.accuracy is not None]
    last = rows[-1]
    series = [{"date": r.taken_on or "",
               "label": (r.test_name or r.taken_on or f"#{i+1}"),
               "score": round(r.score or 0, 1),
               "accuracy": round(r.accuracy, 1) if r.accuracy is not None else 0}
              for i, r in enumerate(rows)]
    return {"tests": len(rows),
            "avg_score": round(sum(scores) / len(scores), 1),
            "best_score": round(max(scores), 1),
            "avg_accuracy": round(sum(accs) / len(accs), 1) if accs else None,
            "last_mistakes": last.wrong if last.wrong is not None else None,
            "series": series}


def _days_until(iso):
    try:
        return (datetime.date.fromisoformat(str(iso)) - datetime.date.today()).days
    except Exception:
        return None


@app.get("/me/mocks", tags=["Tracker"])
def my_mocks(db: Session = Depends(get_db), current_user: DBUser = Depends(get_current_user)):
    """Everything the Test Tracker panel needs: the logged tests, per-stage
    analytics (avg/best score, accuracy, trend series), the user's targets, and
    the days-to-exam countdown."""
    rows = db.query(DBMockScore).filter(DBMockScore.user_id == current_user.id).all()
    rows.sort(key=lambda r: (r.taken_on or "", r.id))     # oldest-first for trend lines
    by_stage = {st: _mock_stats([r for r in rows if (r.stage or "prelims") == st])
                for st in ("prelims", "csat", "mains")}
    goal = db.query(DBExamGoal).filter(DBExamGoal.user_id == current_user.id).first()
    goals = {"target_score": goal.target_score if goal else None,
             "target_accuracy": goal.target_accuracy if goal else None,
             "prelims_date": goal.prelims_date if goal else None,
             "mains_date": goal.mains_date if goal else None}
    countdown = {"prelims": _days_until(goals["prelims_date"]) if goals["prelims_date"] else None,
                 "mains": _days_until(goals["mains_date"]) if goals["mains_date"] else None}
    return {"status": "success",
            "mocks": [_serialize_mock(m) for m in reversed(rows)],   # newest-first for the list
            "by_stage": by_stage, "goals": goals, "countdown": countdown}


@app.post("/me/mocks", tags=["Tracker"])
def add_mock(req: MockScoreIn, db: Session = Depends(get_db),
             current_user: DBUser = Depends(get_current_user)):
    """Log one mock/test-series score. Accuracy is auto-derived from correct/wrong
    when not supplied."""
    stage = (req.stage or "prelims").lower()
    if stage not in ("prelims", "csat", "mains"):
        stage = "prelims"
    acc = req.accuracy
    if acc is None and req.correct is not None and req.wrong is not None:
        att = (req.correct or 0) + (req.wrong or 0)
        acc = round((req.correct or 0) / att * 100, 1) if att else None
    taken = (req.taken_on or "").strip() or datetime.date.today().isoformat()
    m = DBMockScore(
        user_id=current_user.id, stage=stage,
        test_name=(req.test_name or "").strip() or None,
        series=(req.series or "").strip() or None,
        taken_on=taken,
        max_marks=req.max_marks if req.max_marks is not None else 200,
        score=req.score or 0, total_q=req.total_q, correct=req.correct,
        wrong=req.wrong, unattempted=req.unattempted, accuracy=acc,
        weak_areas=(req.weak_areas or "").strip() or None,
        notes=(req.notes or "").strip() or None)
    db.add(m)
    db.commit()
    db.refresh(m)
    return {"status": "success", "mock": _serialize_mock(m)}


@app.delete("/me/mocks/{mock_id}", tags=["Tracker"])
def delete_mock(mock_id: int, db: Session = Depends(get_db),
                current_user: DBUser = Depends(get_current_user)):
    m = db.query(DBMockScore).filter(DBMockScore.id == mock_id,
                                     DBMockScore.user_id == current_user.id).first()
    if not m:
        raise HTTPException(status_code=404, detail="Test not found.")
    db.delete(m)
    db.commit()
    return {"status": "success"}


@app.post("/me/exam-goals", tags=["Tracker"])
def set_exam_goals(req: ExamGoalIn, db: Session = Depends(get_db),
                   current_user: DBUser = Depends(get_current_user)):
    """Upsert the aspirant's target score/accuracy and exam dates."""
    row = db.query(DBExamGoal).filter(DBExamGoal.user_id == current_user.id).first()
    if not row:
        row = DBExamGoal(user_id=current_user.id)
        db.add(row)
    row.target_score = req.target_score
    row.target_accuracy = req.target_accuracy
    row.prelims_date = (req.prelims_date or "").strip() or None
    row.mains_date = (req.mains_date or "").strip() or None
    row.updated_at = datetime.datetime.utcnow()
    db.commit()
    return {"status": "success"}


# ── Phase 3: Exam Simulator (mock composer) ───────────────────────────────────
class SimulatorCompose(BaseModel):
    mode: Optional[str] = "balanced"       # balanced | weak | year
    num_questions: Optional[int] = 100
    year: Optional[str] = None
    duration_minutes: Optional[int] = None
    paper: Optional[str] = "gs"            # gs (General Studies) | csat

@app.post("/me/simulator/compose", tags=["Simulator"])
def simulator_compose(req: SimulatorCompose, db: Session = Depends(get_db),
                      current_user: DBUser = Depends(get_current_user)):
    """Assemble a full-length exam-simulation paper from the VERIFIED PYQ bank.
    Modes: balanced (natural subject mix), weak (60% from your weak subjects),
    year (a specific past paper). Real exam questions only — no AI generation."""
    mode = (req.mode or "balanced").lower()
    paper = (req.paper or "gs").lower()

    # ── CSAT full paper (Paper II) — draw from the AIVORA CSAT bank ─────────────
    if paper == "csat":
        csat_pool = []
        for a in CSAT_AREAS:
            for q in a.get("questions", []):
                qq = dict(q); qq["_area"] = a.get("name", "CSAT")
                csat_pool.append(qq)
        if not csat_pool:
            raise HTTPException(status_code=503, detail="CSAT question bank is unavailable.")
        # Real CSAT = 80 questions / 120 min. Use up to that, capped at what we have.
        target = min(80, len(csat_pool))
        selected = random.sample(csat_pool, target)
        duration = req.duration_minutes or max(10, round(target * 1.5))
        title = "🧪 Exam Simulator • CSAT (Paper II)"
        db_test = DBMockTest(
            title=title, description="Full-length CSAT simulation (Paper II)",
            subject="CSAT", total_questions=len(selected),
            duration_minutes=duration, user_id=current_user.id,
        )
        db.add(db_test); db.commit(); db.refresh(db_test)
        for q in selected:
            db.add(DBQuestion(
                text=q["text"], option_a=q["option_a"], option_b=q["option_b"],
                option_c=q["option_c"], option_d=q["option_d"],
                correct_answer=q["correct_answer"], explanation=q.get("explanation", ""),
                subject="CSAT", topic=q.get("_area"), question_type="csat",
                mock_test_id=db_test.id,
            ))
        db.commit()
        return {
            "status": "success", "mock_test_id": db_test.id, "questions": len(selected),
            "duration_minutes": duration, "mode": "csat", "paper": "csat",
            "distribution": [{"subject": "CSAT", "count": len(selected)}],
        }

    num = max(5, min(int(req.num_questions or 100), 150))
    pool = PYQ_QUESTIONS
    if not pool:
        raise HTTPException(status_code=503, detail="Question bank is unavailable.")

    if mode == "year" and req.year:
        yr = str(req.year).strip()
        ypool = [q for q in pool if str(q.get("year")) == yr]
        if not ypool:
            raise HTTPException(status_code=404, detail="No questions for that year.")
        selected = random.sample(ypool, min(num, len(ypool)))
    elif mode == "weak":
        rows = _answer_rows(db, current_user.id)
        agg = {}
        for ans, q, att in rows:
            s = q.subject or "General"
            a = agg.setdefault(s, {"a": 0, "c": 0})
            a["a"] += 1
            a["c"] += 1 if ans.is_correct else 0
        weak = {s for s, a in agg.items() if a["a"] >= 3 and (a["c"] / a["a"] * 100) < 55}
        weak_pool = [q for q in pool if q.get("subject") in weak]
        rest_pool = [q for q in pool if q.get("subject") not in weak]
        target_weak = int(num * 0.6)
        sel_weak = random.sample(weak_pool, min(target_weak, len(weak_pool))) if weak_pool else []
        remaining = num - len(sel_weak)
        sel_rest = random.sample(rest_pool, min(remaining, len(rest_pool))) if rest_pool else []
        selected = sel_weak + sel_rest
        random.shuffle(selected)
        if not selected:
            selected = random.sample(pool, min(num, len(pool)))
    else:
        mode = "balanced"
        selected = random.sample(pool, min(num, len(pool)))

    dist = {}
    for q in selected:
        s = q.get("subject") or "General"
        dist[s] = dist.get(s, 0) + 1
    duration = req.duration_minutes or max(10, round(len(selected) * 1.2))
    title = "🧪 Exam Simulator • " + mode.title() + (f" • {req.year}" if mode == "year" and req.year else "")
    db_test = DBMockTest(
        title=title, description="Full-length exam simulation (verified PYQs)",
        subject="UPSC Prelims (Mixed)", total_questions=len(selected),
        duration_minutes=duration, user_id=current_user.id,
    )
    db.add(db_test); db.commit(); db.refresh(db_test)
    for q in selected:
        db.add(DBQuestion(
            text=q["text"], option_a=q["option_a"], option_b=q["option_b"],
            option_c=q["option_c"], option_d=q["option_d"],
            correct_answer=q["correct_answer"], explanation=q.get("explanation", ""),
            subject=q.get("subject"), mock_test_id=db_test.id,
        ))
    db.commit()
    return {
        "status": "success", "mock_test_id": db_test.id, "questions": len(selected),
        "duration_minutes": duration, "mode": mode,
        "distribution": [{"subject": k, "count": v} for k, v in sorted(dist.items(), key=lambda x: -x[1])],
    }

@app.get("/me/simulator/history", tags=["Simulator"])
def simulator_history(db: Session = Depends(get_db), current_user: DBUser = Depends(get_current_user)):
    """Net-score across past simulator mocks — your improvement curve over attempts."""
    attempts = (db.query(DBTestAttempt).join(DBMockTest, DBTestAttempt.mock_test_id == DBMockTest.id)
                .filter(DBTestAttempt.user_id == current_user.id,
                        DBMockTest.title.like("%Exam Simulator%"))
                .order_by(DBTestAttempt.completed_at.asc()).all())
    hist = []
    for a in attempts:
        tq = a.mock_test.total_questions if a.mock_test else 0
        if not tq:
            continue
        c = a.score or 0
        attempted = db.query(DBAnswer).filter(DBAnswer.test_attempt_id == a.id).count()
        w = max(0, attempted - c)
        net = c * 2 - w * (2 / 3)
        hist.append({"attempt_id": a.id, "date": a.completed_at.isoformat() if a.completed_at else None,
                     "net": round(net, 1), "net_pct": round(net / (tq * 2) * 100, 1),
                     "score": c, "total": tq, "attempted": attempted})
    return {"status": "success", "count": len(hist), "history": hist}


# ── Phase 3 follow-ups: War Room + Adaptive Difficulty ────────────────────────
def _net_band(net_pct):
    if net_pct >= 55:
        return "Strong"
    if net_pct >= 45:
        return "Approaching"
    if net_pct >= 33:
        return "Developing"
    return "Early"

@app.get("/me/warroom", tags=["Coverage"])
def warroom(db: Session = Depends(get_db), current_user: DBUser = Depends(get_current_user)):
    """Final-stretch command center: countdown, readiness, focus fire, and a plan."""
    now = datetime.datetime.utcnow()
    prof = db.query(DBStudentProfile).filter(DBStudentProfile.user_id == current_user.id).first()
    days_to_exam, exam_label, target_year = None, None, None
    if prof and prof.target_year and str(prof.target_year).strip().isdigit():
        yr = int(prof.target_year)
        target_year = yr
        try:
            exam_date = datetime.date(yr, 5, 25)
            delta = (exam_date - now.date()).days
            if delta >= 0:
                days_to_exam = delta
                exam_label = f"~25 May {yr} (Prelims, approx.)"
        except Exception:
            pass
    rows = _answer_rows(db, current_user.id)
    total = len(rows)
    correct = sum(1 for ans, q, att in rows if ans.is_correct)
    wrong = total - correct
    net_pct = round((correct * 2 - wrong * (2 / 3)) / (total * 2) * 100, 1) if total else 0
    by_sub = {}
    for ans, q, att in rows:
        s = q.subject or "General"
        d = by_sub.setdefault(s, {"a": 0, "c": 0})
        d["a"] += 1
        d["c"] += 1 if ans.is_correct else 0
    focus = []
    for s, d in by_sub.items():
        if d["a"] < 3:
            continue
        acc = round(d["c"] / d["a"] * 100, 1)
        if acc < 60:
            focus.append({"subject": s, "accuracy": acc, "weight": _subject_weight(s), "attempted": d["a"]})
    focus.sort(key=lambda x: (x["accuracy"], -x["weight"]))
    due = db.query(DBReviewItem).filter(
        DBReviewItem.user_id == current_user.id, DBReviewItem.mastered == False,
        DBReviewItem.next_review <= now).count()
    mistakes = (db.query(DBAnswer.question_id)
                .join(DBTestAttempt, DBAnswer.test_attempt_id == DBTestAttempt.id)
                .filter(DBTestAttempt.user_id == current_user.id, DBAnswer.is_correct == False)
                .distinct().count())
    mastered = db.query(DBReviewItem).filter(
        DBReviewItem.user_id == current_user.id, DBReviewItem.mastered == True).count()
    return {
        "status": "success", "has_data": total > 0,
        "days_to_exam": days_to_exam, "exam_label": exam_label, "target_year": target_year,
        "net_pct": net_pct, "band": _net_band(net_pct),
        "focus": focus[:5], "due_revisions": due, "mistakes": mistakes, "mastered": mastered,
    }

@app.get("/me/adaptive/level", tags=["Coverage"])
def adaptive_level(subject: Optional[str] = None, topic: Optional[str] = None,
                   db: Session = Depends(get_db), current_user: DBUser = Depends(get_current_user)):
    """Recommend a difficulty for the next drill from rolling accuracy. With no
    subject, it picks the student's weakest subject to target."""
    rows = _answer_rows(db, current_user.id)
    chosen_subject, chosen_topic = subject, topic
    c = a = 0
    if subject:
        for ans, q, att in rows:
            if (q.subject or "") == subject and (not topic or (q.topic or "") == topic):
                a += 1
                c += 1 if ans.is_correct else 0
    else:
        by_sub = {}
        for ans, q, att in rows:
            s = q.subject or "General"
            d = by_sub.setdefault(s, {"a": 0, "c": 0})
            d["a"] += 1
            d["c"] += 1 if ans.is_correct else 0
        cand = [(s, d) for s, d in by_sub.items() if d["a"] >= 3]
        if cand:
            cand.sort(key=lambda x: x[1]["c"] / x[1]["a"])
            chosen_subject = cand[0][0]
            c, a = cand[0][1]["c"], cand[0][1]["a"]
    acc = round(c / a * 100, 1) if a else None
    if not a or a < 3 or acc is None:
        level, reason = "medium", "Not enough history yet — starting at medium difficulty."
    elif acc < 45:
        level, reason = "easy", f"You're at {acc}% here — locking in fundamentals with easier questions first."
    elif acc < 72:
        level, reason = "medium", f"You're at {acc}% — medium difficulty to push you upward."
    else:
        level, reason = "hard", f"Strong at {acc}% — stepping up to hard, exam-tough questions."
    return {"status": "success", "subject": chosen_subject, "topic": chosen_topic,
            "accuracy": acc, "attempted": a, "level": level, "reason": reason}

@app.get("/me/adaptive/weak-topics", tags=["Coverage"])
def weak_topics(min_attempts: int = 3, top_k: int = 8,
                db: Session = Depends(get_db), current_user: DBUser = Depends(get_current_user)):
    """Finer-grained weakness detection than /me/adaptive/level: ranks the
    student's weakest TOPICS (not just subjects), each with a recommended
    difficulty, so practice can target the exact gaps. Feed a returned
    subject+topic into /mock-tests/generate-verified to drill it."""
    rows = _answer_rows(db, current_user.id)
    agg = {}
    for ans, q, att in rows:
        top = (q.topic or "").strip()
        if not top:
            continue
        d = agg.setdefault((q.subject or "General", top), {"a": 0, "c": 0})
        d["a"] += 1
        d["c"] += 1 if ans.is_correct else 0
    items = []
    for (subj, top), d in agg.items():
        if d["a"] < max(1, min_attempts):
            continue
        acc = round(d["c"] / d["a"] * 100, 1)
        level = "easy" if acc < 45 else ("medium" if acc < 72 else "hard")
        items.append({"subject": subj, "topic": top, "accuracy": acc,
                      "attempted": d["a"], "recommended_level": level})
    items.sort(key=lambda x: (x["accuracy"], -x["attempted"]))
    weak = [x for x in items if x["accuracy"] < 72][:top_k] or items[:top_k]
    return {"status": "success", "count": len(weak),
            "focus_topics": [w["topic"] for w in weak], "weak_topics": weak}

@app.get("/me/trends", tags=["Coverage"])
def my_trends(db: Session = Depends(get_db), current_user: DBUser = Depends(get_current_user)):
    """Progress over time: net-score trajectory, study streak, weekly activity, subject momentum."""
    from collections import Counter
    now = datetime.datetime.utcnow()
    td = datetime.timedelta
    rows = _answer_rows(db, current_user.id)
    by_att, answer_dates, subj_chrono = {}, [], {}
    for ans, q, att in rows:
        a = by_att.setdefault(att.id, {"att": att, "c": 0, "n": 0})
        a["c"] += 1 if ans.is_correct else 0
        a["n"] += 1
        if att.completed_at:
            answer_dates.append(att.completed_at.date())
            subj_chrono.setdefault(q.subject or "General", []).append((att.completed_at, ans.is_correct))
    series = []
    for aid, a in by_att.items():
        att = a["att"]; n = a["n"]; c = a["c"]; wrong = n - c
        total = att.mock_test.total_questions if att.mock_test else n
        net_pct = round((c * 2 - wrong * (2 / 3)) / (total * 2) * 100, 1) if total else 0
        series.append({
            "date": att.completed_at.isoformat() if att.completed_at else None,
            "ts": att.completed_at.timestamp() if att.completed_at else 0,
            "net_pct": net_pct, "accuracy": round(c / n * 100, 1) if n else 0,
            "score": c, "attempted": n,
            "title": att.mock_test.title if att.mock_test else "Test",
        })
    series.sort(key=lambda x: x["ts"])
    nets = [s["net_pct"] for s in series]
    improvement = None
    if len(nets) >= 4:
        k = max(1, len(nets) // 3)
        improvement = {"early": round(sum(nets[:k]) / k, 1), "recent": round(sum(nets[-k:]) / k, 1)}
    dayset = sorted(set(answer_dates))
    cur = 0
    if dayset:
        s = set(dayset)
        if dayset[-1] in (now.date(), now.date() - td(1)):
            check = dayset[-1]
            while check in s:
                cur += 1; check = check - td(1)
    longest = run = 0; prev = None
    for d in dayset:
        run = run + 1 if (prev is not None and (d - prev).days == 1) else 1
        longest = max(longest, run); prev = d
    cnt = Counter(answer_dates)
    weekly = [{"date": (now.date() - td(i)).isoformat(), "count": cnt.get(now.date() - td(i), 0)} for i in range(13, -1, -1)]
    subj_trend = []
    for s, lst in subj_chrono.items():
        if len(lst) < 6:
            continue
        lst.sort(key=lambda x: x[0])
        half = len(lst) // 2
        old, new = lst[:half], lst[half:]
        oacc = round(sum(1 for _, c in old if c) / len(old) * 100, 1)
        nacc = round(sum(1 for _, c in new if c) / len(new) * 100, 1)
        subj_trend.append({"subject": s, "old": oacc, "recent": nacc, "delta": round(nacc - oacc, 1)})
    subj_trend.sort(key=lambda x: -abs(x["delta"]))
    return {
        "status": "success", "series": series[-30:], "improvement": improvement,
        "streak": {"current": cur, "longest": longest, "active_days": len(dayset)},
        "weekly": weekly, "subject_trend": subj_trend[:8], "total_tests": len(series),
    }


# ── Smarter engine: flags · bookmarks · resolved · item difficulty ────────────
class FlagIn(BaseModel):
    reason: Optional[str] = "other"
    note: Optional[str] = None

@app.post("/me/questions/{question_id}/flag", tags=["Learning"])
def flag_question(question_id: int, payload: FlagIn, db: Session = Depends(get_db),
                  current_user: DBUser = Depends(get_current_user)):
    if not db.query(DBQuestion).filter(DBQuestion.id == question_id).first():
        raise HTTPException(status_code=404, detail="Question not found")
    reason = (payload.reason or "other").lower()
    if reason not in ("wrong_answer", "unclear", "wrong_subject", "outdated", "other"):
        reason = "other"
    existing = db.query(DBQuestionFlag).filter(
        DBQuestionFlag.user_id == current_user.id, DBQuestionFlag.question_id == question_id).first()
    if existing:
        existing.reason = reason; existing.note = (payload.note or "")[:500]
    else:
        db.add(DBQuestionFlag(user_id=current_user.id, question_id=question_id,
                              reason=reason, note=(payload.note or "")[:500]))
    db.commit()
    return {"status": "success", "question_id": question_id, "reason": reason}

@app.post("/me/bookmarks/{question_id}", tags=["Learning"])
def toggle_bookmark(question_id: int, db: Session = Depends(get_db),
                    current_user: DBUser = Depends(get_current_user)):
    if not db.query(DBQuestion).filter(DBQuestion.id == question_id).first():
        raise HTTPException(status_code=404, detail="Question not found")
    bm = db.query(DBBookmark).filter(
        DBBookmark.user_id == current_user.id, DBBookmark.question_id == question_id).first()
    if bm:
        db.delete(bm); db.commit()
        return {"status": "success", "bookmarked": False}
    db.add(DBBookmark(user_id=current_user.id, question_id=question_id)); db.commit()
    return {"status": "success", "bookmarked": True}

@app.get("/me/bookmarks", tags=["Learning"])
def list_bookmarks(db: Session = Depends(get_db), current_user: DBUser = Depends(get_current_user)):
    bms = (db.query(DBBookmark, DBQuestion).join(DBQuestion, DBBookmark.question_id == DBQuestion.id)
           .filter(DBBookmark.user_id == current_user.id).order_by(DBBookmark.id.desc()).all())
    items = [{
        "question_id": q.id, "question_text": q.text,
        "options": {l: getattr(q, "option_" + l.lower(), None) for l in ["A", "B", "C", "D"]},
        "correct_answer": q.correct_answer, "explanation": q.explanation or "",
        "subject": q.subject or "General", "topic": q.topic or "",
    } for bm, q in bms]
    return {"status": "success", "count": len(items), "bookmarks": items}


# ══════════════════════════════════════════════════════════════════════════════
#  AIVORA Mastery Loop (AML) — Phase 0 capture
#  Logs every adaptive-flow answer to concept_attempts and moves the student
#  model (concept_mastery + skill_mastery). No selection/teaching logic here.
# ══════════════════════════════════════════════════════════════════════════════

def _aml_difficulty_gain(difficulty, correct):
    d = (difficulty or "medium").lower()
    if correct:
        return {"easy": 0.6, "medium": 1.0, "hard": 1.4}.get(d, 1.0)
    return {"easy": 1.4, "medium": 1.0, "hard": 0.6}.get(d, 1.0)

def _aml_update_mastery(old, correct, difficulty, confidence, base_rate=0.25):
    """Difficulty-weighted exponential update (AML spec §4)."""
    old = 0.5 if old is None else float(old)
    alpha = base_rate * _aml_difficulty_gain(difficulty, correct)
    if not correct:
        if confidence == "sure":
            alpha *= 1.3          # confident-and-wrong = likely misconception → drop harder
        elif confidence == "guess":
            alpha *= 0.7          # guess-and-wrong → softer
    target = 1.0 if correct else 0.0
    return max(0.0, min(1.0, old + alpha * (target - old)))

def _aml_state(m):
    if m is None:
        return "unknown"
    if m < 0.2:  return "introduced"
    if m < 0.5:  return "recall"
    if m < 0.75: return "apply"
    if m < 0.9:  return "mastered"
    return "retained"

def _aml_record(db, user_id, q, correct, selected, response_ms, confidence, context, session_id=None):
    """Shared capture: log one ConceptAttempt and move concept + skill mastery.
    Adds to the session but does NOT commit — the caller owns the transaction."""
    from sqlalchemy import func as _func
    ck = (q.concept_key or q.topic or q.chapter or "general")
    subject = q.subject
    subtopic = q.chapter or q.topic
    pattern = q.pattern or q.question_type
    conf = (confidence or "").strip().lower() or None
    prior = (db.query(_func.count(DBConceptAttempt.id))
             .filter(DBConceptAttempt.user_id == user_id,
                     DBConceptAttempt.question_id == q.id).scalar()) or 0
    db.add(DBConceptAttempt(
        user_id=user_id, question_id=q.id, concept_key=ck, subject=subject, subtopic=subtopic,
        pattern=pattern, correct=bool(correct), selected=selected, difficulty=q.difficulty,
        response_ms=response_ms, confidence=conf, attempt_number=prior + 1, exposure_count=prior + 1,
        attempt_context=context, session_id=session_id))
    cm = (db.query(DBConceptMastery)
          .filter(DBConceptMastery.user_id == user_id,
                  DBConceptMastery.concept_key == ck).first())
    if not cm:
        cm = DBConceptMastery(user_id=user_id, concept_key=ck, subject=subject,
                              subtopic=subtopic, mastery=0.5)
        db.add(cm)
    cm.mastery = _aml_update_mastery(cm.mastery, correct, q.difficulty, conf)
    cm.attempts = (cm.attempts or 0) + 1
    cm.correct = (cm.correct or 0) + (1 if correct else 0)
    cm.confidence_n = (cm.confidence_n or 0) + 1
    cm.streak = ((cm.streak or 0) + 1) if correct else 0
    cm.last_seen = datetime.datetime.utcnow()
    cm.state = _aml_state(cm.mastery)
    if pattern:
        sm = (db.query(DBSkillMastery)
              .filter(DBSkillMastery.user_id == user_id,
                      DBSkillMastery.pattern == pattern).first())
        if not sm:
            sm = DBSkillMastery(user_id=user_id, pattern=pattern, mastery=0.5)
            db.add(sm)
        sm.mastery = _aml_update_mastery(sm.mastery, correct, q.difficulty, conf)
        sm.attempts = (sm.attempts or 0) + 1

class AttemptIn(BaseModel):
    question_id: int
    selected: Optional[str] = None          # A|B|C|D
    correct: Optional[bool] = None          # if omitted, derived from selected vs answer
    response_ms: Optional[int] = None
    confidence: Optional[str] = None        # 'sure' | 'somewhat' | 'guess'
    hint_used: Optional[bool] = False
    attempt_context: Optional[str] = "practice"   # diagnostic|practice|revision|mock|assessment
    session_id: Optional[str] = None

@app.post("/me/attempt", tags=["Learning"])
def record_attempt(payload: AttemptIn, db: Session = Depends(get_db),
                   current_user: DBUser = Depends(get_current_user)):
    """Record one answer and update the student model. Idempotent-safe to call
    once per served question."""
    from sqlalchemy import func as _func
    q = db.query(DBQuestion).filter(DBQuestion.id == payload.question_id).first()
    if not q:
        raise HTTPException(status_code=404, detail="Question not found")

    sel = (payload.selected or "").strip().upper()[:1] or None
    ans = (q.correct_answer or "").strip().upper()[:1] or None
    if payload.correct is not None:
        correct = bool(payload.correct)
    elif sel and ans:
        correct = (sel == ans)
    else:
        correct = False

    ck = (q.concept_key or q.topic or q.chapter or "general")
    subject = q.subject
    subtopic = q.chapter or q.topic
    pattern = q.pattern or q.question_type
    conf = (payload.confidence or "").strip().lower() or None

    prior = (db.query(_func.count(DBConceptAttempt.id))
             .filter(DBConceptAttempt.user_id == current_user.id,
                     DBConceptAttempt.question_id == q.id).scalar()) or 0

    db.add(DBConceptAttempt(
        user_id=current_user.id, question_id=q.id, concept_key=ck,
        subject=subject, subtopic=subtopic, pattern=pattern,
        correct=correct, selected=sel, difficulty=q.difficulty,
        response_ms=payload.response_ms, confidence=conf,
        hint_used=bool(payload.hint_used),
        attempt_number=prior + 1, exposure_count=prior + 1,
        attempt_context=(payload.attempt_context or "practice"),
        session_id=payload.session_id))

    # ── update concept_mastery ──
    cm = (db.query(DBConceptMastery)
          .filter(DBConceptMastery.user_id == current_user.id,
                  DBConceptMastery.concept_key == ck).first())
    if not cm:
        cm = DBConceptMastery(user_id=current_user.id, concept_key=ck,
                              subject=subject, subtopic=subtopic, mastery=0.5)
        db.add(cm)
    cm.mastery = _aml_update_mastery(cm.mastery, correct, q.difficulty, conf)
    cm.attempts = (cm.attempts or 0) + 1
    cm.correct = (cm.correct or 0) + (1 if correct else 0)
    cm.confidence_n = (cm.confidence_n or 0) + 1
    cm.streak = ((cm.streak or 0) + 1) if correct else 0
    cm.last_seen = datetime.datetime.utcnow()
    cm.state = _aml_state(cm.mastery)

    # ── update skill_mastery (by question pattern) ──
    if pattern:
        sm = (db.query(DBSkillMastery)
              .filter(DBSkillMastery.user_id == current_user.id,
                      DBSkillMastery.pattern == pattern).first())
        if not sm:
            sm = DBSkillMastery(user_id=current_user.id, pattern=pattern, mastery=0.5)
            db.add(sm)
        sm.mastery = _aml_update_mastery(sm.mastery, correct, q.difficulty, conf)
        sm.attempts = (sm.attempts or 0) + 1

    db.commit()
    return {"status": "success", "correct": correct, "concept_key": ck,
            "mastery": round(cm.mastery, 3), "attempts": cm.attempts,
            "streak": cm.streak, "state": cm.state}

@app.get("/me/knowledge-state", tags=["Learning"])
def knowledge_state(subject: Optional[str] = None, db: Session = Depends(get_db),
                    current_user: DBUser = Depends(get_current_user)):
    """The student's current knowledge model — per-concept mastery, per-pattern
    exam skill, and a subject rollup. Basis of the self-prep diagnostic report."""
    from collections import defaultdict
    qy = db.query(DBConceptMastery).filter(DBConceptMastery.user_id == current_user.id)
    if subject:
        qy = qy.filter(DBConceptMastery.subject == subject)
    cms = qy.all()
    concepts = [{"concept_key": c.concept_key, "subject": c.subject, "subtopic": c.subtopic,
                 "mastery": round(c.mastery or 0, 3), "state": c.state,
                 "attempts": c.attempts, "correct": c.correct, "streak": c.streak}
                for c in cms]
    skills = [{"pattern": s.pattern, "mastery": round(s.mastery or 0, 3), "attempts": s.attempts}
              for s in db.query(DBSkillMastery)
              .filter(DBSkillMastery.user_id == current_user.id).all()]
    agg = defaultdict(lambda: [0.0, 0])
    for c in cms:
        agg[c.subject or "General"][0] += (c.mastery or 0)
        agg[c.subject or "General"][1] += 1
    subjects = [{"subject": k, "avg_mastery": round(v[0] / v[1], 3), "concepts": v[1]}
                for k, v in agg.items() if v[1]]
    weakest = sorted(concepts, key=lambda x: x["mastery"])[:15]
    return {"status": "success", "concepts_tracked": len(concepts),
            "subjects": sorted(subjects, key=lambda x: x["subject"]),
            "skills": sorted(skills, key=lambda x: x["mastery"]),
            "weakest_concepts": weakest,
            "concepts": concepts[:500]}

@app.get("/admin/flagged-questions", tags=["Admin"])
def admin_flagged(admin: DBUser = Depends(require_admin), db: Session = Depends(get_db)):
    """Questions students have flagged — drives the AI self-improvement loop."""
    rows = (db.query(DBQuestionFlag, DBQuestion)
            .join(DBQuestion, DBQuestionFlag.question_id == DBQuestion.id)
            .order_by(DBQuestionFlag.created_at.desc()).all())
    agg = {}
    for f, q in rows:
        a = agg.setdefault(q.id, {"question_id": q.id, "text": q.text, "subject": q.subject,
                                  "correct_answer": q.correct_answer, "count": 0, "reasons": {}})
        a["count"] += 1
        r = f.reason or "other"
        a["reasons"][r] = a["reasons"].get(r, 0) + 1
    items = sorted(agg.values(), key=lambda x: -x["count"])
    return {"status": "success", "count": len(items), "flagged": items}


# ── Knowledge Base: upload PDFs → chunks (grounding/library) + extracted MCQs ──
def _chunk_text(text, size=1100, overlap=150):
    """Split into ~`size`-char chunks with `overlap` chars carried between them so
    a fact split across a boundary still appears whole in at least one chunk.
    Only affects NEW uploads; existing chunks are untouched."""
    text = " ".join((text or "").split())
    if not text:
        return []
    step = max(1, size - max(0, overlap))
    return [text[i:i + size] for i in range(0, len(text), step)]

def _vec_literal(vec):
    """Format a python float list as a pgvector literal: '[0.12,0.34,...]'."""
    return "[" + ",".join(f"{float(x):.6f}" for x in vec) + "]"


def _rerank_chunks(query, candidates, k):
    """Blend vector similarity with lexical term overlap so passages that also
    contain the query's exact terms surface above purely-semantic near-misses.
    `candidates`: list of (text, vector_similarity). Returns the top-k texts."""
    import re as _re
    qterms = {w for w in _re.findall(r"[a-z0-9]+", (query or "").lower()) if len(w) > 2}
    scored = []
    for text, sim in candidates:
        if qterms:
            tl = (text or "").lower()
            lex = sum(1 for w in qterms if w in tl) / len(qterms)
        else:
            lex = 0.0
        scored.append((0.7 * float(sim) + 0.3 * lex, text))
    scored.sort(key=lambda x: -x[0])
    return [t for _s, t in scored[:k]]


def _semantic_chunks(db, query, subject=None, k=8):
    """Vector-search the knowledge base for the passages closest to `query`, then
    re-rank the top candidates by lexical overlap for sharper grounding. Returns a
    list of text strings (best first). Empty list if pgvector is off, the query
    can't be embedded, or nothing matches — callers fall back to keyword search."""
    if not (query or "").strip():
        return []
    # PRIMARY: local HF embeddings + Qdrant (free, quota-less). Falls through to
    # the legacy pgvector+Gemini path only when Qdrant isn't configured.
    try:
        import hf_embeddings, vector_store
        if vector_store.enabled():
            qv = hf_embeddings.embed_query(query)
            if qv:
                fetch = min(max(k * 3, k), 40)
                cand = vector_store.search(qv, subject=subject, k=fetch)
                if cand:
                    return _rerank_chunks(query, cand, k)
            else:
                print(f"[rag] HF query embedding failed → keyword fallback "
                      f"({getattr(hf_embeddings, 'LAST_ERROR', '') or 'n/a'})")
            return []
    except Exception as e:
        print(f"[rag] qdrant search error → fallback: {type(e).__name__}: {str(e)[:160]}")
    if not VECTOR_OK:
        return []
    try:
        import gemini_service
        qv = gemini_service.embed_query(query)
        if not qv:
            print(f"[rag] semantic search unavailable → keyword fallback "
                  f"(query embedding None; last_error={getattr(gemini_service, 'LAST_EMBED_ERROR', '') or 'n/a'})")
            return []
        fetch = min(max(k * 3, k), 40)   # over-fetch, then re-rank down to k
        sql = ("SELECT text, 1 - (embedding <=> (:qv)::vector) AS sim FROM knowledge_chunks "
               "WHERE embedding IS NOT NULL " +
               ("AND subject = :subj " if subject else "") +
               "ORDER BY embedding <=> (:qv)::vector LIMIT :fetch")
        params = {"qv": _vec_literal(qv), "fetch": int(fetch)}
        if subject:
            params["subj"] = subject
        rows = db.execute(sa_text(sql), params).fetchall()
        cand = [(r[0], float(r[1]) if r[1] is not None else 0.0) for r in rows if r[0]]
        return _rerank_chunks(query, cand, k)
    except Exception as e:
        print(f"[rag] semantic search error → keyword fallback: {type(e).__name__}: {str(e)[:160]}")
        return []


_CONCEPT_SUBJECT_MAP = {
    "Indian Polity": "Polity & Governance",
    "Governance": "Polity & Governance",
    "Ancient History": "History",
    "Medieval History": "History",
    "Modern History": "History",
    "World History": "History",
    "Post-Independence India": "History",
    "Art & Culture": "Art & Culture",
    "Geography": "Geography",
    "Indian Economy": "Economy",
    "Environment & Ecology": "Environment & Ecology",
    "Science & Technology": "Science & Technology",
    "Indian Society": "Society",
    "International Relations": "International Relations",
    "Internal Security & Disaster Management": "Security",
    "Ethics": "Ethics",
    "Current Affairs": "Current Affairs",
}


def _concept_context(db, subject, topic="", limit_chars=2600):
    """Ground generation in the platform concept library (derived concept names +
    public facts only \u2014 never source text). Maps the fine-grained generation
    subject to the library's broader bucket, prefers topic matches, then falls back
    to the highest-frequency (most exam-relevant) concepts for the subject."""
    csub = _CONCEPT_SUBJECT_MAP.get(subject, subject)
    rows = []
    try:
        if topic:
            rows = (db.query(DBConceptInventory)
                    .filter(DBConceptInventory.subject == csub,
                            DBConceptInventory.concept.ilike(f"%{topic}%"))
                    .order_by(DBConceptInventory.frequency.desc()).limit(24).all())
        if len(rows) < 8:
            seen = {r.id for r in rows}
            more = (db.query(DBConceptInventory)
                    .filter(DBConceptInventory.subject == csub)
                    .order_by(DBConceptInventory.frequency.desc()).limit(24).all())
            rows = rows + [r for r in more if r.id not in seen]
    except Exception:
        return ""
    if not rows:
        return ""
    lines, total = [], 0
    for r in rows:
        try:
            facts = json.loads(r.key_facts or "[]")
        except Exception:
            facts = []
        fact_str = "; ".join(str(f) for f in facts[:4]) if isinstance(facts, list) and facts else ""
        line = f"- {r.concept}" + (f" ({r.subtopic})" if r.subtopic else "") + (f": {fact_str}" if fact_str else "")
        if total + len(line) + 1 > limit_chars:
            break
        lines.append(line)
        total += len(line) + 1
    return "\n".join(lines)


def _retrieve_context(db, subject, topic="", limit_chars=2800):
    """Pull the most relevant uploaded book passages for grounding generation.
    Semantic (pgvector) search first; keyword/ILIKE as a fallback."""
    texts = _semantic_chunks(db, f"{subject} {topic}".strip(), subject=subject, k=8)
    if not texts:
        chunks = []
        if topic:
            chunks = (db.query(DBKnowledgeChunk)
                      .filter(DBKnowledgeChunk.subject == subject,
                              DBKnowledgeChunk.text.ilike(f"%{topic}%"))
                      .limit(6).all())
        if len(chunks) < 3:
            more = db.query(DBKnowledgeChunk).filter(DBKnowledgeChunk.subject == subject).limit(6).all()
            chunks = chunks + [c for c in more if c.id not in {x.id for x in chunks}]
        texts = [c.text for c in chunks]
    out, total = [], 0
    for t in texts:
        out.append(t)
        total += len(t or "")
        if total >= limit_chars:
            break
    return ("\n---\n".join(out))[:limit_chars]


def _retrieve_for_query(db, query, k=5, limit_chars=2400):
    """Cross-subject semantic retrieval for the mentor chat — grounds answers in
    whatever the student actually uploaded. Empty string when nothing fits."""
    texts = _semantic_chunks(db, query, subject=None, k=k)
    out, total = [], 0
    for t in texts:
        snippet = (t or "").strip()
        if not snippet:
            continue
        out.append(snippet)
        total += len(snippet)
        if total >= limit_chars:
            break
    return ("\n---\n".join(out))[:limit_chars]


def _embed_and_store_chunks(source_id, batch=128):
    """Embed every still-unembedded chunk of a source and persist the vectors.
    Best-effort and idempotent — safe to re-run.

    PRIMARY: local HF embeddings → Qdrant. LEGACY fallback: Gemini → pgvector."""
    try:
        import hf_embeddings, vector_store
        if vector_store.enabled():
            if not vector_store.ensure_collection(hf_embeddings.EMBED_DIM):
                print(f"[rag] qdrant unavailable for ingest: {vector_store.LAST_ERROR}")
                return 0
            with engine.connect() as conn:
                rows = conn.execute(sa_text(
                    "SELECT id, text, subject FROM knowledge_chunks WHERE source_id = :sid"),
                    {"sid": source_id}).fetchall()
            done = 0
            for i in range(0, len(rows), 64):
                part = rows[i:i + 64]
                have = vector_store.existing_ids([r[0] for r in part])
                todo = [r for r in part if r[0] not in have]
                if not todo:
                    continue
                vecs = hf_embeddings.embed_passages([r[1] for r in todo])
                done += vector_store.upsert_chunks(
                    (r[0], v, {"text": r[1], "subject": r[2], "source_id": int(source_id)})
                    for r, v in zip(todo, vecs))
            return done
    except Exception as e:
        print(f"[rag] qdrant ingest error: {type(e).__name__}: {str(e)[:160]}")
    if not VECTOR_OK:
        return 0
    from gemini_service import embed_texts
    with engine.connect() as conn:
        rows = conn.execute(sa_text(
            "SELECT id, text FROM knowledge_chunks WHERE source_id = :sid AND embedding IS NULL"),
            {"sid": source_id}).fetchall()
    done = 0
    for i in range(0, len(rows), batch):
        part = rows[i:i + batch]
        vecs = embed_texts([r[1] for r in part], task_type="RETRIEVAL_DOCUMENT")
        with engine.begin() as conn:
            for (cid, _txt), vec in zip(part, vecs):
                if not vec:
                    continue
                conn.execute(sa_text(
                    "UPDATE knowledge_chunks SET embedding = (:v)::vector WHERE id = :id"),
                    {"v": _vec_literal(vec), "id": cid})
                done += 1
    return done

def _extract_text_any(file_bytes, filename, path=None):
    """Extract (page_no, text) from ANY supported file. PDFs use PyMuPDF (with
    Gemini-vision OCR fallback for scanned pages); .docx via python-docx; text files
    decoded directly; images OCR'd via Gemini vision. Returns (pages, total, note).

    Pass `path` (a file on disk) instead of `file_bytes` for PDFs to open the file
    MEMORY-MAPPED — pages are read from disk on demand instead of holding the whole
    file in RAM, so large books extract within the free tier's memory budget."""
    name = (filename or "").lower()
    note = ""
    if name.endswith(".pdf"):
        try:
            import fitz  # PyMuPDF
            from gemini_service import ocr_image
            doc = fitz.open(path) if path else fitz.open(stream=file_bytes, filetype="pdf")
            total = doc.page_count
            pages, ocr_used = [], 0
            for i in range(total):
                page = doc.load_page(i)
                txt = (page.get_text() or "").strip()
                if len(txt) < 20 and ocr_used < 400:        # likely scanned → OCR
                    try:
                        # Render CAPPED at ~1600px on the long side. A fixed dpi=150
                        # on an oversized scan (posters, high-res scanner output) can
                        # decompress to a 100MB+ pixmap — one such page OOMs a 512MB
                        # instance (killed the service on 2026-07-04). Capping pixels
                        # keeps any page under ~8MB while staying sharp enough to OCR.
                        rect = page.rect
                        _long = max(rect.width, rect.height) or 1.0
                        zoom = min(150.0 / 72.0, 1600.0 / _long)
                        pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom))
                        png = pix.tobytes("png")
                        pix = None                          # free the page image at once
                        otxt = ocr_image(png, "image/png")
                        png = None
                        if otxt:
                            txt = otxt; ocr_used += 1
                    except Exception:
                        pass
                    if ocr_used and ocr_used % 8 == 0:
                        gc.collect()                        # return freed pixmap memory promptly
                if txt:
                    pages.append((i + 1, txt))
                page = None
            doc.close()
            if ocr_used:
                note = f"OCR used on {ocr_used} scanned page(s)."
            return pages, total, note
        except Exception:
            pass
        try:
            import io as _io
            from pypdf import PdfReader
            # Stream from disk when we have a path — never slurp a (possibly
            # 100MB) file into RAM just for the fallback reader.
            _fh = open(path, "rb") if (file_bytes is None and path) else _io.BytesIO(file_bytes)
            try:
                reader = PdfReader(_fh)
                pages = []
                for i, p in enumerate(reader.pages):
                    t = (p.extract_text() or "").strip()
                    if t:
                        pages.append((i + 1, t))
                return pages, len(reader.pages), note
            finally:
                try:
                    _fh.close()
                except Exception:
                    pass
        except Exception as e:
            return [], 0, f"Could not read PDF: {str(e)[:140]}"
    # Non-PDF types work on bytes; if we were only handed a path, read it now
    # (docx, images and text files are small, so this is cheap).
    if file_bytes is None and path:
        try:
            with open(path, "rb") as _fh:
                file_bytes = _fh.read()
        except Exception:
            file_bytes = b""
    if name.endswith(".docx"):
        try:
            import io as _io
            import docx
            d = docx.Document(_io.BytesIO(file_bytes))
            parts = [p.text for p in d.paragraphs if p.text and p.text.strip()]
            for tbl in d.tables:
                for row in tbl.rows:
                    cells = [c.text for c in row.cells if c.text and c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            t = "\n".join(parts)
            return ([(1, t)] if t.strip() else []), 1, note
        except Exception as e:
            return [], 0, f"Could not read DOCX: {str(e)[:140]}"
    if name.endswith((".txt", ".md", ".csv", ".json")):
        try:
            t = file_bytes.decode("utf-8", "ignore")
            return ([(1, t)] if t.strip() else []), 1, note
        except Exception:
            return [], 0, "Could not decode text file."
    if name.endswith((".png", ".jpg", ".jpeg", ".webp", ".bmp", ".gif", ".tiff", ".tif")):
        from gemini_service import ocr_image
        ext = name.rsplit(".", 1)[-1]
        mime = {"jpg": "image/jpeg", "jpeg": "image/jpeg", "tif": "image/tiff"}.get(ext, "image/" + ext)
        t = ocr_image(file_bytes, mime)
        return ([(1, t)] if t.strip() else []), 1, ("OCR used." if t.strip() else "No readable text found in the image.")
    try:
        t = file_bytes.decode("utf-8", "ignore")
        if t.strip():
            return [(1, t)], 1, "Read as plain text."
    except Exception:
        pass
    return [], 0, "Unsupported file type — could not extract text."


def _process_file(source_id, file_bytes, filename, subject, mode, answer_bytes=None, answer_name=""):
    """Background worker: extract text from any file (OCR for scans/images) → store
    searchable chunks; for MCQ uploads extract MCQs into an importable test; then
    auto-catalogue the content (subjects/topics/tags). Restart-safe (derived data only)."""
    db = SessionLocal()
    try:
        src = db.query(DBKnowledgeSource).filter(DBKnowledgeSource.id == source_id).first()
        if not src:
            return
        page_texts, total_pages, note = _extract_text_any(file_bytes, filename)
        if not page_texts:
            src.status = "error"; src.error = note or "No readable text found."; db.commit(); return
        chunk_count = 0
        if mode in ("book", "both"):
            for pg, txt in page_texts:
                done = False
                for ch in _chunk_text(txt):
                    if chunk_count >= 6000:
                        done = True; break
                    db.add(DBKnowledgeChunk(source_id=source_id, subject=subject, page=pg, text=ch))
                    chunk_count += 1
                if done:
                    break
        src.pages = total_pages; src.chunk_count = chunk_count; db.commit()
        mcq_count = 0
        if mode in ("mcq", "both"):
            groups, buf = [], ""
            for _pg, txt in page_texts:
                buf += "\n" + txt
                if len(buf) > 4000:
                    groups.append(buf); buf = ""
            if buf.strip():
                groups.append(buf)
            groups = groups[:40]
            collected, seen = [], set()
            for g in groups:
                try:
                    qs = extract_mcqs_from_text(g, subject)
                except Exception:
                    qs = []
                for q in qs:
                    key = (q.get("text") or "")[:80]
                    if not key or key in seen:
                        continue
                    seen.add(key); collected.append(q)
                if len(collected) >= 500:
                    break
            if collected:
                db_test = DBMockTest(
                    title=f"📥 Imported • {src.filename}"[:120],
                    description=f"Imported from your upload — {subject}",
                    subject=subject or "Imported", total_questions=len(collected),
                    duration_minutes=max(5, len(collected)), user_id=src.uploaded_by)
                db.add(db_test); db.commit(); db.refresh(db_test)
                for q in collected:
                    db.add(DBQuestion(
                        text=q["text"], option_a=q["option_a"], option_b=q["option_b"],
                        option_c=q["option_c"], option_d=q["option_d"],
                        correct_answer=(q.get("correct_answer") or "A")[:1].upper(),
                        explanation=q.get("explanation", ""), subject=subject or None,
                        topic=q.get("topic") or None, difficulty="medium",
                        question_type="imported", mock_test_id=db_test.id))
                mcq_count = len(collected)
                src.mock_test_id = db_test.id
                db.commit()
        if answer_bytes:
            try:
                ans_pages, _t, _n = _extract_text_any(answer_bytes, answer_name or "answers.pdf")
                added = 0
                for _pg, txt in ans_pages:
                    for ch in _chunk_text(txt):
                        if added >= 400:
                            break
                        db.add(DBKnowledgeChunk(source_id=source_id, subject=subject,
                                                page=_pg, text="[Answer key] " + ch))
                        added += 1; chunk_count += 1
                src.chunk_count = chunk_count; db.commit()
            except Exception:
                pass
        # Embed all chunks for semantic (pgvector) retrieval. Best-effort.
        try:
            _embed_and_store_chunks(source_id)
        except Exception:
            pass
        # Auto-catalogue / taxonomy
        try:
            from gemini_service import catalogue_content
            sample = " ".join(t for _, t in page_texts)[:8000]
            tax = catalogue_content(sample, filename, subject)
            if not isinstance(tax, dict):
                tax = {}
            if note:
                tax["_processing_note"] = note
            if tax:
                src.taxonomy = json.dumps(tax)
                ps = (tax.get("primary_subject") or "").strip()
                if ps and (not subject or subject == "General"):
                    src.subject = ps[:60]
        except Exception:
            pass
        # Processing succeeded — drop the saved raw copy to free DB space.
        src.mcq_count = mcq_count; src.status = "done"; src.raw_b64 = None; db.commit()
    except Exception as e:
        try:
            src = db.query(DBKnowledgeSource).filter(DBKnowledgeSource.id == source_id).first()
            if src:
                # A real error (not a kill) — record it and drop the saved copy so it
                # isn't retried forever. Kills leave status='processing' for auto-resume.
                src.status = "error"; src.error = str(e)[:200]; src.raw_b64 = None; db.commit()
        except Exception:
            pass
    finally:
        db.close()


def _reprocess_source(source_id):
    """Re-run processing for a source from its saved raw copy (resume after a kill).
    Idempotent: clears any partial derived data first so it never duplicates."""
    import base64 as _b64
    db = SessionLocal()
    data = None; fname = "upload.pdf"; subj = "General"; mode = "book"
    try:
        s = db.query(DBKnowledgeSource).filter(DBKnowledgeSource.id == source_id).first()
        if not s or not s.raw_b64:
            return
        try:
            data = _b64.b64decode(s.raw_b64)
        except Exception:
            return
        # Wipe any partial output from the killed run so the re-run is clean.
        db.query(DBKnowledgeChunk).filter(DBKnowledgeChunk.source_id == source_id).delete()
        if s.mock_test_id:
            try:
                db.query(DBQuestion).filter(DBQuestion.mock_test_id == s.mock_test_id).delete()
                db.query(DBMockTest).filter(DBMockTest.id == s.mock_test_id).delete()
            except Exception:
                pass
            s.mock_test_id = None
        s.status = "processing"; s.error = None; s.pages = 0; s.chunk_count = 0; s.mcq_count = 0
        mode = s.proc_mode or "book"; fname = s.filename or "upload.pdf"; subj = s.subject or "General"
        db.commit()
    except Exception:
        db.close(); return
    finally:
        db.close()
    if data is not None:
        _process_file(source_id, data, fname, subj, mode)


@app.on_event("startup")
def _resume_unfinished_uploads():
    """On boot, resume any upload whose processing was interrupted (e.g. the free
    instance spun down mid-job). Only those with a saved raw copy can resume."""
    import threading
    try:
        db = SessionLocal()
        ids = [s.id for s in db.query(DBKnowledgeSource).filter(
            DBKnowledgeSource.status == "processing",
            DBKnowledgeSource.raw_b64.isnot(None)).all()]
        db.close()
        for sid in ids:
            threading.Thread(target=_reprocess_source, args=(sid,), daemon=True).start()
    except Exception:
        pass


# ── Copyright guard: keep third-party test-series out of the raw-copy path ─────
# Known coaching/aggregator brand & watermark markers. Deliberately narrow so
# that generic 'UPSC', 'NCERT', 'prelims' or actual government PYQs do NOT trip it.
_THIRD_PARTY_MARKERS = (
    "visionias", "vision ias", "insightsonindia", "insights ias", "insight ias",
    "forumias", "forum ias", "gs score", "gsscore", "drishti ias", "drishtiias",
    "vajiram", "shankar ias", "shankarias", "iasbaba", "nextias", "next ias",
    "byju", "unacademy", "la excellence", "rau's ias", "raus ias",
    "test booklet series", "examstatic", "freeupsc", "©vision", "© vision",
)


def _is_third_party_testseries(name: str, text_sample: str = "") -> bool:
    """Heuristic: does this look like copyrighted coaching test-series material?
    Matches coaching/aggregator brand or watermark markers in the filename or a
    sample of the text. Conservative on purpose — only coaching brands / test-
    booklet watermarks trigger it, never generic UPSC/NCERT/PYQ content."""
    hay = (str(name or "") + " \n " + str(text_sample or "")).lower()
    return any(mk in hay for mk in _THIRD_PARTY_MARKERS)


def _quick_text_peek(data: bytes, fname: str, max_chars: int = 4000) -> str:
    """Cheap first-pages text peek (no OCR) for the copyright-guard check."""
    name = (fname or "").lower()
    try:
        if name.endswith(".pdf"):
            import fitz
            doc = fitz.open(stream=data, filetype="pdf")
            parts = [doc.load_page(i).get_text() or "" for i in range(min(2, doc.page_count))]
            doc.close()
            return (" ".join(parts))[:max_chars]
        if name.endswith((".txt", ".md", ".csv", ".json")):
            return data[: max_chars * 2].decode("utf-8", "ignore")[:max_chars]
    except Exception:
        return ""
    return ""


def _canonical_subject(s):
    """Map a free-text subject label to ONE canonical top-level UPSC GS subject.
    The LLM writes the subject free-form, so the same subject arrives under many
    labels (Economy/Economics/Indian Economy; Polity/Indian Polity/Governance;
    History/Ancient/Medieval/Modern; ...). Keyword rules collapse them and also
    catch typos and future variants. Finer detail stays in the `subtopic` column."""
    import re as _re
    if not s:
        return "Other"
    t = _re.sub(r"\s+", " ", str(s).strip().lower())
    if not t:
        return "Other"
    _ALIAS = {"ir": "International Relations", "s&t": "Science & Technology",
              "sci-tech": "Science & Technology", "dm": "Security",
              "gs1": "History", "gs2": "Polity & Governance",
              "gs3": "Economy", "gs4": "Ethics"}
    if t in _ALIAS:
        return _ALIAS[t]
    _RULES = (
        ("Art & Culture",          ("art and culture", "art & culture", "art&culture", "culture", "painting", "architecture", "sculpture", "dance", "music", "heritage")),
        ("History",                ("histor", "ancient india", "medieval", "modern india", "post independence", "post-independence", "freedom struggle", "national movement", "world history", "indus", "mauryan", "gupta", "mughal", "maratha", "revolt")),
        ("Geography",              ("geograph", "climat", "oceanograph", "geomorph", "physiograph", "monsoon", "soil", "mineral", "mapping", "places in news")),
        ("Polity & Governance",    ("polit", "governance", "constitution", "parliament", "judiciar", "fundamental right", "panchayat", "federal", "election", "amendment")),
        ("Economy",                ("econom", "fiscal", "monetary", "banking", "budget", "gdp", "inflation", "taxation", "trade", "fintech", "agricultur", "industr", "infrastructure", "poverty", "employment", "subsid")),
        ("Environment & Ecology",  ("environ", "ecolog", "biodiv", "climate change", "pollution", "wildlife", "conservation", "forest", "species", "protected area")),
        ("Science & Technology",   ("science", "technolog", "space", "biotech", "nuclear", "artificial intel", "computer", "robot", "genom", "vaccine")),
        ("International Relations", ("internation", "foreign", "diplomac", "bilateral", "neighbour", "geopolit", "united nations", "summit", "treaty", "multilater")),
        ("Society",                ("societ", "social", "women", "population", "urban", "tribal", "secular", "communal", "health", "education", "welfare")),
        ("Ethics",                 ("ethic", "integrity", "aptitude", "attitude", "moral", "emotional intelligence", "probity")),
        ("Security",               ("secur", "terror", "naxal", "cyber", "border", "armed force", "money launder", "disaster")),
        ("Current Affairs",        ("current affair", "monthly", "daily quiz", "news")),
    )
    for canon, kws in _RULES:
        for kw in kws:
            if kw in t:
                return canon
    return "Other"


def _extract_concept_inventory(text: str, max_workers: int = 3, chunk_size: int = 22000,
                               progress_cb=None, stats=None):
    # chunk_size 22000 (≈5.5k tokens): halves the AI calls per book vs 11000 —
    # free-tier daily quotas are per-REQUEST on Gemini, so bigger windows double
    # the number of books that fit through a day's allowance.
    """LLM concept extraction: text -> RICH concept metadata ONLY (never verbatim).
    Shared by the extraction endpoint and the upload copyright-guard.

    FAST: the document is split into large windows analysed CONCURRENTLY (thread
    pool), so a big PDF finishes in about the time of its slowest window instead of
    the sum of all windows; larger windows also mean far fewer LLM calls.

    POWERFUL: every concept returns subject, finer subtopic, tested pattern,
    difficulty and exam importance plus public key-facts; results are then
    de-duplicated across windows — repeats merged, facts pooled, and a `frequency`
    count records how often the concept is tested, so the most exam-relevant
    concepts sort to the top, ready to drive original question generation."""
    import gemini_service
    import json as _json
    import re as _re
    import time as _time
    from concurrent.futures import ThreadPoolExecutor

    text = text or ""
    if stats is None:
        stats = {}
    stats.setdefault("failed", 0)
    stats.setdefault("chunks", 0)
    stats.setdefault("quota", False)
    if not text.strip():
        return []

    system = (
        "You are an expert UPSC curriculum analyst. The text may contain COPYRIGHTED "
        "practice questions. NEVER reproduce, quote or paraphrase any question, option "
        "or explanation. For every distinct item you detect, output ONLY analytical "
        "metadata: the underlying CONCEPT it tests, the SUBJECT, a finer SUBTOPIC, the "
        "question FORMAT (direct/statement_based/assertion_reason/match_the_following/"
        "pairs/how_many/sequencing/odd_one_out), a DIFFICULTY (easy/medium/hard), an "
        "exam IMPORTANCE (high/medium/low), and 2-5 short PUBLIC, verifiable factual "
        "points a student must know (facts are not copyrightable; original sentences "
        "are — use your OWN words). Merge near-duplicate concepts. STRICT JSON only.")
    schema = ('{"items":[{"concept":str,"subject":str,"subtopic":str,"pattern":str,'
              '"difficulty":"easy|medium|hard","importance":"high|medium|low",'
              '"key_facts":[str]}]}')

    chunks = [text[i:i + chunk_size] for i in range(0, len(text), chunk_size)] or [text]

    # Failure tracking (closures; GIL-safe for counters). `_dead` is a circuit
    # breaker: once one chunk exhausts its rate-limit retries, remaining chunks
    # fail FAST instead of each spending minutes retrying a dead quota.
    _dead = [False]

    def _is_quota(msg):
        m = (msg or "").lower()
        return ("429" in m or "resource_exhausted" in m or "quota" in m
                or "rate limit" in m or "overloaded" in m or "503" in m)

    def _analyse(chunk):
        if _dead[0]:
            stats["failed"] += 1
            return []
        prompt = ("Extract concept metadata ONLY (never the question wording). "
                  f"Return STRICT JSON exactly in this shape: {schema}\n\nTEXT:\n{chunk}")
        raw = None
        wait = 35.0
        for attempt in range(3):
            try:
                raw = gemini_service.gen_text(prompt, json_mode=True, prefer="deepseek", system=system)
                break
            except Exception as e:
                msg = str(e)
                if attempt < 2 and _is_quota(msg) and not _dead[0]:
                    m = _re.search(r"retry in ([0-9.]+)s", msg)
                    try:
                        hint = float(m.group(1)) + 2.0 if m else wait
                    except Exception:
                        hint = wait
                    _time.sleep(min(70.0, hint))
                    wait = min(70.0, wait * 2)
                    continue
                stats["failed"] += 1
                if _is_quota(msg):
                    stats["quota"] = True
                    _dead[0] = True       # open the circuit — quota is gone
                return []
        m = _re.search(r"\{.*\}", raw or "", _re.DOTALL)
        if not m:
            return []
        try:
            return _json.loads(m.group(0)).get("items", []) or []
        except Exception:
            return []

    # Analyse windows CONCURRENTLY (LLM calls are I/O-bound) — the big wall-clock win.
    total_chunks = len(chunks)
    stats["chunks"] = total_chunks
    def _report(done):
        if progress_cb:
            try:
                progress_cb(done, total_chunks)
            except Exception:
                pass
    raw_items = []
    done = 0
    if total_chunks == 1:
        raw_items = _analyse(chunks[0]); done = 1; _report(done)
    else:
        try:
            with ThreadPoolExecutor(max_workers=min(max_workers, total_chunks)) as ex:
                for res in ex.map(_analyse, chunks):
                    raw_items.extend(res or [])
                    done += 1; _report(done)
        except Exception:
            for ch in chunks:                       # fallback: sequential
                raw_items.extend(_analyse(ch)); done += 1; _report(done)

    # De-duplicate + merge across windows; pool facts; count frequency.
    def _norm(s):
        return _re.sub(r"[^a-z0-9]+", " ", (s or "").lower()).strip()

    _PAT_OK = {"direct", "statement_based", "assertion_reason", "match_the_following",
               "pairs", "how_many", "sequencing", "odd_one_out"}
    _DIFF_OK = {"easy", "medium", "hard"}
    _IMP_OK = {"high", "medium", "low"}

    merged, order = {}, []
    for it in raw_items:
        concept = (it.get("concept") or "").strip()
        key = _norm(concept)
        if not key:
            continue
        if key not in merged:
            merged[key] = {"concept": concept, "subject": (it.get("subject") or "").strip(),
                           "subtopic": (it.get("subtopic") or "").strip(),
                           "pattern": (it.get("pattern") or "direct"),
                           "difficulty": (it.get("difficulty") or "medium"),
                           "importance": (it.get("importance") or "medium"),
                           "key_facts": [], "frequency": 0}
            order.append(key)
        rec = merged[key]
        rec["frequency"] += 1
        if not rec["subject"] and it.get("subject"):
            rec["subject"] = it["subject"].strip()
        if not rec["subtopic"] and it.get("subtopic"):
            rec["subtopic"] = it["subtopic"].strip()
        seen = {_norm(f) for f in rec["key_facts"]}
        for f in (it.get("key_facts") or []):
            fs = str(f)[:300].strip()
            if fs and _norm(fs) not in seen:
                rec["key_facts"].append(fs)
                seen.add(_norm(fs))
        rec["key_facts"] = rec["key_facts"][:6]

    items = []
    for k in order:
        rec = merged[k]
        if rec["pattern"] not in _PAT_OK:
            rec["pattern"] = "direct"
        if rec["difficulty"] not in _DIFF_OK:
            rec["difficulty"] = "medium"
        if rec["importance"] not in _IMP_OK:
            rec["importance"] = "medium"
        items.append(rec)

    _imp = {"high": 0, "medium": 1, "low": 2}
    items.sort(key=lambda x: (-x["frequency"], _imp.get(x["importance"], 1), x["concept"].lower()))
    return items


# Upload categories -> (processing mode, friendly label)
KNOWLEDGE_CATEGORIES = {
    "mcq_no_answers":   ("mcq",  "MCQs (no answers)"),
    "mcq_with_answers": ("mcq",  "MCQs (with answers)"),
    "book":             ("book", "Book"),
    "current_affairs":  ("both", "Current Affairs"),
    "strategy":         ("book", "Strategy"),
    "techniques":       ("book", "Techniques"),
    "mnemonics":        ("book", "Mnemonics"),
    "other":            ("book", "Other"),
}

@app.post("/admin/knowledge/upload", tags=["Knowledge"])
async def knowledge_upload(background: BackgroundTasks, file: UploadFile = File(...),
                           answer_file: Optional[UploadFile] = File(None),
                           subject: str = Form(""), category: str = Form("book"),
                           kind: str = Form(""), month: str = Form(""),
                           description: str = Form(""),
                           admin: DBUser = Depends(require_admin), db: Session = Depends(get_db)):
    # SAFE-BY-DESIGN (copyright): this endpoint no longer stores raw files or verbatim
    # text. The raw-copy knowledge uploader (which kept raw_b64 + verbatim chunks) was
    # RETIRED. Every upload — regardless of source — is now routed to the zero-footprint
    # concept path: the file is read in memory, only concept metadata is returned, and
    # nothing is written to disk or the database. This is the single safe upload path.
    fname = (file.filename or "upload").strip() or "upload"
    data = await file.read()
    if not data:
        raise HTTPException(status_code=400, detail="The uploaded file is empty.")
    pages, total, note = _extract_text_any(data, fname)
    _text = "\n\n".join(t for _p, t in pages)
    del data  # never keep the source bytes
    if not _text.strip():
        return {"status": "success", "mode": "concept_extraction", "file": fname,
                "note": note or "No extractable text.", "item_count": 0,
                "footprint": "none", "concepts": []}
    concepts = _extract_concept_inventory(_text)
    # Zero-footprint guarantee: the source file and its extracted text lived only in
    # memory for this request. Release them; only the concept metadata below survives.
    _text = ""
    pages = None
    import gc as _gc
    _gc.collect()
    return {
        "status": "success",
        "mode": "concept_extraction",
        "file": fname,
        "pages": total,
        "note": note or "",
        "item_count": len(concepts),
        "message": ("Concepts extracted. Nothing was stored — no raw copy and no verbatim text. "
                    "Use these concepts to generate ORIGINAL questions."),
        "footprint": ("none — the source file and its text were held only in memory for this "
                      "request, never written to disk or the database; only the concept metadata "
                      "below persists."),
        "concepts": concepts,
    }

@app.get("/admin/knowledge/sources", tags=["Knowledge"])
def knowledge_sources(admin: DBUser = Depends(require_admin), db: Session = Depends(get_db)):
    srcs = db.query(DBKnowledgeSource).order_by(DBKnowledgeSource.id.desc()).all()
    total_chunks = db.query(DBKnowledgeChunk).count()
    return {"status": "success", "total_chunks": total_chunks,
            "sources": [{
                "id": s.id, "filename": s.filename, "subject": s.subject, "kind": s.kind,
                "pages": s.pages or 0, "chunk_count": s.chunk_count or 0, "mcq_count": s.mcq_count or 0,
                "status": s.status, "error": s.error, "mock_test_id": s.mock_test_id,
                "description": s.description, "file_type": s.file_type,
                "resumable": bool(s.raw_b64),
                "taxonomy": (json.loads(s.taxonomy) if s.taxonomy else None),
                "created_at": s.created_at.isoformat() if s.created_at else None,
            } for s in srcs]}


@app.post("/admin/knowledge/purge-third-party", tags=["Knowledge"])
def purge_third_party(dry_run: bool = True, deep: bool = False,
                      admin: DBUser = Depends(require_admin), db: Session = Depends(get_db)):
    """Retroactive copyright guard for ALREADY-uploaded files: scan knowledge
    sources for third-party (copyrighted coaching) test-series material and remove
    its footprint — the stored raw copy (raw_b64) AND the verbatim chunks.
    dry_run=true (default) only reports what WOULD be purged; dry_run=false performs
    the removal, leaving an audit stub (filename + status) with no content retained.

    Blob-free AND set-based: it never loads the (potentially huge) raw_b64 column,
    and it detects/purges with a handful of aggregate queries instead of one round
    trip per source — so it stays fast even on a slow free-tier DB with a single
    worker (the old per-source loop could hold the worker for minutes)."""
    # 1) All sources (id, filename, has_raw) — one small query, no blobs.
    srcs = db.execute(sa_text(
        "SELECT id, filename, (raw_b64 IS NOT NULL) FROM knowledge_sources")).fetchall()

    dialect = db.bind.dialect.name if db.bind is not None else "sqlite"

    # 2) OPTIONAL deep scan (deep=true): one text sample per source so watermarked
    #    files that were RENAMED (brand stripped from the filename) still get caught.
    #    This reads a slice of every chunk, so it is heavy on a throttled free-tier
    #    DB — it is OFF by default. The default path below uses filenames only,
    #    which catches coaching-branded uploads (e.g. "VisionIAS - ...pdf") instantly
    #    and returns reliably instead of stalling behind a proxy timeout.
    samples = {}
    if deep:
        left_fn = "left(text, 800)" if dialect.startswith("postgres") else "substr(text, 1, 800)"
        try:
            for sid, txt in db.execute(sa_text(
                    f"SELECT source_id, MIN({left_fn}) FROM knowledge_chunks "
                    "GROUP BY source_id")).fetchall():
                samples[sid] = txt or ""
        except Exception:
            samples = {}

    # 3) Detect. Filename (+ optional deep sample). No chunk scan unless deep=true.
    result = []
    flagged_ids = []
    for sid, fname, has_raw in srcs:
        if not _is_third_party_testseries(fname or "", samples.get(sid, "")):
            continue
        flagged_ids.append(int(sid))
        result.append({"id": sid, "filename": fname,
                       "chunk_count": 0, "had_raw_copy": bool(has_raw)})

    # 4) Chunk counts for the FLAGGED sources only — bounded and fast (no full scan).
    if flagged_ids:
        _idl = ",".join(str(i) for i in flagged_ids)
        try:
            counts = {sid: int(n) for sid, n in db.execute(sa_text(
                f"SELECT source_id, COUNT(*) FROM knowledge_chunks "
                f"WHERE source_id IN ({_idl}) GROUP BY source_id")).fetchall()}
            for item in result:
                item["chunk_count"] = counts.get(item["id"], 0)
        except Exception:
            pass

    purged = 0
    if not dry_run and flagged_ids:
        # ids come straight from the DB as integers — safe to inline for a bulk purge.
        id_list = ",".join(str(i) for i in flagged_ids)
        db.execute(sa_text(f"DELETE FROM knowledge_chunks WHERE source_id IN ({id_list})"))
        db.execute(sa_text(
            "UPDATE knowledge_sources SET raw_b64 = NULL, chunk_count = 0, "
            "status = 'removed_third_party', "
            "error = 'Removed as third-party copyrighted material (zero-footprint policy).' "
            f"WHERE id IN ({id_list})"))
        db.commit()
        purged = len(flagged_ids)

    return {"status": "success", "dry_run": dry_run,
            "detection": "filename+content" if deep else "filename",
            "flagged": len(result), "purged": purged, "items": result,
            "note": ("Dry run — nothing changed. Re-run with dry_run=false to purge."
                     if dry_run else
                     "Raw copies and verbatim chunks removed for flagged sources; audit stubs kept.")}


@app.post("/admin/knowledge/reset-stuck", tags=["Knowledge"])
def reset_stuck_sources(admin: DBUser = Depends(require_admin), db: Session = Depends(get_db)):
    """Mark uploads stuck in 'processing' (their background job was killed by a
    server restart) as 'error', so they stop showing as in-progress. The raw file
    isn't retained, so these must be re-uploaded to finish processing."""
    rows = db.query(DBKnowledgeSource).filter(DBKnowledgeSource.status == "processing").all()
    ids = []
    for s in rows:
        s.status = "error"
        s.error = "Processing was interrupted by a server restart. Please delete and re-upload this file."
        ids.append(s.id)
    db.commit()
    return {"status": "success", "reset": len(ids), "ids": ids}


@app.post("/admin/knowledge/resume", tags=["Knowledge"])
def resume_processing(admin: DBUser = Depends(require_admin), db: Session = Depends(get_db)):
    """Manually re-trigger processing for uploads stuck in 'processing' that still
    have a saved raw copy. Uploads without a saved copy (older ones) can't resume —
    those are reported so the admin knows to delete + re-upload them."""
    import threading
    resumable = [s.id for s in db.query(DBKnowledgeSource).filter(
        DBKnowledgeSource.status == "processing", DBKnowledgeSource.raw_b64.isnot(None)).all()]
    no_copy = db.query(DBKnowledgeSource).filter(
        DBKnowledgeSource.status == "processing", DBKnowledgeSource.raw_b64.is_(None)).count()
    for sid in resumable:
        threading.Thread(target=_reprocess_source, args=(sid,), daemon=True).start()
    return {"status": "success", "resuming": resumable,
            "stuck_without_saved_file": no_copy}


@app.delete("/admin/knowledge/source/{source_id}", tags=["Knowledge"])
def delete_knowledge_source(source_id: int, admin: DBUser = Depends(require_admin),
                            db: Session = Depends(get_db)):
    """Delete a knowledge source and everything derived from it (chunks, and the
    imported MCQ test + its questions if one was created)."""
    s = db.query(DBKnowledgeSource).filter(DBKnowledgeSource.id == source_id).first()
    if not s:
        raise HTTPException(status_code=404, detail="Source not found")
    db.query(DBKnowledgeChunk).filter(DBKnowledgeChunk.source_id == source_id).delete()
    if s.mock_test_id:
        try:
            db.query(DBQuestion).filter(DBQuestion.mock_test_id == s.mock_test_id).delete()
            db.query(DBMockTest).filter(DBMockTest.id == s.mock_test_id).delete()
        except Exception:
            pass
    db.delete(s)
    db.commit()
    try:                       # drop this source's vectors from Qdrant too
        import vector_store
        if vector_store.enabled():
            vector_store.delete_source(source_id)
    except Exception:
        pass
    return {"status": "success", "deleted": source_id, "filename": s.filename}

@app.post("/admin/embeddings/backfill", tags=["Knowledge"])
def backfill_embeddings(admin: DBUser = Depends(require_admin), db: Session = Depends(get_db)):
    """Embed chunks that don't yet have a vector (e.g. uploaded before RAG existed).
    Processes a bounded batch per call to stay within request limits — call again
    while `still_pending` > 0 until it reaches 0."""
    # PRIMARY: local HF embeddings → Qdrant.
    try:
        import hf_embeddings, vector_store
        if vector_store.enabled():
            if not vector_store.ensure_collection(hf_embeddings.EMBED_DIM):
                return {"status": "error", "backend": "qdrant",
                        "reason": vector_store.LAST_ERROR or "qdrant unavailable"}
            total = db.execute(sa_text("SELECT COUNT(*) FROM knowledge_chunks")).scalar() or 0
            rows = db.execute(sa_text(
                "SELECT id, text, subject, source_id FROM knowledge_chunks ORDER BY id")).fetchall()
            done = 0
            for i in range(0, len(rows), 100):
                part = rows[i:i + 100]
                have = vector_store.existing_ids([r[0] for r in part])
                todo = [r for r in part if r[0] not in have]
                if not todo:
                    continue
                vecs = hf_embeddings.embed_passages([r[1] for r in todo])
                done += vector_store.upsert_chunks(
                    (r[0], v, {"text": r[1], "subject": r[2], "source_id": int(r[3] or 0)})
                    for r, v in zip(todo, vecs))
                if done >= 600:      # bounded per call
                    break
            in_store = vector_store.count() or 0
            return {"status": "success", "backend": "qdrant+hf",
                    "embed_model": hf_embeddings.HF_EMBED_MODEL,
                    "embedded_this_run": done, "chunks_total": int(total),
                    "points_in_qdrant": in_store,
                    "still_pending": max(0, int(total) - in_store),
                    "done": in_store >= int(total),
                    "last_error": hf_embeddings.LAST_ERROR or vector_store.LAST_ERROR or ""}
    except Exception as e:
        return {"status": "error", "backend": "qdrant", "reason": f"{type(e).__name__}: {str(e)[:200]}"}
    if not VECTOR_OK:
        return {"status": "skipped",
                "reason": "pgvector is not enabled on this database; semantic search is using keyword fallback."}
    import gemini_service
    from gemini_service import embed_texts
    was_pending = db.execute(sa_text(
        "SELECT COUNT(*) FROM knowledge_chunks WHERE embedding IS NULL")).scalar() or 0
    rows = db.execute(sa_text(
        "SELECT id, text FROM knowledge_chunks WHERE embedding IS NULL ORDER BY id LIMIT 300")).fetchall()
    done = 0
    for i in range(0, len(rows), 50):
        part = rows[i:i + 50]
        vecs = embed_texts([r[1] for r in part], task_type="RETRIEVAL_DOCUMENT")
        for (cid, _t), vec in zip(part, vecs):
            if not vec:
                continue
            db.execute(sa_text("UPDATE knowledge_chunks SET embedding = (:v)::vector WHERE id = :id"),
                       {"v": _vec_literal(vec), "id": cid})
            done += 1
        db.commit()
    still_pending = db.execute(sa_text(
        "SELECT COUNT(*) FROM knowledge_chunks WHERE embedding IS NULL")).scalar() or 0
    try:
        provider = gemini_service.embed_provider()
    except Exception:
        provider = "gemini"
    return {"status": "success", "embedded_this_run": done,
            "was_pending": was_pending, "still_pending": still_pending,
            "embed_provider": provider,
            "embed_model": getattr(gemini_service, "EMBED_MODEL", ""),
            "last_error": getattr(gemini_service, "LAST_EMBED_ERROR", ""),
            "done": still_pending == 0}


@app.post("/admin/embeddings/reset", tags=["Knowledge"])
def reset_embeddings(admin: DBUser = Depends(require_admin), db: Session = Depends(get_db)):
    """Clear ALL stored chunk vectors (set embedding = NULL). Needed once when you
    SWITCH embedding providers (e.g. Gemini → OpenAI): vectors from different models
    live in different spaces, so the whole index must be rebuilt with the new one.
    After calling this, run /admin/embeddings/backfill until still_pending = 0."""
    if not VECTOR_OK:
        return {"status": "skipped", "reason": "pgvector not enabled."}
    try:
        provider = __import__("gemini_service").embed_provider()
    except Exception:
        provider = "gemini"
    db.execute(sa_text("UPDATE knowledge_chunks SET embedding = NULL WHERE embedding IS NOT NULL"))
    db.commit()
    pending = db.execute(sa_text("SELECT COUNT(*) FROM knowledge_chunks")).scalar() or 0
    return {"status": "success", "cleared": True, "now_active_provider": provider,
            "chunks_to_reembed": pending}


@app.post("/admin/rescue/import-old-db", tags=["Knowledge"])
def rescue_import_old_db(what: str = "concepts", offset: int = 0,
                         admin: DBUser = Depends(require_admin), db: Session = Depends(get_db)):
    """RESCUE: merge data from the OLD (suspended Render) database into this one.

    Set the OLD_DATABASE_URL secret (the old DB's EXTERNAL connection string),
    then call repeatedly while still_pending > 0:
      what=concepts  → merges concept_inventory (frequencies/sources/facts pooled)
      what=users     → copies accounts that don't exist here yet (same passwords)
    Bounded batches (2000 rows/call) so each request stays quick. Idempotent —
    safe to re-run; nothing is ever deleted from either side."""
    import json as _json
    old_url = os.getenv("OLD_DATABASE_URL", "").strip()
    if not old_url:
        return {"status": "skipped", "reason": "OLD_DATABASE_URL secret is not set."}
    if old_url.startswith("postgres://"):
        old_url = old_url.replace("postgres://", "postgresql://", 1)
    if "render.com" in old_url and "sslmode=" not in old_url:
        old_url += ("&" if "?" in old_url else "?") + "sslmode=require"
    from sqlalchemy import create_engine as _ce
    try:
        old_eng = _ce(old_url, pool_pre_ping=True, connect_args={"connect_timeout": 15})
        with old_eng.connect() as oc:
            if what == "users":
                cols = [c.name for c in DBUser.__table__.columns if c.name != "id"]
                rows = oc.execute(sa_text(
                    "SELECT " + ", ".join(cols) + " FROM users ORDER BY id "
                    "LIMIT 2000 OFFSET :off"), {"off": int(offset)}).mappings().all()
                total = oc.execute(sa_text("SELECT COUNT(*) FROM users")).scalar() or 0
                existing = {e for (e,) in db.query(DBUser.email).all()}
                added = 0
                for m in rows:
                    if (m.get("email") or "").lower() in {x.lower() for x in existing}:
                        continue
                    db.add(DBUser(**{k: m[k] for k in cols}))
                    existing.add(m.get("email") or "")
                    added += 1
                db.commit()
                nxt = int(offset) + len(rows)
                return {"status": "success", "what": "users", "imported_this_run": added,
                        "old_total": int(total), "next_offset": nxt,
                        "still_pending": max(0, int(total) - nxt), "done": nxt >= int(total)}
            # concepts (default)
            total = oc.execute(sa_text("SELECT COUNT(*) FROM concept_inventory")).scalar() or 0
            rows = oc.execute(sa_text(
                "SELECT key, concept, subject, subtopic, pattern, difficulty, importance, "
                "key_facts, frequency, sources FROM concept_inventory ORDER BY id "
                "LIMIT 2000 OFFSET :off"), {"off": int(offset)}).fetchall()
            keys = [r[0] for r in rows if r[0]]
            existing = {e.key: e for e in db.query(DBConceptInventory)
                        .filter(DBConceptInventory.key.in_(keys)).all()} if keys else {}
            added = merged = 0
            from datetime import datetime as _dt
            for (k, concept, subject, subtopic, pattern, difficulty,
                 importance, key_facts, frequency, sources) in rows:
                if not k:
                    continue
                row = existing.get(k)
                if row:
                    row.frequency = (row.frequency or 1) + int(frequency or 1)
                    row.sources = (row.sources or 1) + int(sources or 1)
                    try:
                        old_f = _json.loads(key_facts or "[]")
                        new_f = _json.loads(row.key_facts or "[]")
                        row.key_facts = _json.dumps(list(dict.fromkeys([*new_f, *old_f]))[:12])
                    except Exception:
                        pass
                    row.updated_at = _dt.utcnow()
                    merged += 1
                else:
                    db.add(DBConceptInventory(
                        key=k, concept=concept, subject=subject, subtopic=subtopic,
                        pattern=pattern, difficulty=difficulty, importance=importance,
                        key_facts=key_facts or "[]", frequency=int(frequency or 1),
                        sources=int(sources or 1)))
                    added += 1
            db.commit()
            nxt = int(offset) + len(rows)
            return {"status": "success", "what": "concepts", "new_concepts": added,
                    "merged_into_existing": merged, "old_total": int(total),
                    "next_offset": nxt, "still_pending": max(0, int(total) - nxt),
                    "done": nxt >= int(total)}
    except Exception as e:
        return {"status": "error",
                "reason": f"{type(e).__name__}: {str(e)[:300]}",
                "hint": "Old DB still suspended/unreachable, or wrong OLD_DATABASE_URL."}


@app.post("/admin/concepts/consolidate-subjects", tags=["Knowledge"])
def consolidate_subjects(after_id: int = 0, limit: int = 1000, apply: bool = False,
                         admin: DBUser = Depends(require_admin),
                         db: Session = Depends(get_db)):
    """Consolidate free-text concept_inventory.subject into canonical GS subjects
    and MERGE the duplicate rows that result. Because the dedup key is
    concept|subject, the same concept under two labels (e.g. "Economy" vs
    "Indian Economy") is stored as two rows; canonicalising the subject makes
    their keys collide, and this endpoint pools them (frequency/sources/facts)
    like the normal upsert instead of leaving duplicates.

    Call repeatedly with after_id = last_id from the previous response until
    done=true. apply=false is a DRY RUN (writes nothing); apply=true performs the
    relabel+merge for the batch. Idempotent: already-canonical rows are skipped,
    so it is safe to re-run."""
    import json as _json
    from datetime import datetime as _dt
    rows = (db.query(DBConceptInventory)
            .filter(DBConceptInventory.id > int(after_id))
            .order_by(DBConceptInventory.id)
            .limit(max(1, min(int(limit), 5000))).all())
    if not rows:
        return {"status": "success", "done": True, "last_id": int(after_id),
                "scanned": 0, "relabeled": 0, "merged": 0, "apply": bool(apply)}
    relabeled = merged = skipped = 0
    last_id = int(after_id)
    for r in rows:
        last_id = r.id
        canon = _canonical_subject(r.subject)
        new_key = ((r.concept or "").strip().lower() + "|" + canon.lower())[:400]
        if new_key == (r.key or ""):
            skipped += 1
            continue
        target = (db.query(DBConceptInventory)
                  .filter(DBConceptInventory.key == new_key,
                          DBConceptInventory.id != r.id).first())
        if target is not None:
            merged += 1
            if apply:
                target.frequency = (target.frequency or 1) + (r.frequency or 1)
                target.sources = (target.sources or 1) + (r.sources or 1)
                try:
                    tf = _json.loads(target.key_facts or "[]")
                    rf = _json.loads(r.key_facts or "[]")
                    target.key_facts = _json.dumps(list(dict.fromkeys([*tf, *rf]))[:12])
                except Exception:
                    pass
                if not target.subtopic and r.subtopic:
                    target.subtopic = r.subtopic
                if not target.importance and r.importance:
                    target.importance = r.importance
                target.updated_at = _dt.utcnow()
                db.delete(r)
                db.flush()          # so a later row in this batch sees the merge
        else:
            relabeled += 1
            if apply:
                r.subject = canon
                r.key = new_key
                r.updated_at = _dt.utcnow()
                db.flush()          # make the new key visible to later rows' lookups
    if apply:
        db.commit()
    else:
        db.rollback()
    return {"status": "success", "done": len(rows) < max(1, min(int(limit), 5000)),
            "last_id": last_id, "scanned": len(rows), "relabeled": relabeled,
            "merged": merged, "skipped_already_canonical": skipped,
            "apply": bool(apply),
            "hint": "re-call with after_id=last_id until done=true; set apply=true to write"}


def _plan_subject_consolidation(batch, existing_by_key):
    """Pure planning logic for the FAST consolidation endpoint (unit-testable).
    batch: list of row-dicts (id, concept, subject, key, subtopic, importance,
    frequency, sources, key_facts) in id order. existing_by_key: row-dicts for
    every row in the WHOLE table whose key equals one of the batch's computed
    canonical keys. Returns (updates, delete_ids, stats): identical semantics to
    the per-row endpoint - relabel when the canonical key is free, merge
    (pool frequency/sources/key_facts, backfill subtopic/importance) when it is
    taken. Groups whose key-holder is itself being relabeled this batch are
    DEFERRED (ordering inside one UPDATE is undefined); a final sweep from
    after_id=0 catches them."""
    import json as _json
    changed, skipped = [], 0
    for r in batch:
        canon = _canonical_subject(r["subject"])
        nkey = ((r["concept"] or "").strip().lower() + "|" + canon.lower())[:400]
        if nkey == (r["key"] or ""):
            skipped += 1
            continue
        changed.append((r, canon, nkey))
    changed_ids = {r["id"] for r, _c, _k in changed}
    groups = {}
    for item in changed:
        groups.setdefault(item[2], []).append(item)

    def _facts(*rows):
        pooled = []
        for rr in rows:
            try:
                pooled.extend(_json.loads(rr.get("key_facts") or "[]"))
            except Exception:
                pass
        return _json.dumps(list(dict.fromkeys(pooled))[:12], ensure_ascii=False)

    updates, deletes = [], []
    relabeled = merged = deferred = 0
    for nkey, members in groups.items():
        members.sort(key=lambda it: it[0]["id"])
        ext = existing_by_key.get(nkey)
        if ext is not None and ext["id"] in changed_ids:
            deferred += len(members)
            continue
        canon = members[0][1]
        if ext is not None:
            surv, rest = ext, [it[0] for it in members]
        else:
            surv, rest = members[0][0], [it[0] for it in members[1:]]
        updates.append({
            "id": surv["id"], "subject": canon, "key": nkey,
            "frequency": (surv.get("frequency") or 1) + sum((rr.get("frequency") or 1) for rr in rest),
            "sources": (surv.get("sources") or 1) + sum((rr.get("sources") or 1) for rr in rest),
            "key_facts": _facts(surv, *rest),
            "subtopic": surv.get("subtopic") or next((rr.get("subtopic") for rr in rest if rr.get("subtopic")), None),
            "importance": surv.get("importance") or next((rr.get("importance") for rr in rest if rr.get("importance")), None),
        })
        deletes.extend(rr["id"] for rr in rest)
        merged += len(rest)
        if ext is None:
            relabeled += 1
    return updates, deletes, {"relabeled": relabeled, "merged": merged,
                              "deferred": deferred, "skipped": skipped}


@app.post("/admin/concepts/consolidate-subjects-fast", tags=["Knowledge"])
def consolidate_subjects_fast(after_id: int = 0, limit: int = 20000,
                              apply: bool = False, token: str = "",
                              db: Session = Depends(get_db)):
    """SET-BASED version of consolidate-subjects: identical relabel+merge
    semantics, but a whole id-range per request using a handful of bulk SQL
    statements instead of 1-2 network round-trips per row (the per-row version
    needs ~12h for a 200k library over a remote Postgres; this needs minutes).
    Token-guarded with the BACKUP_TOKEN secret (like export-backup) so an
    automated runner can drive it without a 12h admin session. Postgres only.
    Call repeatedly with after_id=last_id until done=true; apply=false is a
    free dry run. Idempotent and safe to re-run over any range."""
    import os as _os
    import time as _time
    t0 = _time.time()
    expected = _os.getenv("BACKUP_TOKEN", "").strip()
    if not expected or token != expected:
        raise HTTPException(status_code=403, detail="forbidden - bad or missing token")
    if db.get_bind().dialect.name != "postgresql":
        raise HTTPException(status_code=400, detail="postgres only - use consolidate-subjects")
    limit = max(100, min(int(limit), 20000))
    cols = (DBConceptInventory.id, DBConceptInventory.concept, DBConceptInventory.subject,
            DBConceptInventory.key, DBConceptInventory.subtopic, DBConceptInventory.importance,
            DBConceptInventory.frequency, DBConceptInventory.sources, DBConceptInventory.key_facts)
    names = ("id", "concept", "subject", "key", "subtopic", "importance",
             "frequency", "sources", "key_facts")
    raw = (db.query(*cols).filter(DBConceptInventory.id > int(after_id))
           .order_by(DBConceptInventory.id).limit(limit).all())
    if not raw:
        return {"status": "success", "done": True, "last_id": int(after_id), "scanned": 0,
                "relabeled": 0, "merged": 0, "deferred": 0, "apply": bool(apply)}
    batch = [dict(zip(names, r)) for r in raw]
    last_id = batch[-1]["id"]
    # ONE indexed query fetches every possible collision target for the batch.
    nkeys = set()
    for r in batch:
        canon = _canonical_subject(r["subject"])
        nkeys.add(((r["concept"] or "").strip().lower() + "|" + canon.lower())[:400])
    existing_by_key, nk = {}, list(nkeys)
    for i in range(0, len(nk), 5000):
        for row in db.query(*cols).filter(DBConceptInventory.key.in_(nk[i:i + 5000])).all():
            d = dict(zip(names, row))
            existing_by_key[d["key"]] = d
    updates, deletes, stats = _plan_subject_consolidation(batch, existing_by_key)
    if apply and (updates or deletes):
        cur = db.connection().connection.cursor()   # raw driver cursor: true bulk statements
        if deletes:
            cur.execute("DELETE FROM concept_inventory WHERE id = ANY(%s)", (deletes,))
        CH = 3000                                    # 8 params/row - stay far below the 65k cap
        for i in range(0, len(updates), CH):
            part = updates[i:i + CH]
            vals = ",".join(["(%s,%s,%s,%s,%s,%s,%s,%s)"] * len(part))
            params = []
            for u in part:
                params += [u["id"], u["subject"], u["key"], int(u["frequency"]),
                           int(u["sources"]), u["key_facts"], u["subtopic"], u["importance"]]
            cur.execute(
                "UPDATE concept_inventory AS c SET"
                " subject=v.subject, key=v.key, frequency=v.frequency::int,"
                " sources=v.sources::int, key_facts=v.key_facts,"
                " subtopic=COALESCE(c.subtopic, v.subtopic),"
                " importance=COALESCE(c.importance, v.importance), updated_at=NOW()"
                " FROM (VALUES " + vals + ") AS v(id, subject, key, frequency, sources,"
                " key_facts, subtopic, importance) WHERE c.id = v.id",
                params)
        db.commit()
    else:
        db.rollback()
    return {"status": "success", "done": len(batch) < limit, "last_id": last_id,
            "scanned": len(batch), "relabeled": stats["relabeled"], "merged": stats["merged"],
            "deferred": stats["deferred"], "skipped_already_canonical": stats["skipped"],
            "apply": bool(apply), "ms": int((_time.time() - t0) * 1000)}


@app.get("/admin/embeddings/health", tags=["Knowledge"])
def embeddings_health(admin: DBUser = Depends(require_admin), db: Session = Depends(get_db)):
    """Diagnose semantic-search health of the knowledge base.

    Reports the active (pinned) embedding provider and chunk coverage, and runs a
    SAME-SPACE self-check: it takes one stored chunk, re-embeds that chunk's own
    text through the real query path, and measures cosine against its stored
    vector. ~1.0 means the index and queries share one vector space; a low value
    means the stored vectors were made by a different provider than queries now
    use — the silent failure mode. Fix: set EMBED_PROVIDER, then reset + backfill."""
    # New backend status first
    _new = {}
    try:
        import hf_embeddings, vector_store
        if vector_store.enabled():
            pts = vector_store.count()
            _new = {"rag_backend": "qdrant+hf",
                    "hf_embed_model": hf_embeddings.HF_EMBED_MODEL,
                    "hf_model_loaded": hf_embeddings.available(),
                    "qdrant_collection": vector_store.COLLECTION,
                    "qdrant_points": pts,
                    "qdrant_error": vector_store.LAST_ERROR or "",
                    "hf_error": hf_embeddings.LAST_ERROR or ""}
    except Exception as e:
        _new = {"rag_backend_error": f"{type(e).__name__}: {str(e)[:160]}"}
    import gemini_service
    try:
        provider = gemini_service.embed_provider()
    except Exception:
        provider = "unknown"
    total = db.execute(sa_text("SELECT COUNT(*) FROM knowledge_chunks")).scalar() or 0
    embedded = 0
    if VECTOR_OK:
        embedded = db.execute(sa_text(
            "SELECT COUNT(*) FROM knowledge_chunks WHERE embedding IS NOT NULL")).scalar() or 0
    out = {
        "vector_enabled": bool(VECTOR_OK),
        "embed_provider": provider,
        "embed_provider_pinned": (os.getenv("EMBED_PROVIDER", "").strip().lower() or None),
        "embed_model": getattr(gemini_service, "EMBED_MODEL", ""),
        "chunks_total": int(total),
        "chunks_embedded": int(embedded),
        "chunks_missing": int(total) - int(embedded),
        "last_embed_error": getattr(gemini_service, "LAST_EMBED_ERROR", ""),
    }
    out.update(_new)
    if not _new.get("rag_backend"):
        out["rag_backend"] = "pgvector+api" if VECTOR_OK else "keyword-only"
    if VECTOR_OK and embedded:
        try:
            row = db.execute(sa_text(
                "SELECT id, text FROM knowledge_chunks "
                "WHERE embedding IS NOT NULL AND length(text) > 40 ORDER BY id LIMIT 1")).fetchone()
            if row:
                cid, ctext = row[0], row[1]
                qv = gemini_service.embed_query((ctext or "")[:2000])
                if not qv:
                    out["self_check"] = {"error": "query embedding returned None (provider unavailable)"}
                else:
                    sim = db.execute(sa_text(
                        "SELECT 1 - (embedding <=> (:qv)::vector) FROM knowledge_chunks WHERE id = :id"),
                        {"qv": _vec_literal(qv), "id": cid}).scalar()
                    sim = float(sim) if sim is not None else None
                    ok = sim is not None and sim >= 0.70
                    out["self_check"] = {
                        "chunk_id": int(cid),
                        "self_cosine": round(sim, 4) if sim is not None else None,
                        "same_space": ok,
                        "verdict": ("OK — index and queries share one space"
                                    if ok else
                                    "MISMATCH — set EMBED_PROVIDER, then /admin/embeddings/reset + backfill"),
                    }
        except Exception as e:
            out["self_check"] = {"error": f"{type(e).__name__}: {str(e)[:160]}"}
    return out


@app.post("/admin/concepts/extract-pdf", tags=["Knowledge"])
async def extract_concepts_pdf(file: UploadFile = File(...),
                               admin: DBUser = Depends(require_admin)):
    """COPYRIGHT-SAFE concept extraction from a LAWFULLY-OBTAINED file.

    Reads the file (OCR for scans), then uses the LLM to identify ONLY the concept
    each item tests plus the underlying PUBLIC fact, subject, difficulty and
    question format. The source's verbatim questions, options and explanations are
    NEVER returned or stored — the extracted text is used transiently to classify
    and then discarded. Use the returned concept inventory to seed ORIGINAL
    question generation; do not re-store third-party question text.

    Intended for commercial test-series PDFs you have lawfully obtained; get legal
    review of your workflow before relying on it commercially."""
    import gemini_service
    import json as _json
    import re as _re

    fname = (file.filename or "upload").strip() or "upload"
    data = await file.read()
    if not data:
        raise HTTPException(status_code=400, detail="The uploaded file is empty.")
    pages, total, note = _extract_text_any(data, fname)
    text = "\n\n".join(t for _pno, t in pages)
    del data  # do not keep the source bytes
    if not text.strip():
        return {"status": "success", "file": fname, "note": note or "No extractable text.",
                "item_count": 0, "concepts": []}

    items = _extract_concept_inventory(text)

    # Zero-footprint guarantee: the source file and its extracted text lived only
    # in memory for this request. Actively release them; nothing is written to
    # disk or the database, and only the concept metadata below survives.
    text = ""
    pages = None
    import gc as _gc
    _gc.collect()

    return {"status": "success", "file": fname, "pages": total, "note": note or "",
            "item_count": len(items),
            "footprint": "none — source file and its text were held only in memory for this "
                         "request, never written to disk or the database; only the concept "
                         "metadata below persists.",
            "concepts": items}


# ── Background concept-extraction JOBS ───────────────────────────────────────
# Files are uploaded, held as a TEMPORARY base64 blob, and extracted on the
# SERVER in a background thread — so extraction continues even if the admin closes
# the tab or logs out. The temp blob is deleted the moment a file is done (or
# errors), leaving only the derived concept metadata. Jobs survive a restart
# (the blob is in the DB) and auto-resume on the next boot.

import tempfile as _tempfile
_CONCEPT_TMP = os.path.join(_tempfile.gettempdir(), "aivora_concept_jobs")
# Hard ceiling on a single upload. Not the old restrictive cap — this is a SAFETY
# limit: a file much larger than this cannot be read/OCR'd within the 512MB free
# tier and would crash the whole service. Raise it once on a bigger (paid) plan.
_MAX_CONCEPT_UPLOAD = 100 * 1024 * 1024   # 100 MB

def _concept_tmp_path(job_id):
    return os.path.join(_CONCEPT_TMP, f"job_{job_id}.bin")

def _process_concept_job(job_id: int):
    import gc as _gc, json as _json, os as _os, base64 as _b64
    db = SessionLocal()
    tmp_path = None
    try:
        from datetime import datetime as _dt
        job = db.query(DBConceptJob).filter(DBConceptJob.id == job_id).first()
        if not job or not job.raw_b64:
            return
        job.status = "processing"
        job.stage = "reading"
        job.progress = 5
        job.started_at = _dt.utcnow()
        db.commit()
        raw = job.raw_b64
        data = None
        if raw.startswith("file://"):
            tmp_path = raw[7:]
            if not _os.path.exists(tmp_path):
                # Temp source is gone (e.g. lost when a free-tier instance restarted).
                job.status = "error"
                job.error = "Source file was no longer available — please re-upload it."
                job.raw_b64 = None
                db.commit()
                return
        else:                                   # legacy base64 blob
            try:
                data = _b64.b64decode(raw)
            except Exception:
                data = b""
            if not data:
                job.status = "error"
                job.error = "Source was no longer available — please re-upload it."
                job.raw_b64 = None
                db.commit()
                return
        try:
            # A PDF handed as a temp path is opened MEMORY-MAPPED (pages read from
            # disk on demand), so even a large book stays within the memory budget.
            if tmp_path:
                pages, total, note = _extract_text_any(None, job.filename or "upload", path=tmp_path)
            else:
                pages, total, note = _extract_text_any(data, job.filename or "upload")
                del data
            text = "\n\n".join(t for _p, t in pages)
            pages = None
            _gc.collect()
            # Text is read → now extracting concepts (the long, LLM-bound stage).
            _j = db.query(DBConceptJob).filter(DBConceptJob.id == job_id).first()
            if _j:
                _j.stage = "extracting"; _j.progress = 35; db.commit()
            _last = [35]
            def _cb(done, total):
                pct = 35 + int(60 * done / max(1, total))     # 35% → 95% across the chunks
                if pct - _last[0] >= 3 or done >= total:
                    _last[0] = pct
                    try:
                        _jj = db.query(DBConceptJob).filter(DBConceptJob.id == job_id).first()
                        if _jj:
                            _jj.progress = min(97, pct); db.commit()
                    except Exception:
                        db.rollback()
            _stats = {}
            items = _extract_concept_inventory(text, progress_cb=_cb, stats=_stats) if text.strip() else []
            text = ""
            _gc.collect()
            _failed = int(_stats.get("failed") or 0)
            _tot = int(_stats.get("chunks") or 0)
            _quota = bool(_stats.get("quota"))
            job = db.query(DBConceptJob).filter(DBConceptJob.id == job_id).first()
            if job and _failed and not items:
                # Every section failed (e.g. Gemini free-tier quota exhausted).
                # Fail LOUDLY instead of reporting "done, 0 concepts".
                job.status = "error"
                job.stage = "error"
                job.progress = 100
                job.pages = int(total or 0)
                job.item_count = 0
                job.raw_b64 = None
                job.error = ("AI provider unavailable — "
                             + ("free-tier quota exhausted (429). " if _quota else "")
                             + f"0 of {_tot} sections analysed. Re-upload this file once "
                             "the quota resets (or after adding an OpenAI key).")
                db.commit()
            elif job:
                job.concepts = _json.dumps(items)
                job.item_count = len(items)
                job.pages = int(total or 0)
                job.status = "done"
                job.stage = "done"
                job.progress = 100
                job.raw_b64 = None            # drop the pointer; temp file deleted below
                # Partial coverage still counts as done, but say what was skipped.
                job.error = (f"Partial: {_failed} of {_tot} sections hit the AI rate "
                             "limit and were skipped — re-upload later to fill the gaps."
                             ) if _failed else None
                db.commit()
                try:                          # persist into the permanent concept library
                    _merge_concepts_into_inventory(db, items)
                except Exception as _me:
                    db.rollback()
                    print(f"[concept-inventory] merge failed for job {job_id}: {str(_me)[:160]}")
        except Exception as e:
            db.rollback()
            job = db.query(DBConceptJob).filter(DBConceptJob.id == job_id).first()
            if job:
                job.status = "error"
                job.stage = "error"
                job.error = f"{type(e).__name__}: {str(e)[:400]}"
                job.raw_b64 = None            # never retain the source, even on failure
                db.commit()
    finally:
        if tmp_path:                          # delete the temp file + its footprint
            try:
                if _os.path.exists(tmp_path):
                    _os.remove(tmp_path)
            except Exception:
                pass
        db.close()


# One extraction at a time, server-wide. Each upload request starts its own
# worker thread, so without this gate several 100+ page extractions would run
# CONCURRENTLY and together exceed the 512MB instance memory (this exact OOM
# killed the service three times on 2026-07-04). Threads queue on the lock —
# waiting jobs stay visible as "queued" and each begins the moment the previous
# one finishes. Throughput is the same; peak memory is one job, not N.
import threading as _gate_threading
_CONCEPT_JOB_GATE = _gate_threading.Lock()


def _merge_concepts_into_inventory(db, items):
    """Upsert one job's extracted concepts into the PERMANENT concept_inventory
    table (deduped by concept+subject). Returns how many brand-new concepts were
    added. Existing rows get their frequency/sources bumped and fact lists merged,
    so repeat uploads enrich the library instead of duplicating it."""
    import json as _json
    from datetime import datetime as _dt
    normed = {}
    for it in (items or []):
        c = (it.get("concept") or "").strip()
        if not c:
            continue
        it["subject"] = _canonical_subject(it.get("subject"))   # canonical GS subject
        k = (c.lower() + "|" + it["subject"].lower())[:400]
        normed[k] = it
    if not normed:
        return 0
    added = 0
    existing = {r.key: r for r in db.query(DBConceptInventory)
                .filter(DBConceptInventory.key.in_(list(normed.keys()))).all()}
    for k, it in normed.items():
        row = existing.get(k)
        if row:
            row.frequency = (row.frequency or 1) + int(it.get("frequency") or 1)
            row.sources = (row.sources or 1) + 1
            try:
                old = _json.loads(row.key_facts or "[]")
            except Exception:
                old = []
            merged = list(dict.fromkeys([*old, *(it.get("key_facts") or [])]))[:12]
            row.key_facts = _json.dumps(merged)
            row.updated_at = _dt.utcnow()
        else:
            db.add(DBConceptInventory(
                key=k,
                concept=(it.get("concept") or "").strip()[:300],
                subject=(it.get("subject") or None),
                subtopic=(it.get("subtopic") or None),
                pattern=(it.get("pattern") or None),
                difficulty=(it.get("difficulty") or None),
                importance=(it.get("importance") or None),
                key_facts=_json.dumps(it.get("key_facts") or []),
                frequency=int(it.get("frequency") or 1),
                sources=1))
            added += 1
    db.commit()
    return added


def _backfill_concept_inventory():
    """One-time recovery: if the permanent inventory is empty, merge in the
    concepts of every already-finished job (loads ONE job row at a time to stay
    within the free tier's memory budget)."""
    try:
        import json as _json
        db = SessionLocal()
        if db.query(DBConceptInventory.id).first() is not None:
            db.close()
            return                                    # already populated — nothing to do
        ids = [r[0] for r in db.query(DBConceptJob.id)
               .filter(DBConceptJob.status == "done").order_by(DBConceptJob.id).all()]
        total_new = 0
        for jid in ids:
            row = db.query(DBConceptJob).filter(DBConceptJob.id == jid).first()
            items = []
            if row and row.concepts:
                try:
                    items = _json.loads(row.concepts)
                except Exception:
                    items = []
            db.expunge_all()
            total_new += _merge_concepts_into_inventory(db, items)
        db.close()
        print(f"[concept-inventory] backfilled {total_new} concepts from {len(ids)} finished jobs")
    except Exception as e:
        print(f"[concept-inventory] backfill skipped: {type(e).__name__}: {str(e)[:160]}")


def _process_concept_batch(job_ids):
    for jid in job_ids:
        try:
            with _CONCEPT_JOB_GATE:
                _process_concept_job(jid)
        except Exception as e:
            print(f"[concept-job] {jid} failed: {type(e).__name__}: {str(e)[:160]}")


# ── Stay-awake while extracting ────────────────────────────────────────────────
# Render's free tier spins the instance down after ~15 min without INBOUND web
# traffic — background threads don't count, so an extraction left running with
# the admin tab closed used to die mid-job. While any concept job is queued or
# processing, this thread pings our own public /health endpoint every 4 minutes,
# which registers as inbound traffic and keeps the instance alive. The moment
# the queue is empty the thread exits and normal spin-down behaviour resumes.
_KEEPALIVE_FLAG = {"on": False}
_KEEPALIVE_LOCK = _gate_threading.Lock()


def _concept_keepalive():
    import time as _time
    base = (os.environ.get("RENDER_EXTERNAL_URL") or "https://ias-mentor-ai.onrender.com").rstrip("/")
    try:
        while True:
            db = SessionLocal()
            try:
                active = db.query(DBConceptJob).filter(
                    DBConceptJob.status.in_(["queued", "processing"])).count()
            except Exception:
                active = 0
            finally:
                db.close()
            if not active:
                break
            try:
                urlopen(base + "/health", timeout=20).read(128)
            except Exception:
                pass                        # a missed ping is harmless; try again next round
            for _ in range(24):             # 24 × 10s = 4 min, in short steps so we exit fast
                _time.sleep(10)
    finally:
        with _KEEPALIVE_LOCK:
            _KEEPALIVE_FLAG["on"] = False


def _ensure_keepalive():
    with _KEEPALIVE_LOCK:
        if _KEEPALIVE_FLAG["on"]:
            return
        _KEEPALIVE_FLAG["on"] = True
    _gate_threading.Thread(target=_concept_keepalive, daemon=True).start()


def _resume_concept_jobs():
    """On boot, do NOT auto-retry unfinished jobs. A huge file can crash the worker,
    and — critically — retrying it would reload its (possibly hundreds-of-MB) temp
    blob into memory and OOM again, crash-looping the whole service. So we bulk-mark
    every unfinished job as errored and NULL its temp blob using a single SQL UPDATE
    that NEVER loads the row (or its blob) into memory. Nothing loops; the admin just
    re-uploads if needed. (Any orphaned on-disk temp files are on ephemeral storage
    and vanish with the restart.)"""
    try:
        db = SessionLocal()
        db.query(DBConceptJob).filter(
            DBConceptJob.status.in_(["queued", "processing"])).update(
            {DBConceptJob.status: "error",
             DBConceptJob.error: "Interrupted by a server restart — please re-upload this file.",
             DBConceptJob.raw_b64: None},
            synchronize_session=False)
        db.commit()
        db.close()
    except Exception as e:
        print(f"[concept-job] resume cleanup skipped: {type(e).__name__}: {str(e)[:160]}")


@app.on_event("startup")
def _on_startup_resume_jobs():
    _resume_concept_jobs()
    # Recover concepts from finished jobs into the permanent inventory (no-op once
    # populated). Runs in a background thread so boot stays fast.
    _gate_threading.Thread(target=_backfill_concept_inventory, daemon=True).start()


@app.post("/admin/concepts/extract-jobs", tags=["Knowledge"])
async def create_concept_jobs(files: List[UploadFile] = File(...),
                              admin: DBUser = Depends(require_admin),
                              db: Session = Depends(get_db)):
    """Queue one or more files for SERVER-SIDE concept extraction. Each upload is
    STREAMED to a temporary file on disk (never held whole in memory + base64, which
    used to OOM the free instance on large books), then deleted the moment the file
    is processed. Returns a batch id; poll /admin/concepts/jobs for progress."""
    import threading, uuid as _uuid
    os.makedirs(_CONCEPT_TMP, exist_ok=True)
    batch_id = _uuid.uuid4().hex[:12]
    jobs, ids = [], []
    for f in files:
        fname = (f.filename or "upload").strip() or "upload"
        # Create the row first so we have an id for the temp filename.
        job = DBConceptJob(batch_id=batch_id, filename=fname[:200], status="queued",
                           uploaded_by=admin.id)
        db.add(job); db.commit(); db.refresh(job)
        path = _concept_tmp_path(job.id)
        try:
            total = 0
            too_big = False
            with open(path, "wb") as out:
                while True:
                    chunk = await f.read(1024 * 1024)   # stream 1 MB at a time (low memory)
                    if not chunk:
                        break
                    total += len(chunk)
                    if total > _MAX_CONCEPT_UPLOAD:      # reject before it can OOM the server
                        too_big = True
                        break
                    out.write(chunk)
            if too_big:
                try:
                    os.remove(path)
                except Exception:
                    pass
                job.status = "error"
                job.error = ("File is over the 100 MB limit for the current plan — a file this large "
                             "runs the free-tier server out of memory. Split it into smaller PDFs, or "
                             "upgrade the server for more RAM.")
                db.commit()
                jobs.append({"id": job.id, "filename": fname, "status": "error"})
                continue
            if total <= 0:
                try:
                    os.remove(path)
                except Exception:
                    pass
                db.delete(job); db.commit()
                continue
            job.raw_b64 = "file://" + path        # pointer to the temp file (not base64)
            db.commit()
            ids.append(job.id)
            jobs.append({"id": job.id, "filename": fname, "status": "queued"})
        except Exception as e:
            try:
                if os.path.exists(path):
                    os.remove(path)
            except Exception:
                pass
            job.status = "error"
            job.error = f"Upload failed: {str(e)[:200]}"
            db.commit()
            jobs.append({"id": job.id, "filename": fname, "status": "error"})
    if ids:
        threading.Thread(target=_process_concept_batch, args=(ids,), daemon=True).start()
        _ensure_keepalive()               # keep the free instance awake until the queue drains
    return {"status": "success", "batch_id": batch_id, "count": len(jobs), "jobs": jobs,
            "message": ("Extraction is running on the server — you can close this tab or log "
                        "out and it will finish on its own.")}


@app.get("/admin/concepts/jobs", tags=["Knowledge"])
def list_concept_jobs(batch_id: Optional[str] = None, limit: int = 60,
                      admin: DBUser = Depends(require_admin), db: Session = Depends(get_db)):
    """Status + results for concept-extraction jobs. Poll this to watch progress;
    finished jobs carry their concept metadata (the source file is already gone)."""
    # IMPORTANT: select ONLY light columns. Loading every job's `concepts` TEXT
    # (hundreds of KB each × 200 rows) allocated tens of MB per poll and OOM'd
    # the 512MB instance on 2026-07-06 once the library grew large. Concepts are
    # served by /admin/concepts/inventory/export (streamed) instead.
    q = db.query(DBConceptJob.id, DBConceptJob.batch_id, DBConceptJob.filename,
                 DBConceptJob.status, DBConceptJob.item_count, DBConceptJob.pages,
                 DBConceptJob.error, DBConceptJob.stage, DBConceptJob.progress,
                 DBConceptJob.started_at, DBConceptJob.created_at
                 ).order_by(DBConceptJob.id.desc())
    if batch_id:
        q = q.filter(DBConceptJob.batch_id == batch_id)
    rows = q.limit(max(1, min(limit, 200))).all()
    out = [{
        "id": r.id, "batch_id": r.batch_id, "filename": r.filename,
        "status": r.status, "item_count": r.item_count or 0, "pages": r.pages or 0,
        "error": r.error, "concepts": [],
        "stage": r.stage,
        "progress": int(r.progress or 0),
        "started_at": r.started_at.isoformat() if r.started_at else None,
        "created_at": r.created_at.isoformat() if r.created_at else None,
    } for r in rows]
    busy = sum(1 for r in rows if r.status in ("queued", "processing"))
    return {"status": "success", "processing": busy, "jobs": out}


@app.get("/admin/concepts/inventory/export", tags=["Knowledge"])
def export_concept_inventory(admin: DBUser = Depends(require_admin)):
    """Stream the ENTIRE permanent concept library as a JSON download. Streamed
    row-by-row so even a 100k-concept library never builds a big buffer in RAM."""
    from fastapi.responses import StreamingResponse

    def _gen():
        _db = SessionLocal()
        try:
            yield "["
            first = True
            for row in _db.query(DBConceptInventory).yield_per(500):
                try:
                    facts = json.loads(row.key_facts or "[]")
                except Exception:
                    facts = []
                item = {"concept": row.concept, "subject": row.subject,
                        "subtopic": row.subtopic, "pattern": row.pattern,
                        "difficulty": row.difficulty, "importance": row.importance,
                        "key_facts": facts, "frequency": row.frequency or 1,
                        "sources": row.sources or 1}
                yield ("" if first else ",") + json.dumps(item, ensure_ascii=False)
                first = False
            yield "]"
        finally:
            _db.close()

    return StreamingResponse(_gen(), media_type="application/json",
                             headers={"Content-Disposition":
                                      "attachment; filename=concept_inventory.json"})


@app.get("/admin/concepts/export-backup", tags=["Knowledge"])
def export_backup(token: str = ""):
    """Token-guarded FULL concept-library export for automated off-site backup.
    Open (no admin login) but requires ?token=<BACKUP_TOKEN secret>, so a scheduled
    job can pull a JSON snapshot without a session. Streamed row-by-row (safe for a
    200k+ library). Returns 403 if the token is unset or wrong."""
    import os as _os
    from fastapi.responses import StreamingResponse, JSONResponse
    expected = _os.getenv("BACKUP_TOKEN", "").strip()
    if not expected or token != expected:
        return JSONResponse({"error": "forbidden - bad or missing token"}, status_code=403)

    def _gen():
        _db = SessionLocal()
        try:
            yield "["
            first = True
            for row in _db.query(DBConceptInventory).yield_per(500):
                try:
                    facts = json.loads(row.key_facts or "[]")
                except Exception:
                    facts = []
                item = {"concept": row.concept, "subject": row.subject,
                        "subtopic": row.subtopic, "pattern": row.pattern,
                        "difficulty": row.difficulty, "importance": row.importance,
                        "key_facts": facts, "frequency": row.frequency or 1,
                        "sources": row.sources or 1}
                yield ("" if first else ",") + json.dumps(item, ensure_ascii=False)
                first = False
            yield "]"
        finally:
            _db.close()

    return StreamingResponse(_gen(), media_type="application/json",
                             headers={"Content-Disposition":
                                      "attachment; filename=concept_inventory_backup.json"})


@app.post("/admin/concepts/jobs/clear", tags=["Knowledge"])
def clear_concept_jobs(batch_id: Optional[str] = None, only_finished: bool = True,
                       admin: DBUser = Depends(require_admin), db: Session = Depends(get_db)):
    """Delete extraction jobs (and their stored concept metadata). By default only
    finished/errored jobs are removed; running jobs are left alone."""
    q = db.query(DBConceptJob)
    if batch_id:
        q = q.filter(DBConceptJob.batch_id == batch_id)
    if only_finished:
        q = q.filter(DBConceptJob.status.in_(["done", "error"]))
    n = q.count()
    q.delete(synchronize_session=False)
    db.commit()
    return {"status": "success", "cleared": int(n)}


@app.get("/knowledge/search", tags=["Knowledge"])
def knowledge_search(q: str = "", subject: Optional[str] = None, limit: int = 20,
                     current_user: DBUser = Depends(get_current_user), db: Session = Depends(get_db)):
    qq = (q or "").strip()
    if len(qq) < 2:
        return {"status": "success", "count": 0, "results": []}
    query = (db.query(DBKnowledgeChunk, DBKnowledgeSource)
             .join(DBKnowledgeSource, DBKnowledgeChunk.source_id == DBKnowledgeSource.id)
             .filter(DBKnowledgeChunk.text.ilike(f"%{qq}%")))
    if subject and subject.lower() not in ("", "all"):
        query = query.filter(DBKnowledgeChunk.subject == subject)
    rows = query.limit(max(1, min(limit, 50))).all()
    return {"status": "success", "count": len(rows),
            "results": [{"text": c.text, "subject": c.subject, "page": c.page,
                         "source": s.filename} for c, s in rows]}

@app.get("/knowledge/subjects", tags=["Knowledge"])
def knowledge_subjects(current_user: DBUser = Depends(get_current_user), db: Session = Depends(get_db)):
    subs = [r[0] for r in db.query(DBKnowledgeChunk.subject).distinct().all() if r[0]]
    return {"status": "success", "subjects": sorted(subs), "has_library": bool(subs)}


# ── Concept Explorer: the student-facing window into the concept library ─────
def _serialize_concept(r):
    try:
        facts = json.loads(r.key_facts or "[]")
    except Exception:
        facts = []
    return {"id": r.id, "concept": r.concept, "subject": r.subject,
            "subtopic": r.subtopic, "pattern": r.pattern, "difficulty": r.difficulty,
            "importance": r.importance, "key_facts": facts[:8],
            "frequency": r.frequency or 1}


@app.get("/concepts/subjects", tags=["Knowledge"])
def concept_explorer_subjects(current_user: DBUser = Depends(get_current_user),
                              db: Session = Depends(get_db)):
    """Clean canonical subject list + library size, for the explorer's filters."""
    from sqlalchemy import func as _f
    rows = (db.query(DBConceptInventory.subject, _f.count(DBConceptInventory.id))
            .group_by(DBConceptInventory.subject).all())
    subs = sorted([{"subject": s or "Other", "count": int(n)} for s, n in rows],
                  key=lambda x: -x["count"])
    return {"status": "success", "subjects": subs,
            "total_concepts": sum(x["count"] for x in subs)}


@app.get("/concepts/explore", tags=["Knowledge"])
def concepts_explore(q: str = "", subject: Optional[str] = None, limit: int = 24,
                     current_user: DBUser = Depends(get_current_user),
                     db: Session = Depends(get_db)):
    """Search/browse the concept library, most-tested first. With no query it
    surfaces the highest-frequency (most exam-relevant) concepts — a syllabus
    'greatest hits' view; with a query it's a straight concept search."""
    query = db.query(DBConceptInventory)
    qq = (q or "").strip()
    if qq:
        query = query.filter(DBConceptInventory.concept.ilike(f"%{qq}%"))
    if subject and subject.lower() not in ("", "all"):
        query = query.filter(DBConceptInventory.subject == subject)
    rows = (query.order_by(DBConceptInventory.frequency.desc(), DBConceptInventory.id)
            .limit(max(1, min(int(limit), 50))).all())
    return {"status": "success", "count": len(rows),
            "results": [_serialize_concept(r) for r in rows]}


@app.get("/concepts/related", tags=["Knowledge"])
def concepts_related(concept_id: int, limit: int = 8,
                     current_user: DBUser = Depends(get_current_user),
                     db: Session = Depends(get_db)):
    """Concepts linked to this one: same subtopic first (the strongest signal in
    the library), then concepts sharing a significant word of the name. This is
    the interlinking backbone v1 — 'you're studying X, the exam also tests Y'."""
    import re as _re
    base = db.query(DBConceptInventory).filter(DBConceptInventory.id == int(concept_id)).first()
    if not base:
        raise HTTPException(status_code=404, detail="Concept not found")
    out, seen = [], {base.id}

    def take(rows):
        for r in rows:
            if r.id not in seen:
                seen.add(r.id)
                out.append(_serialize_concept(r))
            if len(out) >= limit:
                return True
        return False

    if base.subtopic:
        rows = (db.query(DBConceptInventory)
                .filter(DBConceptInventory.subtopic == base.subtopic,
                        DBConceptInventory.subject == base.subject,
                        DBConceptInventory.id != base.id)
                .order_by(DBConceptInventory.frequency.desc()).limit(limit * 2).all())
        take(rows)
    if len(out) < limit:
        stop = {"the", "and", "for", "with", "from", "into", "india", "indian", "of", "in"}
        words = [w for w in _re.findall(r"[a-zA-Z]{4,}", base.concept or "")
                 if w.lower() not in stop][:3]
        for w in words:
            if len(out) >= limit:
                break
            rows = (db.query(DBConceptInventory)
                    .filter(DBConceptInventory.concept.ilike(f"%{w}%"),
                            DBConceptInventory.id != base.id)
                    .order_by(DBConceptInventory.frequency.desc()).limit(limit).all())
            if take(rows):
                break
    return {"status": "success", "concept": _serialize_concept(base),
            "related": out[:limit]}


# ── Phase 4: Content Depth (notes · flashcards · mnemonics · mind map · CA) ────
class ContentRequest(BaseModel):
    topic: str
    subject: Optional[str] = ""
    count: Optional[int] = 10

class CARequest(BaseModel):
    event: str

@app.post("/content/notes", tags=["Content"])
def content_notes(req: ContentRequest, current_user: DBUser = Depends(get_current_user)):
    if not (req.topic or "").strip():
        raise HTTPException(status_code=400, detail="Topic is required")
    try:
        return {"status": "success", "notes": generate_study_notes(req.topic.strip(), req.subject or "")}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Generation error: {str(e)}")

@app.post("/content/flashcards", tags=["Content"])
def content_flashcards(req: ContentRequest, current_user: DBUser = Depends(get_current_user)):
    if not (req.topic or "").strip():
        raise HTTPException(status_code=400, detail="Topic is required")
    try:
        cards = generate_flashcards(req.topic.strip(), req.subject or "", req.count or 10)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Generation error: {str(e)}")
    if not cards:
        raise HTTPException(status_code=502, detail="Could not generate flashcards. Try again.")
    return {"status": "success", "count": len(cards), "cards": cards}

@app.post("/content/mnemonics", tags=["Content"])
def content_mnemonics(req: ContentRequest, current_user: DBUser = Depends(get_current_user)):
    if not (req.topic or "").strip():
        raise HTTPException(status_code=400, detail="Topic is required")
    try:
        return {"status": "success", "mnemonics": generate_mnemonics(req.topic.strip(), req.subject or "")}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Generation error: {str(e)}")

@app.post("/content/mindmap", tags=["Content"])
def content_mindmap(req: ContentRequest, current_user: DBUser = Depends(get_current_user)):
    if not (req.topic or "").strip():
        raise HTTPException(status_code=400, detail="Topic is required")
    try:
        return {"status": "success", "tree": generate_mindmap(req.topic.strip(), req.subject or "")}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Generation error: {str(e)}")

@app.post("/content/current-affairs", tags=["Content"])
def content_current_affairs(req: CARequest, current_user: DBUser = Depends(get_current_user)):
    if not (req.event or "").strip():
        raise HTTPException(status_code=400, detail="Event/topic is required")
    try:
        return {"status": "success", "analysis": current_affairs_analysis(req.event.strip())}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Generation error: {str(e)}")


# ── Leaderboard ───────────────────────────────────────────────────────────────
@app.get("/leaderboard/", tags=["Results"])
def get_leaderboard(db: Session = Depends(get_db), current_user: DBUser = Depends(get_current_user)):
    users = db.query(DBUser).all()
    board = []
    for u in users:
        attempts = db.query(DBTestAttempt).filter(DBTestAttempt.user_id == u.id).all()
        if not attempts:
            continue
        avg = sum(
            (a.score / a.mock_test.total_questions * 100) for a in attempts if a.mock_test.total_questions
        ) / len(attempts)
        board.append({
            "name": u.name or u.email.split('@')[0],
            "email": u.email,
            "tests_taken": len(attempts),
            "avg_score": round(avg, 1),
            "is_me": u.id == current_user.id,
        })
    board.sort(key=lambda x: x['avg_score'], reverse=True)
    for i, entry in enumerate(board):
        entry['rank'] = i + 1
    return {"status": "success", "leaderboard": board}

# ── Admin dashboard ───────────────────────────────────────────────────────────
def _attempt_metrics(db, attempt):
    total = attempt.mock_test.total_questions or 0 if attempt.mock_test else 0
    attempted = db.query(DBAnswer).filter(DBAnswer.test_attempt_id == attempt.id).count()
    correct = attempt.score or 0
    wrong = max(0, attempted - correct)
    pct = round(correct / total * 100, 1) if total else 0.0
    net = round(correct * 2.0 - wrong * (2.0 / 3.0), 2)
    net_pct = round(net / (total * 2.0) * 100, 1) if total else 0.0
    accuracy = round(correct / attempted * 100, 1) if attempted else 0.0
    return {"attempted": attempted, "correct": correct, "wrong": wrong, "total": total,
            "pct": pct, "net": net, "net_pct": net_pct, "accuracy": accuracy}

@app.get("/admin/overview", tags=["Admin"])
def admin_overview(admin: DBUser = Depends(require_admin), db: Session = Depends(get_db)):
    from collections import Counter
    users = db.query(DBUser).filter(DBUser.email != SYSTEM_USER_EMAIL).all()
    attempts = db.query(DBTestAttempt).all()
    pcts = [(a.score or 0) / a.mock_test.total_questions * 100
            for a in attempts if a.mock_test and a.mock_test.total_questions]
    cutoff = datetime.datetime.utcnow() - datetime.timedelta(days=7)
    active7 = len({a.user_id for a in attempts if a.completed_at and a.completed_at >= cutoff})
    cnt = Counter(a.mock_test.title for a in attempts if a.mock_test)
    return {
        "status": "success",
        "total_candidates": len(users),
        "total_attempts": len(attempts),
        "avg_percentage": round(sum(pcts) / len(pcts), 1) if pcts else 0.0,
        "active_last_7_days": active7,
        "most_attempted": [{"title": t, "attempts": c} for t, c in cnt.most_common(5)],
    }

@app.get("/admin/prelims-stats", tags=["Admin"])
def admin_prelims_stats(admin: DBUser = Depends(require_admin), db: Session = Depends(get_db)):
    """Exam-oriented question-count datasets for the Prelims Dashboard:
    (1) NCERT bookwise (with chapters), (2) reference-book chapterwise,
    (3) subject-topicwise."""
    from sqlalchemy import text
    from collections import OrderedDict
    def rows(sql):
        return db.execute(text(sql)).fetchall()
    total = db.execute(text("SELECT count(*) FROM questions")).scalar() or 0
    NCERT = "book IS NOT NULL AND book LIKE '%(Class%'"
    REF = "book IS NOT NULL AND book NOT LIKE '%(Class%'"

    ncert = OrderedDict()
    for book, subj, cnt in rows(
        "SELECT book, coalesce(subject,'-') s, count(*) c FROM questions "
        "WHERE " + NCERT + " GROUP BY book, subject ORDER BY c DESC"):
        ncert.setdefault(book, {"book": book, "subject": subj, "total": 0, "chapters": []})
        ncert[book]["total"] += cnt
    for book, ch, cnt in rows(
        "SELECT book, coalesce(chapter,'(whole book)') ch, count(*) c FROM questions "
        "WHERE " + NCERT + " GROUP BY book, chapter"):
        if book in ncert:
            ncert[book]["chapters"].append({"name": ch, "count": cnt})
    for v in ncert.values():
        v["chapters"].sort(key=lambda x: -x["count"])

    ref = OrderedDict()
    for book, ch, cnt in rows(
        "SELECT book, coalesce(chapter,'(no chapter)') ch, count(*) c FROM questions "
        "WHERE " + REF + " GROUP BY book, chapter ORDER BY book, c DESC"):
        ref.setdefault(book, {"book": book, "total": 0, "chapters": []})
        ref[book]["total"] += cnt
        ref[book]["chapters"].append({"name": ch, "count": cnt})

    subj = OrderedDict()
    for s, t, cnt in rows(
        "SELECT coalesce(subject,'(no subject)') s, coalesce(topic,'(general / untagged)') t, "
        "count(*) c FROM questions GROUP BY 1,2 ORDER BY 1, c DESC"):
        subj.setdefault(s, {"subject": s, "total": 0, "topics": []})
        subj[s]["total"] += cnt
        subj[s]["topics"].append({"name": t, "count": cnt})
    subjects = sorted(subj.values(), key=lambda x: -x["total"])

    return {
        "status": "success",
        "generated_at": datetime.datetime.utcnow().isoformat() + "Z",
        "total": total,
        "ncert_total": sum(v["total"] for v in ncert.values()),
        "ref_total": sum(v["total"] for v in ref.values()),
        "ncert_books": list(ncert.values()),
        "reference_books": list(ref.values()),
        "subjects": subjects,
    }


# ── NCERT page-citation (server-side, no external shell/secret needed) ────────
# Appends a page-accurate line like
#   Source: NCERT · Class 8 · Our Pasts III · Ch.5 When People Rebel · p.61
# to each question's explanation, by matching the question text against the stored
# chapter PDF (PyMuPDF) and detecting the printed page number (folio).
_CITE_MARK = "Source: NCERT"
_CITE_DOT = "·"
_CITE_STOP = set("""a an the of to in on at by for and or nor but so yet is are was were be been being
as it its this that these those with from into within without about above below over under which who
whom whose what when where why how all any both each few more most other some such only own same than
too very can will just should now not no do does did done has have had having their they them he she his
her him you your we our us my mine ours yours also may might must shall would could there here then once
during before after between against because while across upon per via out up down off following include
including etc vs versus among along around due much many one two three four them then""".split())

def _cite_tokens(s):
    s = (s or "").lower()
    return [t for t in re.findall(r"[a-z][a-z0-9\-']{2,}", s) if t not in _CITE_STOP]

def _cite_correct_opt(q):
    m = {"a": q.option_a, "b": q.option_b, "c": q.option_c, "d": q.option_d}
    return m.get((q.correct_answer or "").strip().lower()[:1], "") or ""

def _cite_bare_book(name):
    return re.sub(r"\s*\(Class[^)]*\)\s*$", "", name or "").strip()

def _cite_load_chapter_bytes(db, book_key, ch_index):
    row = (db.query(DBNcertPdf)
           .filter(DBNcertPdf.book_key == book_key, DBNcertPdf.chapter_index == ch_index)
           .first())
    if not row:
        return None
    if row.data:
        return bytes(row.data)
    if row.src_url:
        data = _ncert_cache_get(row.src_url)
        if data is None:
            try:
                data = _http_get(row.src_url)
            except Exception:
                return None
            if not data or data[:4] != b"%PDF":
                return None
            _ncert_cache_put(row.src_url, data)
        return data
    return None

def _cite_page_folio(page, H):
    """Printed page number (folio) from the top/bottom edge zones of a page."""
    try:
        words = page.get_text("words")   # (x0,y0,x1,y1, word, block, line, wordno)
    except Exception:
        return None
    best = None
    for w in words:
        y0, y1, tok = w[1], w[3], (w[4] or "").strip()
        if not (tok.isascii() and tok.isdigit()):   # reject unicode super/subscript digits
            continue
        try:
            val = int(tok)
        except ValueError:
            continue
        if val <= 0 or val > 2000:
            continue
        if not (y0 < H * 0.11 or y1 > H * 0.89):
            continue
        score = 1 if (y1 > H * 0.89) else 0     # prefer bottom folios
        if best is None or score > best[0]:
            best = (score, val)
    return best[1] if best else None

def _cite_chapter_pages(pdf_bytes):
    """[{seq, folio, toks}] for a chapter PDF, folios reconciled to page sequence so
    a few mis-detections don't corrupt citations."""
    import fitz
    from collections import Counter as _C
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    pages = []
    try:
        for i in range(doc.page_count):
            pg = doc.load_page(i)
            H = pg.rect.height or 1.0
            pages.append({"seq": i + 1, "folio": _cite_page_folio(pg, H),
                          "toks": _cite_tokens(pg.get_text() or "")})
            pg = None
    finally:
        doc.close()
    offs = _C()
    for p in pages:
        if p["folio"] is not None:
            offs[p["folio"] - p["seq"]] += 1
    off = offs.most_common(1)[0][0] if offs else 0
    for p in pages:                              # derive folio where detection missing/outlier
        exp = p["seq"] + off
        if p["folio"] is None or p["folio"] != exp:
            p["folio"] = exp
    return pages

def _cite_best_folio(pages, q_toks):
    """Best-matching page's folio, scoring pages by idf-weighted token overlap."""
    if not q_toks or not pages:
        return None, 0.0
    import math
    df = {}
    for p in pages:
        for t in set(p["toks"]):
            df[t] = df.get(t, 0) + 1
    N = len(pages)
    qset = set(q_toks)
    best = (None, 0.0)
    for p in pages:
        common = qset & set(p["toks"])
        if not common:
            continue
        sc = sum(math.log((N + 1) / df.get(t, 1)) for t in common)
        if sc > best[1]:
            best = (p["folio"], sc)
    return best

class CiteChapterRequest(BaseModel):
    book_key: str
    chapter_index: int
    overwrite: Optional[bool] = False

@app.post("/admin/cite-chapter", tags=["Admin"])
def admin_cite_chapter(req: CiteChapterRequest,
                       admin: DBUser = Depends(require_admin), db: Session = Depends(get_db)):
    """Page-cite every question of ONE stored NCERT chapter. Idempotent: skips
    questions that already carry a 'Source: NCERT' line unless overwrite=true."""
    book = syllabus_data.get_ncert_book(req.book_key)
    if not book:
        raise HTTPException(status_code=404, detail="Unknown NCERT book key.")
    chapters = book.get("chapters") or []
    if req.chapter_index < 0 or req.chapter_index >= len(chapters):
        raise HTTPException(status_code=400, detail="chapter_index out of range.")
    title = chapters[req.chapter_index]
    book_disp = book["book"]
    grade = book.get("grade", "")
    bare = _cite_bare_book(book_disp)

    pdf = _cite_load_chapter_bytes(db, req.book_key, req.chapter_index)
    if not pdf:
        return {"status": "skipped", "reason": "chapter PDF not stored",
                "book": book_disp, "chapter": title, "chapter_index": req.chapter_index,
                "total": 0, "cited": 0, "already": 0, "weak": 0}
    try:
        pages = _cite_chapter_pages(pdf)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF parse failed: {str(e)[:120]}")
    folios = [p["folio"] for p in pages if p["folio"] is not None]
    mid_folio = pages[len(pages) // 2]["folio"] if pages else None

    qs = (db.query(DBQuestion)
          .filter(DBQuestion.book == book_disp, DBQuestion.chapter == title)
          .all())
    cited = already = weak = 0
    for q in qs:
        expl = q.explanation or ""
        if (_CITE_MARK in expl) and not req.overwrite:
            already += 1
            continue
        qtoks = _cite_tokens((q.text or "") + " " + _cite_correct_opt(q) + " " + (q.topic or ""))
        folio, sc = _cite_best_folio(pages, qtoks)
        if folio is None:                        # zero overlap → cite chapter mid-page
            folio = mid_folio
            weak += 1
        if folio is None:
            continue
        line = (f"{_CITE_MARK} {_CITE_DOT} {grade} {_CITE_DOT} {bare} {_CITE_DOT} "
                f"Ch.{req.chapter_index + 1} {title} {_CITE_DOT} p.{folio}")
        if req.overwrite and (_CITE_MARK in expl):
            expl = re.sub(r"\n*Source: NCERT.*$", "", expl, flags=re.S).rstrip()
        q.explanation = (expl.rstrip() + "\n\n" + line) if expl.strip() else line
        cited += 1
    db.commit()
    return {"status": "success", "book": book_disp, "chapter": title,
            "chapter_index": req.chapter_index, "grade": grade,
            "pages": len(pages),
            "folio_range": [min(folios), max(folios)] if folios else [None, None],
            "total": len(qs), "cited": cited, "already": already, "weak": weak}

def _cite_book_pages(db, book_key, book):
    """All stored pages of a book, each tagged with its chapter index + title,
    for whole-book (chapter-less) matching."""
    chapters = book.get("chapters") or []
    idxs = [r[0] for r in db.query(DBNcertPdf.chapter_index)
            .filter(DBNcertPdf.book_key == book_key)
            .order_by(DBNcertPdf.chapter_index).all()]
    allpages = []
    for ci in idxs:
        if ci < 0 or ci >= len(chapters):
            continue
        pdf = _cite_load_chapter_bytes(db, book_key, ci)
        if not pdf:
            continue
        try:
            for p in _cite_chapter_pages(pdf):
                p["ch"] = ci
                p["title"] = chapters[ci]
                allpages.append(p)
        except Exception:
            continue
    return allpages

def _cite_best_page(pages, q_toks):
    """Best-matching page object across a set of pages (idf-weighted overlap)."""
    if not q_toks or not pages:
        return None, 0.0
    import math
    df = {}
    for p in pages:
        for t in set(p["toks"]):
            df[t] = df.get(t, 0) + 1
    N = len(pages)
    qset = set(q_toks)
    best = (None, 0.0)
    for p in pages:
        common = qset & set(p["toks"])
        if not common:
            continue
        sc = sum(math.log((N + 1) / df.get(t, 1)) for t in common)
        if sc > best[1]:
            best = (p, sc)
    return best

class CiteLooseRequest(BaseModel):
    book_key: str
    overwrite: Optional[bool] = False

@app.post("/admin/cite-loose", tags=["Admin"])
def admin_cite_loose(req: CiteLooseRequest,
                     admin: DBUser = Depends(require_admin), db: Session = Depends(get_db)):
    """Page-cite a book's CHAPTER-LESS questions (chapter NULL/blank) by scanning
    every stored chapter of the book and attributing each question to its best page
    (and that page's chapter). Idempotent unless overwrite=true."""
    from sqlalchemy import or_
    book = syllabus_data.get_ncert_book(req.book_key)
    if not book:
        raise HTTPException(status_code=404, detail="Unknown NCERT book key.")
    book_disp = book["book"]
    grade = book.get("grade", "")
    bare = _cite_bare_book(book_disp)
    pages = _cite_book_pages(db, req.book_key, book)
    if not pages:
        return {"status": "skipped", "reason": "no stored chapter PDFs",
                "book": book_disp, "total": 0, "cited": 0, "already": 0, "weak": 0}
    mid = pages[len(pages) // 2]
    qs = (db.query(DBQuestion)
          .filter(DBQuestion.book == book_disp,
                  or_(DBQuestion.chapter.is_(None), DBQuestion.chapter == ""))
          .all())
    cited = already = weak = 0
    for q in qs:
        expl = q.explanation or ""
        if (_CITE_MARK in expl) and not req.overwrite:
            already += 1
            continue
        qtoks = _cite_tokens((q.text or "") + " " + _cite_correct_opt(q) + " " + (q.topic or ""))
        pg, sc = _cite_best_page(pages, qtoks)
        if pg is None:
            pg = mid
            weak += 1
        line = (f"{_CITE_MARK} {_CITE_DOT} {grade} {_CITE_DOT} {bare} {_CITE_DOT} "
                f"Ch.{pg['ch'] + 1} {pg['title']} {_CITE_DOT} p.{pg['folio']}")
        if req.overwrite and (_CITE_MARK in expl):
            expl = re.sub(r"\n*Source: NCERT.*$", "", expl, flags=re.S).rstrip()
        q.explanation = (expl.rstrip() + "\n\n" + line) if expl.strip() else line
        cited += 1
    db.commit()
    return {"status": "success", "book": book_disp, "grade": grade,
            "book_pages": len(pages), "total": len(qs),
            "cited": cited, "already": already, "weak": weak}

@app.get("/admin/cite-coverage", tags=["Admin"])
def admin_cite_coverage(admin: DBUser = Depends(require_admin), db: Session = Depends(get_db)):
    """Per-NCERT-book count of questions vs. how many carry a 'Source: NCERT' line."""
    from sqlalchemy import text as _t
    rows = db.execute(_t(
        "SELECT book, count(*) tot, "
        "count(*) FILTER (WHERE explanation LIKE '%Source: NCERT%') cited "
        "FROM questions WHERE book LIKE '%(Class%' GROUP BY book ORDER BY book")).fetchall()
    out = [{"book": r[0], "total": r[1], "cited": r[2]} for r in rows]
    return {"status": "success",
            "total": sum(r["total"] for r in out),
            "total_cited": sum(r["cited"] for r in out),
            "books": out}


@app.get("/admin/db-usage", tags=["Admin"])
def admin_db_usage(admin: DBUser = Depends(require_admin), db: Session = Depends(get_db)):
    """Exact database storage: total size + per-table breakdown (size incl. indexes
    & TOAST, plus an approximate row count). Read-only — for monitoring vs plan limits."""
    from sqlalchemy import text as _t
    tot = db.execute(_t(
        "SELECT pg_size_pretty(pg_database_size(current_database())), "
        "pg_database_size(current_database())")).fetchone()
    rows = db.execute(_t(
        "SELECT c.relname, "
        "pg_size_pretty(pg_total_relation_size(c.oid)) AS total_size, "
        "pg_total_relation_size(c.oid) AS bytes, "
        "c.reltuples::bigint AS est_rows "
        "FROM pg_class c JOIN pg_namespace n ON n.oid = c.relnamespace "
        "WHERE n.nspname = 'public' AND c.relkind = 'r' "
        "ORDER BY pg_total_relation_size(c.oid) DESC LIMIT 25")).fetchall()
    return {
        "status": "success",
        "database_size": tot[0],
        "database_bytes": int(tot[1]),
        "tables": [{"table": r[0], "size": r[1], "bytes": int(r[2]),
                    "approx_rows": int(r[3]) if r[3] is not None else None} for r in rows],
    }


# ── Concept exam-track classification (Prelims vs Mains/Interview) ────────────
def _ensure_exam_track_col(db):
    from sqlalchemy import text as _t
    db.execute(_t("ALTER TABLE concept_inventory ADD COLUMN IF NOT EXISTS exam_track VARCHAR"))
    db.commit()

@app.get("/admin/concepts/subtopics", tags=["Knowledge"])
def admin_concept_subtopics(admin: DBUser = Depends(require_admin), db: Session = Depends(get_db)):
    """Distinct (subject, subtopic) with concept counts + current exam_track — drives
    the Prelims/Mains/Interview classification."""
    from sqlalchemy import text as _t
    _ensure_exam_track_col(db)
    rows = db.execute(_t(
        "SELECT coalesce(subject,'(none)') s, coalesce(subtopic,'(none)') st, count(*) n, "
        "max(exam_track) trk FROM concept_inventory GROUP BY 1,2 ORDER BY 1, n DESC")).fetchall()
    tally = db.execute(_t(
        "SELECT coalesce(exam_track,'(untagged)') t, count(*) n FROM concept_inventory GROUP BY 1")).fetchall()
    return {"status": "success",
            "distinct_subtopics": len(rows),
            "track_tally": {r[0]: int(r[1]) for r in tally},
            "subtopics": [{"subject": r[0], "subtopic": r[1], "n": int(r[2]), "track": r[3]} for r in rows]}

class TagTracksRequest(BaseModel):
    reset_to: Optional[str] = None            # set ALL rows to this first (e.g. 'prelims')
    subject_defaults: Optional[dict] = None   # {subject: track} applied after reset
    rules: Optional[list] = None              # [{subject, subtopic, track}] per-subtopic overrides

@app.post("/admin/concepts/tag-tracks", tags=["Knowledge"])
def admin_tag_tracks(req: TagTracksRequest,
                     admin: DBUser = Depends(require_admin), db: Session = Depends(get_db)):
    """Write exam_track on concept_inventory. Order: reset_to -> subject_defaults ->
    per-subtopic rules (later steps override earlier). Tracks: prelims|mains|both|interview."""
    from sqlalchemy import text as _t
    _ensure_exam_track_col(db)
    # A track is any short lowercase word (e.g. prelims|mains|both|interview|csat).
    # Values are always bound as query parameters, so this is only a sanity gate.
    valid = lambda x: isinstance(x, str) and bool(re.match(r"^[a-z]{3,12}$", x))
    updated = 0
    if valid(req.reset_to):
        r = db.execute(_t("UPDATE concept_inventory SET exam_track=:t"), {"t": req.reset_to})
        updated += r.rowcount or 0
    for s, trk in (req.subject_defaults or {}).items():
        if valid(trk):
            r = db.execute(_t("UPDATE concept_inventory SET exam_track=:t WHERE subject=:s"),
                           {"t": trk, "s": s})
            updated += r.rowcount or 0
    for rule in (req.rules or []):
        trk = (rule or {}).get("track")
        if not valid(trk):
            continue
        s = rule.get("subject"); st = rule.get("subtopic")
        if st is None:
            r = db.execute(_t("UPDATE concept_inventory SET exam_track=:t WHERE subject=:s"),
                           {"t": trk, "s": s})
        else:
            r = db.execute(_t("UPDATE concept_inventory SET exam_track=:t "
                              "WHERE subject=:s AND subtopic=:st"), {"t": trk, "s": s, "st": st})
        updated += r.rowcount or 0
    db.commit()
    tally = db.execute(_t(
        "SELECT coalesce(exam_track,'(untagged)') t, count(*) n FROM concept_inventory GROUP BY 1")).fetchall()
    return {"status": "success", "rows_updated": updated,
            "track_tally": {r[0]: int(r[1]) for r in tally}}


@app.get("/admin/candidates", tags=["Admin"])
def admin_candidates(admin: DBUser = Depends(require_admin), db: Session = Depends(get_db)):
    rows = []
    for u in db.query(DBUser).filter(DBUser.email != SYSTEM_USER_EMAIL).all():
        atts = db.query(DBTestAttempt).filter(DBTestAttempt.user_id == u.id).all()
        pcts = [(a.score or 0) / a.mock_test.total_questions * 100
                for a in atts if a.mock_test and a.mock_test.total_questions]
        last = max([a.completed_at for a in atts if a.completed_at], default=None)
        rows.append({
            "id": u.id, "name": u.name or "", "email": u.email,
            "registered": u.created_at.isoformat() if u.created_at else None,
            "tests_taken": len(atts),
            "avg_pct": round(sum(pcts) / len(pcts), 1) if pcts else 0.0,
            "best_pct": round(max(pcts), 1) if pcts else 0.0,
            "last_active": last.isoformat() if last else None,
        })
    rows.sort(key=lambda r: r["last_active"] or "", reverse=True)
    return {"status": "success", "count": len(rows), "candidates": rows}

@app.get("/admin/candidates/{user_id}", tags=["Admin"])
def admin_candidate_detail(user_id: int, admin: DBUser = Depends(require_admin), db: Session = Depends(get_db)):
    u = db.query(DBUser).filter(DBUser.id == user_id).first()
    if not u:
        raise HTTPException(status_code=404, detail="Candidate not found")
    atts = (db.query(DBTestAttempt).filter(DBTestAttempt.user_id == u.id)
            .order_by(DBTestAttempt.completed_at.desc()).all())
    attempts = []
    for a in atts:
        m = _attempt_metrics(db, a)
        attempts.append({"attempt_id": a.id, "test": a.mock_test.title if a.mock_test else "",
                         "date": a.completed_at.isoformat() if a.completed_at else None, **m})
    rows = (db.query(DBQuestion.subject, DBAnswer.is_correct)
            .join(DBAnswer, DBAnswer.question_id == DBQuestion.id)
            .join(DBTestAttempt, DBAnswer.test_attempt_id == DBTestAttempt.id)
            .filter(DBTestAttempt.user_id == u.id).all())
    subj = {}
    for s, ok in rows:
        if not s or s == "General":
            continue
        e = subj.setdefault(s, {"attempted": 0, "correct": 0})
        e["attempted"] += 1
        if ok:
            e["correct"] += 1
    by_subject = [{"subject": s, "attempted": e["attempted"], "correct": e["correct"],
                   "accuracy": round(e["correct"] / e["attempted"] * 100, 1) if e["attempted"] else 0.0}
                  for s, e in sorted(subj.items())]
    return {
        "status": "success",
        "candidate": {"id": u.id, "name": u.name or "", "email": u.email,
                      "registered": u.created_at.isoformat() if u.created_at else None,
                      "tests_taken": len(atts)},
        "attempts": attempts, "by_subject": by_subject,
    }

@app.get("/admin/activity", tags=["Admin"])
def admin_activity(admin: DBUser = Depends(require_admin), db: Session = Depends(get_db)):
    events = []
    for u in db.query(DBUser).filter(DBUser.email != SYSTEM_USER_EMAIL).all():
        if u.created_at:
            events.append({"type": "register", "time": u.created_at.isoformat(),
                           "name": u.name or u.email, "email": u.email, "detail": "registered"})
    for a in (db.query(DBTestAttempt).order_by(DBTestAttempt.completed_at.desc()).limit(100).all()):
        t = a.mock_test.total_questions or 0 if a.mock_test else 0
        pct = round((a.score or 0) / t * 100, 1) if t else 0.0
        u = a.user
        events.append({"type": "attempt", "time": a.completed_at.isoformat() if a.completed_at else None,
                       "name": (u.name or u.email) if u else "", "email": u.email if u else "",
                       "detail": "scored " + str(pct) + "% on " + (a.mock_test.title if a.mock_test else "")})
    events = [e for e in events if e["time"]]
    events.sort(key=lambda e: e["time"], reverse=True)
    return {"status": "success", "activity": events[:60]}

class AdminEmailReq(BaseModel):
    email: str

@app.get("/admin/admins", tags=["Admin"])
def list_admins(admin: DBUser = Depends(require_admin), db: Session = Depends(get_db)):
    config_set = set(ADMIN_EMAILS)
    out = [{"email": e, "source": "built-in", "removable": False} for e in sorted(config_set)]
    for a in db.query(DBAdminEmail).order_by(DBAdminEmail.email).all():
        if (a.email or "").lower() in config_set:
            continue
        out.append({"email": a.email, "source": "added", "removable": True,
                    "added_by": a.added_by,
                    "created_at": a.created_at.isoformat() if a.created_at else None})
    return {"status": "success", "admins": out}

@app.post("/admin/admins", tags=["Admin"])
def add_admin(req: AdminEmailReq, admin: DBUser = Depends(require_admin), db: Session = Depends(get_db)):
    email = (req.email or "").strip().lower()
    if "@" not in email or "." not in email.split("@")[-1]:
        raise HTTPException(status_code=400, detail="Enter a valid email address")
    if email in set(ADMIN_EMAILS):
        raise HTTPException(status_code=400, detail="That email is already a built-in admin")
    if db.query(DBAdminEmail).filter(DBAdminEmail.email == email).first():
        raise HTTPException(status_code=400, detail="That email is already an admin")
    db.add(DBAdminEmail(email=email, added_by=admin.email))
    db.commit()
    return {"status": "success", "message": email + " added as admin",
            "note": "They get admin access once they register/log in with this email."}

@app.delete("/admin/admins", tags=["Admin"])
def remove_admin(email: str, admin: DBUser = Depends(require_admin), db: Session = Depends(get_db)):
    email = (email or "").strip().lower()
    if email in set(ADMIN_EMAILS):
        raise HTTPException(status_code=400, detail="Built-in admins can't be removed here")
    a = db.query(DBAdminEmail).filter(DBAdminEmail.email == email).first()
    if not a:
        raise HTTPException(status_code=404, detail="Admin email not found")
    db.delete(a)
    db.commit()
    return {"status": "success", "message": email + " removed from admins"}

# ── AI ────────────────────────────────────────────────────────────────────────
@app.post("/ai/generate-questions/", tags=["AI Features"])
def ai_generate_questions(request: AIGenerateRequest, current_user: DBUser = Depends(get_current_user)):
    try:
        result = generate_questions(request.subject, request.topic, request.num_questions)
        return {"status": "success", "generated_questions": result}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI Error: {str(e)}")

@app.post("/ai/explain/", tags=["AI Features"])
def ai_explain(request: AIExplainRequest, current_user: DBUser = Depends(get_current_user)):
    try:
        result = explain_concept(request.topic, request.context)
        return {"status": "success", "explanation": result}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI Error: {str(e)}")


class MissionLearnRequest(BaseModel):
    target: str
    topic: Optional[str] = ""
    exam_label: Optional[str] = None
    mastery: Optional[float] = None
    retention: Optional[float] = None
    consistency: Optional[float] = None
    reading_speed: Optional[float] = None
    attempts: Optional[int] = None
    style: Optional[str] = None
    attempt: Optional[int] = 0        # re-teach counter (0 = first lesson)


# Map a mission target (often a concept, e.g. "Fundamental Rights") to the generation
# subject the concept library is keyed by, so grounding actually matches. Fixes the
# grounded:false gap for concept-level targets.
_TARGET_SUBJECT = {
    "fundamental rights": "Indian Polity", "directive principles": "Indian Polity",
    "fundamental duties": "Indian Polity", "preamble": "Indian Polity",
    "federalism": "Indian Polity", "parliament": "Indian Polity",
    "judiciary": "Indian Polity", "polity": "Indian Polity", "constitution": "Indian Polity",
    "president": "Indian Polity", "governor": "Indian Polity", "amendment": "Indian Polity",
    "inflation": "Indian Economy", "deflation": "Indian Economy", "gdp": "Indian Economy",
    "fiscal policy": "Indian Economy", "monetary policy": "Indian Economy", "budget": "Indian Economy",
}


def _resolve_target(target):
    """Return (generation_subject, topic) for grounding lookups. Prefers an exact
    subject the library knows; else maps a known concept to its subject."""
    t = (target or "").strip()
    tl = t.lower()
    for k in _CONCEPT_SUBJECT_MAP:
        kl = k.lower()
        if kl in tl or tl in kl:
            return k, ("" if kl == tl else t)
    for k, subj in _TARGET_SUBJECT.items():
        if k in tl:
            return subj, t
    return t, ""


def _csat_ground(target, n=4):
    """Ground CSAT/aptitude lessons in the app's REAL CSAT practice bank (exam-pattern
    questions with verified answers + explanations) instead of generic LLM filler."""
    tl = (target or "").lower()
    if not any(w in tl for w in ("csat", "reasoning", "comprehension", "aptitude", "quant", "decision")):
        return "", None
    words = [w for w in tl.replace("csat", "").split() if len(w) > 3]
    area = None
    for a in CSAT_AREAS:
        nm = (a.get("name", "") + " " + a.get("code", "")).lower()
        if any(w in nm for w in words):
            area = a
            break
    if not area and CSAT_AREAS:
        area = next((a for a in CSAT_AREAS if "reason" in a.get("name", "").lower()), CSAT_AREAS[0])
    if not area:
        return "", None
    lines = []
    for q in area.get("questions", [])[:n]:
        opts = " | ".join(x for x in [q.get("option_a"), q.get("option_b"),
                                      q.get("option_c"), q.get("option_d")] if x)
        lines.append(f"Q: {q.get('text','')}\nOptions: {opts}\nCorrect: {q.get('correct_answer','')} — "
                     f"{q.get('explanation','')}")
    if not lines:
        return "", None
    return ("Real CSAT exam-pattern questions from the platform bank (verified answers) — "
            "teach the SOLVING TECHNIQUE using these, and how to eliminate wrong options:\n"
            + "\n\n".join(lines)), area.get("name")


# Map a lesson target to the real-PYQ 'type' bucket in csat_pyq.json.
_PYQ_TYPE_MAP = [
    (("comprehen", "passage", "reading", "rc "), "comprehension"),
    (("data interpret", "data-interpret", "graph", "chart", "table"), "data-interpretation"),
    (("sufficien", "statement"), "data-sufficiency"),
    (("decision", "situation", "ethical"), "decision-making"),
    (("quant", "math", "arithmetic", "percentage", "ratio", "speed", "number", "average"), "quant"),
    (("reason", "logic", "direction", "blood", "series", "coding", "syllog", "seat", "arrange", "rank"), "reasoning"),
]
_PYQ_SUBFAMS = ["direction", "blood", "series", "coding", "syllog", "seat", "arrange",
                "rank", "calendar", "clock", "cube", "venn"]


def _pyq_ground(target, attempt=0):
    """Pull ONE real UPSC CSAT PYQ (verified answer) to anchor the worked example.
    The moat: teach the technique on the ACTUAL exam UPSC set, not a coaching toy.
    Returns (pyq_block, label, question_dict) or ("", None, None)."""
    tl = (target or "").lower()
    if not any(w in tl for w in ("csat", "reason", "logic", "aptitude", "quant",
                                 "decision", "comprehen", "interpret", "sufficien",
                                 "direction", "blood", "series", "coding", "syllog")):
        return "", None, None
    qtype = "reasoning"
    for kws, t in _PYQ_TYPE_MAP:
        if any(k in tl for k in kws):
            qtype = t
            break
    pool = []
    for p in CSAT_PYQ_BANK.get("papers", []):
        yr = p.get("year")
        # Passage lookup for comprehension questions (passage text lives on the paper).
        pmap = {str(x.get("id")): (x.get("text") or "") for x in (p.get("passages") or [])}
        for q in p.get("questions", []):
            if q.get("type") == qtype and q.get("text") and q.get("correct_answer"):
                pool.append((yr, q, pmap))
    if not pool:
        return "", None, None
    # Narrow to a single sub-family when the target names one (teach ONE thing deeply).
    sub = next((s for s in _PYQ_SUBFAMS if s in tl), None)
    if sub:
        narrowed = [x for x in pool if sub in (x[1].get("text", "").lower())]
        if narrowed:
            pool = narrowed
    # Prefer a teachable length (not a one-liner, not a monster); deterministic pick
    # that rotates gently on re-teach so a second attempt gets a fresh question.
    pool.sort(key=lambda x: len(x[1].get("text", "")))
    mid = pool[len(pool) // 4: max(1, (len(pool) * 3) // 4)] or pool
    year, qd, pmap = mid[(len(tl) + int(attempt or 0)) % len(mid)]
    opts = "  ".join(f"({L}) {qd.get('option_' + L.lower())}"
                     for L in ["A", "B", "C", "D"] if qd.get("option_" + L.lower()))
    # For comprehension, include the passage so the mentor can teach reading + elimination.
    passage = pmap.get(str(qd.get("passage_id"))) if qd.get("passage_id") else ""
    pre = (f"PASSAGE:\n{passage.strip()}\n\n" if passage else "")
    block = (f"REAL UPSC CSAT PYQ ({year or 'recent'}): {pre}QUESTION: {qd.get('text', '').strip()}\n"
             f"Options: {opts}\nVerified correct answer: {qd.get('correct_answer', '')}")
    return block, (f"UPSC CSAT PYQ {year}" if year else "UPSC CSAT PYQ"), qd


def _teaching_aggregate_priors(db, min_n=3):
    """Cross-learner priors: the best strategy per barrier across ALL learners —
    ranked by EFFICIENCY (win-rate per minute), so the engine prefers the FASTEST
    reliable teaching, not merely one that works (time-to-mastery, §11)."""
    priors = {}
    try:
        rows = db.query(DBTeachingEvent.barrier, DBTeachingEvent.strategy,
                        DBTeachingEvent.passed, DBTeachingEvent.seconds).limit(8000).all()
        agg = {}
        for b, s, p, sec in rows:
            if not b or not s:
                continue
            a = agg.setdefault((b, s), [0, 0, 0, 0])   # wins, total, sec_sum, sec_n
            a[1] += 1
            if p:
                a[0] += 1
            if sec and sec > 0:
                a[2] += sec
                a[3] += 1
        by_barrier = {}
        for (b, s), (w, tot, ssum, sn) in agg.items():
            if tot >= min_n:
                wr = w / tot
                avg_min = (ssum / sn / 60.0) if sn else 6.0   # assume ~6 min if unknown
                eff = wr / max(0.5, avg_min)                  # mastery per minute
                if b not in by_barrier or eff > by_barrier[b][1]:
                    by_barrier[b] = (s, eff)
        priors = {b: sv[0] for b, sv in by_barrier.items()}
    except Exception:
        priors = {}
    return priors


def _learner_reason_counts(db, user_id, limit=60):
    """Best-effort: how this learner tends to miss (conceptual/factual/careless/
    misread/guess) — the raw signal the Teaching Engine reads for behaviour barriers."""
    counts = {}
    try:
        rows = (db.query(DBAnswer.wrong_reason)
                .join(DBTestAttempt, DBAnswer.test_attempt_id == DBTestAttempt.id)
                .filter(DBTestAttempt.user_id == user_id,
                        DBAnswer.is_correct == False,
                        DBAnswer.wrong_reason.isnot(None))
                .order_by(DBAnswer.id.desc()).limit(limit).all())
        for (r,) in rows:
            r = (r or "").lower().strip()
            if r:
                counts[r] = counts.get(r, 0) + 1
    except Exception:
        return {}
    return counts


@app.post("/me/mission/learn", tags=["Planner"])
def mission_learn(request: MissionLearnRequest, db: Session = Depends(get_db),
                  current_user: DBUser = Depends(get_current_user)):
    """The Learn phase — the TEACHING ENGINE (see TEACHING_ENGINE.md) in action.
    First diagnose WHY this learner is blocked and pick a teaching strategy; then
    render that strategy over content GROUNDED in the verified concept library
    (Content ≠ Pedagogy — the engine chooses how, the library supplies the facts)."""
    target = (request.target or "").strip() or "today's topic"
    topic = (request.topic or "").strip()

    # ── Content: ground in the verified concept library, then book passages ──
    g_subject, g_topic = _resolve_target(target)
    g_topic = g_topic or topic
    concept_ctx = ""
    source_label = "your concept library"
    try:
        concept_ctx = _concept_context(db, g_subject, g_topic) or ""
        if len(concept_ctx) < 40:
            book_ctx = _retrieve_context(db, g_subject, g_topic) or ""
            concept_ctx = (concept_ctx + "\n" + book_ctx).strip()
    except Exception:
        concept_ctx = ""
    if len(concept_ctx) < 40:                          # CSAT/aptitude → real CSAT bank
        csat_ctx, csat_area = _csat_ground(target)
        if csat_ctx:
            concept_ctx = csat_ctx
            source_label = "the CSAT practice bank" + (f" · {csat_area}" if csat_area else "")
    grounded = len(concept_ctx) >= 40
    # Real UPSC PYQ to anchor the worked example (the moat) — CSAT/aptitude targets.
    pyq_block, pyq_label, _pyq_q = _pyq_ground(target, attempt=(request.attempt or 0))

    # ── Pedagogy: diagnose the barrier + choose the strategy for THIS learner ──
    # Teaching-Strategy memory (Evidence Engine §12): what has worked for THIS learner.
    strat_hist = []
    try:
        for ev in (db.query(DBTeachingEvent)
                   .filter(DBTeachingEvent.user_id == current_user.id)
                   .order_by(DBTeachingEvent.id.desc()).limit(40).all()):
            strat_hist.append({"barrier": ev.barrier, "strategy": ev.strategy,
                               "passed": bool(ev.passed), "seconds": (ev.seconds or 0)})
        strat_hist.reverse()
    except Exception:
        strat_hist = []
    signals = {
        "target": target, "exam_label": request.exam_label,
        "mastery": request.mastery, "retention": request.retention,
        "consistency": request.consistency, "reading_speed": request.reading_speed,
        "attempts": request.attempts or 0, "style": request.style,
        "attempt": request.attempt or 0, "strategy_history": strat_hist,
        "aggregate_priors": _teaching_aggregate_priors(db),
        "reasons": _learner_reason_counts(db, current_user.id), "grounded": grounded,
    }
    try:
        tplan = teaching.plan(signals)
    except Exception:
        tplan = None

    # ── Render pedagogy × content into the lesson ──
    if tplan:
        strat = tplan["strategy"]
        recipe = " → ".join(strat.get("recipe", []))
        objectives = "; ".join(tplan["objectives"]["objectives"][:4])
        barrier = tplan["barrier"]
        mm = tplan.get("mental_model")
        mm_line = (f"FIRST install this mental model (the structure, before any detail): "
                   f"{' → '.join(mm)}. Then hang the detail on it.\n") if mm else ""
        fix = tplan.get("model_fix")
        fix_line = (f"REPAIR THE MODEL FIRST. The learner very likely believes: \"{fix['misconception']}\" "
                    f"That is the root of the confusion. Before any detail, explicitly replace it with the "
                    f"correct model: {fix['correction']} Name the wrong idea, flip it, then teach on the "
                    f"corrected model.\n") if fix else ""
        cog = tplan.get("cognitive_state") or {}
        cog_line = (f"The learner's current state is '{cog.get('state')}' — {cog.get('action')}\n"
                    if cog.get("action") else "")
        dep = tplan.get("depth") or {}
        dep_line = (f"TARGET DEPTH: by the end the learner must be able to {dep.get('label')} — "
                    f"teach and phrase everything to reach that depth, not just 'understood'.\n"
                    if dep.get("label") else "")
        bmix = tplan.get("build_mix") or {}
        mix_line = ("Internally weight the lesson toward: "
                    + ", ".join(f"{k.replace('_',' ')} {v}%" for k, v in bmix.items() if v) + ".\n") if bmix else ""
        # ── Mentor 'gold' the ENGINE supplies (not LLM-invented) ──
        et = tplan.get("exam_tricks")
        trick_line = (f"USE THIS EXAM TRICK — do not invent your own. Heuristic (the 'aha'): "
                      f"\"{et['heuristic']}\" Trick: \"{et['trick']}\" Teach these as the memorable rule "
                      f"and the exam move.\n") if et else ""
        # Personalise the opener from how THIS learner actually loses marks.
        _REASON_PHRASE = {
            "conceptual": "you understand the idea loosely, but the exam's exact wording trips you",
            "factual": "the fact is there, but it slips at the moment you need it",
            "careless": "you know it, but speed or a rushed read costs you the mark",
            "misread": "you solve a slightly different question than the one actually asked",
            "guess": "when unsure you guess, instead of eliminating your way to the answer",
            "silly": "you know it, but speed or a rushed read costs you the mark",
        }
        reasons = (signals.get("reasons") or {})
        top_reason = max(reasons, key=reasons.get) if reasons else ""
        lose_phrase = ""
        for _k, _phrase in _REASON_PHRASE.items():
            if _k in top_reason:
                lose_phrase = _phrase
                break
        lose_line = (f"This learner's real leak: {lose_phrase}. Open by naming it warmly, in one line, "
                     f"so they think 'yes — that's exactly me.'\n") if lose_phrase else ""
        # Mastery band → adaptive depth of the SAME topic (owner #10): different lesson,
        # not the same lesson for everyone.
        _m = signals.get("mastery")
        try:
            _m = float(_m)
        except (TypeError, ValueError):
            _m = 0.3
        if _m < 0.35:
            band_line = ("LEVEL: near-beginner. Teach the ONE core habit/idea plainly; do not overload. "
                         "Goal is a first clean win.\n")
        elif _m < 0.7:
            band_line = ("LEVEL: intermediate. They know the basics — teach speed, elimination and the "
                         "trap that still costs them; skip beginner definitions.\n")
        else:
            band_line = ("LEVEL: advanced. Teach the subtle assumption UPSC hides and the fastest path; "
                         "assume fluency with basics.\n")
        # Real UPSC PYQ for the worked example (the moat) — passed in from _pyq_ground.
        pyq_line = (f"USE THIS REAL UPSC PYQ for the 'Watch Me Think' worked example — narrate solving "
                    f"it to its verified answer; do NOT invent a toy question:\n{pyq_block}\n"
                    if pyq_block else
                    "For 'Watch Me Think', use one realistic UPSC/CSAT-style question (or one grounded "
                    "above) — never an artificial toy.\n")
        pedagogy = (
            f"You are this learner's personal UPSC MENTOR — teacher + guide + friend in one. Not a "
            f"textbook, not a lecturer. You are mentoring them on '{target}'. Speak directly TO them "
            "('you'), warm and human, like a sharp senior who cracked this exam and wants them to.\n"
            f"What's holding them back: {barrier['label']} — {', '.join(barrier.get('why', []))}.\n"
            + fix_line + cog_line + mm_line + trick_line + lose_line + band_line + dep_line + mix_line
            + pyq_line +
            "HARD RULES (this is a Today's Mission, ~8 min — laser-focused, not a chapter):\n"
            "• Whole lesson must be readable in UNDER 4 minutes. Ruthless brevity — if a line doesn't "
            "change how they SOLVE, cut it. No syllabus restating, no padding, no 'in conclusion'.\n"
            "• Teach exactly ONE sub-family / ONE habit today (e.g. only direction-sense, or only "
            "'classify before solving') — depth over breadth. Mention nothing you won't teach.\n"
            "• Teach how UPSC ACTUALLY asks this. UPSC rarely asks hard logic — it hides simple logic "
            "inside long wording. Teach that behaviour, not generic aptitude theory.\n"
            "• Every logical step correct. If something 'cannot be determined', say exactly that — never "
            "fake certainty. A wrong worked example is worse than none.\n\n"
            "WRITE THE LESSON IN THIS EXACT MENTOR STRUCTURE (bold headings, keep each tiny):\n"
            "**Did you know?** — one surprising, true, motivating hook about how UPSC sets this (1 line).\n"
            "**Why you lose marks here** — the personal diagnosis above, in 1–2 lines. Make them nod.\n"
            "**Today's one goal** — state the SINGLE habit/skill you'll fix today. Only one.\n"
            "**The big insight** — the one sentence that changes how they think, plus ONE vivid, "
            "specific analogy that sticks (e.g. 'reasoning is detective work — you don't guess, you "
            "eliminate suspects until one remains'). No limp analogies like 'puzzle pieces'.\n"
            "**See it** — a tiny ASCII diagram / flowchart / table (e.g. Question → Identify type → "
            "Strip the wording → Find pattern → Eliminate → Answer). The brain remembers pictures.\n"
            "**Watch me think** — solve the REAL PYQ above by THINKING ALOUD like a topper, not "
            "presenting. Model the inner voice: 'Stop — don't calculate everything. First check if the "
            "jump is constant… no… second differences… no… now try squares… ah.' Show the METHOD.\n"
            "**Where UPSC traps you** — name the exact trap here + one timing rule ('never spend >90s "
            "on this — mark it and move', or 'if you see X, solve backward from the options').\n"
            "**Remember forever** — one tiny memory anchor: a mnemonic, a one-line rule, or a 3-line "
            "diagram they'll recall in the hall (e.g. 'Right turn = clockwise, Left turn = anti-clockwise').\n"
            "**Now you try** — pose ONE similar question and STOP. Do not solve it. (The app checks them "
            "next.)\n"
            "**Today's win** — end SPECIFIC and actionable, tied to the one goal: 'Today isn't about "
            "mastering all reasoning — it's about classifying every question before you solve. Do that "
            "consistently and your accuracy jumps.' Then hand off warmly to the practice set.")
    else:
        pedagogy = ("For a UPSC aspirant. Give the core concept, 2 short worked examples, "
                    "the 2-3 most common mistakes, and one quick approach tip. Concise.")
    if pyq_block:
        # CSAT/aptitude: the constraint is TECHNIQUE + a real PYQ, not fact-fidelity.
        # Give the concept bank as reference but let the mentor teach the method freely.
        ref = ("Reference material from the platform (use for accuracy, don't just recite):\n"
               + concept_ctx + "\n\n") if grounded else ""
        ctx = ref + pedagogy
    elif grounded:
        ctx = ("Base this lesson STRICTLY on these verified concepts and facts from the "
               "student's own library — do not introduce outside facts:\n" + concept_ctx +
               "\n\n" + pedagogy)
    else:
        ctx = pedagogy
    try:
        explanation = explain_concept(target, ctx)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI Error: {str(e)}")

    # Concept names to show only apply to ConceptInventory grounding — the CSAT bank
    # context is Q/Options/answers, not concept names, so don't surface it as "concepts".
    concepts = []
    if source_label.startswith("your concept library"):
        for ln in concept_ctx.split("\n"):
            ln = ln.strip().lstrip("-").strip()
            if ln and not ln.lower().startswith(("base ", "teach ", "use ", "diagnosed ", "you are ")):
                concepts.append(ln.split(":")[0].split("(")[0].strip())
            if len(concepts) >= 6:
                break
    # Provenance: a real PYQ anchor is itself grounding worth showing, even if the
    # concept bank was thin.
    if pyq_block:
        src = ((source_label + " · " if grounded else "") + (pyq_label or "real UPSC PYQ"))
        out_grounded = True
    else:
        src = source_label if grounded else "general knowledge"
        out_grounded = grounded
    return {"status": "success", "explanation": explanation, "grounded": out_grounded,
            "source": src, "concepts": concepts, "plan": tplan}


class MissionCheckRequest(BaseModel):
    target: str
    topic: Optional[str] = ""
    difficulty: Optional[str] = "easy"


@app.post("/me/mission/check", tags=["Planner"])
def mission_check(request: MissionCheckRequest, db: Session = Depends(get_db),
                  current_user: DBUser = Depends(get_current_user)):
    """The immediate understanding check after a lesson — a few grounded questions on
    exactly what was taught. This is how the Teaching Engine learns whether the
    teaching worked *right now* (not in three days). Fail → re-teach differently."""
    target = (request.target or "").strip() or "today's topic"
    topic = (request.topic or "").strip()
    g_subject, g_topic = _resolve_target(target)
    g_topic = g_topic or topic
    ctx = ""
    try:
        ctx = _concept_context(db, g_subject, g_topic) or ""
        if len(ctx) < 40:
            ctx = (ctx + "\n" + (_retrieve_context(db, g_subject, g_topic) or "")).strip()
    except Exception:
        ctx = ""
    try:
        qs = generate_verified_questions(subject=g_subject, topic=g_topic, num_questions=3,
                                         difficulty=(request.difficulty or "easy"),
                                         question_type="all", source_context=ctx) or []
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Could not build the check: {str(e)}")
    qs = qs[:3]
    if not qs:
        raise HTTPException(status_code=502, detail="Could not build the check. Please try again.")
    # Persist the check as a small graded set so, once taken, it can count toward
    # mastery in the Digital Twin (the same pipeline as any verified test). Titled so
    # it can be excluded from the "mocks" metric — it's a check, not a full mock.
    mock_id, dbqs = None, []
    try:
        mt = DBMockTest(title=f"🔎 Mission check — {target}"[:120],
                        description="Teaching Engine understanding check", subject=g_subject,
                        total_questions=len(qs), duration_minutes=max(2, len(qs)), user_id=current_user.id)
        db.add(mt); db.commit(); db.refresh(mt); mock_id = mt.id
        for q in qs:
            cans = (q.get("correct_answer") or "A").strip().upper()
            db.add(DBQuestion(
                text=q.get("text"), option_a=q.get("option_a"), option_b=q.get("option_b"),
                option_c=q.get("option_c"), option_d=q.get("option_d"),
                correct_answer=cans[:1] if cans[:1] in ("A", "B", "C", "D") else cans,
                explanation=q.get("explanation", ""), subject=g_subject, topic=g_topic or None,
                difficulty=(request.difficulty or "easy"),
                question_type=q.get("question_type", "direct"), mock_test_id=mt.id))
        db.commit()
        dbqs = db.query(DBQuestion).filter(DBQuestion.mock_test_id == mt.id).order_by(DBQuestion.id.asc()).all()
    except Exception:
        db.rollback(); mock_id, dbqs = None, []
    out = []
    for i, q in enumerate(qs):
        ans = (q.get("correct_answer") or "").strip().upper()
        out.append({"id": (dbqs[i].id if i < len(dbqs) else None), "text": q.get("text"),
                    "options": {"A": q.get("option_a"), "B": q.get("option_b"),
                                "C": q.get("option_c"), "D": q.get("option_d")},
                    "answer": ans[:1] if ans[:1] in ("A", "B", "C", "D") else ans,
                    "explanation": q.get("explanation", "")})
    return {"status": "success", "questions": out, "grounded": len(ctx) >= 40,
            "mock_test_id": mock_id}


class MissionCheckResult(BaseModel):
    target: str
    barrier: Optional[str] = None
    layer: Optional[str] = None
    strategy: Optional[str] = None
    attempt: Optional[int] = 0
    correct: Optional[int] = 0
    total: Optional[int] = 0
    passed: Optional[bool] = False
    stage: Optional[str] = None
    seconds: Optional[int] = 0            # lesson→check duration (time-to-mastery)
    mock_test_id: Optional[int] = None
    results: Optional[list] = None        # [{question_id, selected}]


@app.post("/me/mission/check/record", tags=["Planner"])
def mission_check_record(request: MissionCheckResult, db: Session = Depends(get_db),
                         current_user: DBUser = Depends(get_current_user)):
    """Record a taught-then-checked outcome. Two writes:
    (1) the Teaching Engine's memory (§12) — which strategy worked; and
    (2) the check as a small GRADED ATTEMPT, so it flows through the normal pipeline
        (_gather_answers → build_knowledge_map) and the learner's mastery in the
        Digital Twin actually moves when the teaching lands. Grading is server-side
        authoritative (never trusts the client)."""
    try:
        db.add(DBTeachingEvent(
            user_id=current_user.id, target=(request.target or "")[:200],
            barrier=request.barrier, layer=request.layer, strategy=request.strategy,
            attempt=int(request.attempt or 0), correct=int(request.correct or 0),
            total=int(request.total or 0), passed=bool(request.passed), stage=request.stage,
            seconds=max(0, min(int(request.seconds or 0), 7200)),
        ))
        db.commit()
    except Exception:
        db.rollback()

    # (2) Persist the check as a graded attempt → counts toward mastery (the Twin).
    try:
        if request.mock_test_id and request.results:
            answer_objs, score = [], 0
            for r in (request.results or []):
                try:
                    qid = int(r.get("question_id"))
                except Exception:
                    continue
                sel = (str(r.get("selected") or "").strip().upper()[:1])
                dq = (db.query(DBQuestion)
                      .filter(DBQuestion.id == qid, DBQuestion.mock_test_id == request.mock_test_id).first())
                if not dq:
                    continue
                is_c = bool(dq.correct_answer and dq.correct_answer.strip().upper() == sel)
                if is_c:
                    score += 1
                answer_objs.append(DBAnswer(selected_option=(sel or "?"), is_correct=is_c, question_id=qid))
            if answer_objs:
                pct = round(100 * score / len(answer_objs))
                db.add(DBTestAttempt(score=pct, time_taken_seconds=0, user_id=current_user.id,
                                     mock_test_id=request.mock_test_id, answers=answer_objs))
                db.commit()
    except Exception:
        db.rollback()
    return {"status": "success"}


@app.get("/me/teaching/insights", tags=["Planner"])
def teaching_insights(db: Session = Depends(get_db), current_user: DBUser = Depends(get_current_user)):
    """What the Teaching Engine has learned about THIS learner (Evidence Engine §12,
    surfaced): which strategies work for them, how often teaching lands, and the
    barriers they hit most. Makes the engine's memory visible and mentor-like."""
    try:
        evs = (db.query(DBTeachingEvent)
               .filter(DBTeachingEvent.user_id == current_user.id)
               .order_by(DBTeachingEvent.id.desc()).limit(300).all())
    except Exception:
        evs = []
    total = len(evs)
    passed = sum(1 for e in evs if e.passed)
    concepts = len({(e.target or "").lower() for e in evs if e.target})
    strat, barr = {}, {}
    for e in evs:
        if e.strategy:
            s = strat.setdefault(e.strategy, {"wins": 0, "total": 0})
            s["total"] += 1
            if e.passed:
                s["wins"] += 1
        if e.barrier and not e.passed:
            barr[e.barrier] = barr.get(e.barrier, 0) + 1
    best = sorted(({"strategy": k, "wins": v["wins"], "total": v["total"]}
                   for k, v in strat.items() if v["wins"] > 0),
                  key=lambda x: (-x["wins"], -x["wins"] / max(1, x["total"])))[:3]
    barriers = sorted(({"barrier": k, "n": v} for k, v in barr.items()), key=lambda x: -x["n"])[:3]
    return {"status": "success", "events": total, "concepts": concepts,
            "pass_rate": (round(100 * passed / total) if total else None),
            "best_strategies": best, "barriers": barriers,
            "has_enough": total >= 3}


@app.post("/ai/analyze/{attempt_id}/", tags=["AI Features"])
def ai_analyze_performance(attempt_id: int, db: Session = Depends(get_db), current_user: DBUser = Depends(get_current_user)):
    """AI Mentor Report: compute exact performance stats (net score with UPSC negative
    marking, attempt/accuracy split, subject-wise accuracy) and let the AI interpret them."""
    db_attempt = db.query(DBTestAttempt).filter(DBTestAttempt.id == attempt_id, DBTestAttempt.user_id == current_user.id).first()
    if not db_attempt:
        raise HTTPException(status_code=404, detail="Attempt not found")

    questions = db.query(DBQuestion).filter(DBQuestion.mock_test_id == db_attempt.mock_test_id).all()
    ans_map = {a.question_id: a for a in db.query(DBAnswer).filter(DBAnswer.test_attempt_id == attempt_id).all()}

    total = len(questions)
    attempted = sum(1 for q in questions if q.id in ans_map)
    correct = sum(1 for q in questions if q.id in ans_map and ans_map[q.id].is_correct)
    wrong = attempted - correct
    blank = total - attempted
    accuracy = round(correct / attempted * 100, 1) if attempted else 0.0

    # UPSC Prelims marking: 2 marks per question, -1/3 of 2 (~0.66) per wrong answer.
    MARKS_PER_Q, NEG = 2.0, 2.0 / 3.0
    max_marks = round(total * MARKS_PER_Q, 2)
    net_marks = round(correct * MARKS_PER_Q - wrong * NEG, 2)
    net_pct = round(net_marks / max_marks * 100, 1) if max_marks else 0.0

    def _correct_text(q):
        return {"A": q.option_a, "B": q.option_b, "C": q.option_c, "D": q.option_d}.get(q.correct_answer, "")

    # Subject-wise breakdown (skip untagged questions so they don't form a fake subject)
    subj = {}
    this_test_missed = []  # detailed misses on THIS test, for topic extraction
    for q in questions:
        a = ans_map.get(q.id)
        if not (a and a.is_correct):  # wrong or blank
            this_test_missed.append({
                "subject": q.subject or "—", "text": q.text,
                "correct": _correct_text(q), "exp": (q.explanation or "")[:200],
            })
        s = q.subject
        if not s or s == "General":
            continue
        e = subj.setdefault(s, {"total": 0, "correct": 0, "attempted": 0})
        e["total"] += 1
        if a:
            e["attempted"] += 1
            if a.is_correct:
                e["correct"] += 1
    by_subject = []
    for s, e in sorted(subj.items()):
        acc = round(e["correct"] / e["total"] * 100, 1) if e["total"] else 0.0
        by_subject.append({"subject": s, "total": e["total"], "correct": e["correct"],
                           "attempted": e["attempted"], "accuracy": acc})

    # ── Candidate-wide history (across ALL of this user's attempts) ──────────────
    all_attempts = (db.query(DBTestAttempt)
                    .filter(DBTestAttempt.user_id == current_user.id)
                    .order_by(DBTestAttempt.completed_at).all())
    attempt_ids = [a.id for a in all_attempts]

    # Cumulative subject accuracy across every answer the candidate has ever given.
    overall_subj = {}
    if attempt_ids:
        rows = (db.query(DBQuestion.subject, DBAnswer.is_correct)
                .join(DBAnswer, DBAnswer.question_id == DBQuestion.id)
                .filter(DBAnswer.test_attempt_id.in_(attempt_ids)).all())
        for s, ok in rows:
            if not s or s == "General":
                continue  # ignore untagged questions in subject analytics
            e = overall_subj.setdefault(s, {"attempted": 0, "correct": 0})
            e["attempted"] += 1
            if ok:
                e["correct"] += 1
    overall_by_subject = []
    for s, e in sorted(overall_subj.items()):
        acc = round(e["correct"] / e["attempted"] * 100, 1) if e["attempted"] else 0.0
        overall_by_subject.append({"subject": s, "attempted": e["attempted"], "correct": e["correct"], "accuracy": acc})

    # Recurring misses across ALL the candidate's tests (for topic-pattern detection).
    history_missed = []
    if attempt_ids:
        for s, qtext in (db.query(DBQuestion.subject, DBQuestion.text)
                         .join(DBAnswer, DBAnswer.question_id == DBQuestion.id)
                         .filter(DBAnswer.test_attempt_id.in_(attempt_ids),
                                 DBAnswer.is_correct == False).limit(60).all()):  # noqa: E712
            if s and s != "General":
                history_missed.append((s, qtext))

    # Per-attempt trend (net % and accuracy over time).
    ans_counts = {}
    if attempt_ids:
        from sqlalchemy import func as sa_func
        for aid, cnt in (db.query(DBAnswer.test_attempt_id, sa_func.count(DBAnswer.id))
                         .filter(DBAnswer.test_attempt_id.in_(attempt_ids))
                         .group_by(DBAnswer.test_attempt_id).all()):
            ans_counts[aid] = cnt
    trend = []
    for a in all_attempts:
        a_total = a.mock_test.total_questions or 0
        a_att = ans_counts.get(a.id, 0)
        a_correct = a.score or 0
        a_wrong = max(0, a_att - a_correct)
        a_net = round(a_correct * MARKS_PER_Q - a_wrong * NEG, 2)
        a_max = round(a_total * MARKS_PER_Q, 2)
        trend.append({
            "title": a.mock_test.title,
            "date": a.completed_at.strftime("%d %b %Y") if a.completed_at else "",
            "net_percentage": round(a_net / a_max * 100, 1) if a_max else 0.0,
            "accuracy": round(a_correct / a_att * 100, 1) if a_att else 0.0,
        })
    overall_attempted = sum(ans_counts.values())
    overall_correct = sum(a.score or 0 for a in all_attempts)
    overall_accuracy = round(overall_correct / overall_attempted * 100, 1) if overall_attempted else 0.0

    candidate = {
        "name": current_user.name or (current_user.email.split("@")[0] if current_user.email else "Aspirant"),
        "tests_taken": len(all_attempts),
        "overall_accuracy": overall_accuracy,
        "overall_by_subject": overall_by_subject,
        "trend": trend[-8:],
    }

    stats = {
        "score": db_attempt.score, "total_questions": total,
        "attempted": attempted, "correct": correct, "wrong": wrong, "blank": blank,
        "accuracy": accuracy, "net_marks": net_marks, "max_marks": max_marks, "net_percentage": net_pct,
        "by_subject": by_subject,
        "candidate": candidate,
    }

    # Build a detailed stats summary for the AI (it must work from the ACTUAL missed
    # questions to name specific topics — not give generic subject advice).
    single_subject = len(by_subject) == 1
    persistent_weak = sorted([b for b in overall_by_subject if b["attempted"] >= 3], key=lambda x: x["accuracy"])[:3]
    trend_line = " -> ".join(f"{t['net_percentage']}%" for t in candidate["trend"]) or "n/a"

    this_missed_block = "\n".join(
        f"  - [{m['subject']}] {m['text'][:160]} | Correct: {m['correct'][:70]}"
        + (f" | Why: {m['exp']}" if m['exp'] else "")
        for m in this_test_missed[:14]
    ) or "  (none — all correct)"
    history_block = "\n".join(f"  - [{s}] {t[:120]}" for s, t in history_missed[:25]) or "  (no prior misses on record)"

    summary = (
        f"Candidate: {candidate['name']}\n"
        f"Tests taken: {candidate['tests_taken']} | Overall accuracy: {overall_accuracy}%\n"
        f"Net-score% trend (oldest->newest): {trend_line}\n\n"
        f"THIS TEST ({db_attempt.mock_test.title}): "
        + (f"single-subject test on {by_subject[0]['subject']}. " if single_subject else "")
        + f"net {net_pct}%, accuracy {accuracy}% (attempted {attempted}/{total}, {correct} correct, {wrong} wrong, {blank} blank), "
        f"negative-marking cost {round(wrong*NEG,1)} marks.\n"
        + ("" if single_subject else
           "This test by subject: " + "; ".join(f"{b['subject']} {b['accuracy']}%" for b in by_subject) + "\n")
        + "\nCUMULATIVE subject accuracy (all tests): "
        + ("; ".join(f"{b['subject']} {b['accuracy']}% (n={b['attempted']})" for b in overall_by_subject) or "only this test so far")
        + f"\nPersistent weak subjects (>=3 attempts): {', '.join(b['subject'] for b in persistent_weak) or 'not enough data yet'}\n\n"
        f"QUESTIONS MISSED ON THIS TEST (use these to name the exact topics/concepts they fail):\n{this_missed_block}\n\n"
        f"RECURRING MISSES ACROSS THEIR TESTS (find repeating themes/topics here):\n{history_block}"
    )

    try:
        report = generate_mentor_report(summary, candidate_name=candidate["name"])
    except Exception as e:
        report = None
        ai_error = str(e)
    if report is None:
        return {"status": "partial", "stats": stats, "report": None,
                "message": f"Stats computed, but AI narrative failed: {ai_error}"}
    return {"status": "success", "stats": stats, "report": report}

def _student_context(db, user):
    """A concise, factual snapshot of the student for the mentor to personalise on."""
    today = datetime.date.today()
    name = user.name or (user.email or "student").split("@")[0]
    parts = [f"Name: {name}"]
    prof = db.query(DBStudentProfile).filter(DBStudentProfile.user_id == user.id).first()
    if prof:
        if prof.target_year:
            try:
                exam = study_planner.prelims_date(int(str(prof.target_year)[:4]))
                dleft = (exam - today).days
                if dleft > 0:
                    parts.append(f"Target: Prelims {prof.target_year} (~{dleft} days left)")
            except Exception:
                pass
        if prof.study_hours:
            parts.append(f"Study time/day: {prof.study_hours} hrs")
        if prof.optional_subject:
            parts.append(f"Optional: {prof.optional_subject}")
        if prof.medium:
            parts.append(f"Medium: {prof.medium}")
        if prof.mains_language:
            parts.append(f"Mains language: {prof.mains_language}")
        if prof.coaching_status:
            parts.append(f"Prep mode: {prof.coaching_status}")
        # Educational background — helps gauge baseline knowledge.
        if prof.education:
            edu = prof.education
            if prof.graduation_stream:
                edu += f" ({prof.graduation_stream})"
            parts.append(f"Education: {edu}")
        if prof.schooling_medium:
            parts.append(f"Schooling medium: {prof.schooling_medium}")
        if prof.degree_percentage:
            parts.append(f"Graduation score: {prof.degree_percentage}")
        if prof.additional_qualification:
            parts.append(f"Additional qualifications: {prof.additional_qualification}")
        # Working commitment — governs how much time they realistically have.
        if prof.working_professional:
            wk = "Working"
            if prof.prep_intensity:
                wk += f" {prof.prep_intensity}"
            if prof.work_experience:
                wk += f" — {prof.work_experience}"
            parts.append(wk)
        # Self-rated level + diagnostic — where to pitch difficulty.
        if prof.prep_level:
            parts.append(f"Prep stage: {prof.prep_level}")
        if prof.knowledge_level:
            parts.append(f"Self-rated knowledge: {prof.knowledge_level}")
        if prof.comprehension_skill:
            parts.append(f"CSAT/reasoning: {prof.comprehension_skill}")
        if prof.diagnostic_gs is not None or prof.diagnostic_csat is not None:
            parts.append(f"Diagnostic — Knowledge {prof.diagnostic_gs}%, "
                         f"Reasoning {prof.diagnostic_csat}%")
        if prof.reading_speed:
            parts.append(f"Reading style: {prof.reading_speed}")
        if prof.learning_style:
            parts.append(f"Learning style: {prof.learning_style}")
        # Past attempts — where they fell short before.
        if prof.attempts:
            parts.append(f"Previous attempts: {prof.attempts}")
        if prof.failure_stage:
            parts.append(f"Fell short at: {prof.failure_stage}")
        if prof.failure_reason:
            parts.append(f"What held them back: {prof.failure_reason}")
        if prof.strong_subjects:
            parts.append(f"Self-declared strong areas: {prof.strong_subjects}")
        if prof.weak_subjects:
            parts.append(f"Self-declared weak areas: {prof.weak_subjects}")
    answers = _gather_answers(db, user.id)
    km = prepos.build_knowledge_map(answers, today) if answers else []
    if km:
        parts.append("Weakest subjects (mastery / retention): " +
                     ", ".join(f"{s['subject']} {s['mastery']}%/{s['retention']}%" for s in km[:3]))
        strong = sorted(km, key=lambda x: -x["mastery"])[:2]
        parts.append("Strongest: " + ", ".join(f"{s['subject']} {s['mastery']}%" for s in strong))
        review_items = [{"times_seen": r.times_seen, "times_correct": r.times_correct,
                         "mastered": r.mastered, "next_review": r.next_review}
                        for r in db.query(DBReviewItem).filter(DBReviewItem.user_id == user.id).all()]
        attempts = [{"completed_at": a.completed_at, "score": a.score}
                    for a in db.query(DBTestAttempt).filter(DBTestAttempt.user_id == user.id).all()]
        syl = db.query(DBSyllabusProgress).filter(DBSyllabusProgress.user_id == user.id).count()
        cov = round(100 * syl / max(1, syllabus_tracker_data.total_topics()))
        sc = prepos.compute_scores(answers, review_items, attempts, cov, today)
        parts.append(f"Scores — Knowledge {sc['knowledge']}%, Readiness {sc['readiness']}%, "
                     f"Retention {sc['retention']}%, Consistency {sc['consistency']}%, "
                     f"Success probability {sc['success_probability']}%. Answered {sc['answered']} questions.")
    else:
        parts.append("No practice attempts logged yet — they are just getting started.")
    try:
        review_due = db.query(DBReviewItem).filter(
            DBReviewItem.user_id == user.id, DBReviewItem.next_review <= datetime.datetime.utcnow(),
            DBReviewItem.mastered == False).count()  # noqa: E712
        hb = (prof.study_hours if prof else None) or "2-4"
        mission = prepos.daily_mission(km, review_due, hb, None, today, _mission_profile(prof))
        if mission["tasks"]:
            parts.append("Today's planned mission: " + "; ".join(t["title"] for t in mission["tasks"][:4]))
    except Exception:
        pass
    try:
        mc = (db.query(DBAnswer).join(DBTestAttempt, DBAnswer.test_attempt_id == DBTestAttempt.id)
              .filter(DBTestAttempt.user_id == user.id, DBAnswer.is_correct == False).count())  # noqa: E712
        if mc:
            parts.append(f"Total mistakes in their notebook: {mc}")
    except Exception:
        pass
    try:
        _recent = (db.query(DBMentorTopic).filter(DBMentorTopic.user_id == user.id)
                   .order_by(DBMentorTopic.last_seen.desc()).limit(6).all())
        if _recent:
            parts.append("Recently studied with the mentor: " +
                         ", ".join(f"{t.topic} [{t.level}]" for t in _recent))
        _due = (db.query(DBMentorTopic).filter(
                    DBMentorTopic.user_id == user.id, DBMentorTopic.revised == False,
                    DBMentorTopic.revise_due != None,
                    DBMentorTopic.revise_due <= datetime.datetime.utcnow()).all())
        if _due:
            parts.append("Due for revision now (bring these up naturally): " +
                         ", ".join(t.topic for t in _due))
    except Exception:
        pass
    return "\n".join("- " + p for p in parts), name, km


def _mentor_suggestions(km, name):
    s = ["What should I do right now?"]
    if km:
        s.append(f"Teach me {km[0]['subject']}")
        s.append("Explain my most recent mistake")
        s.append("Am I on track to clear Prelims?")
    else:
        s.append("How do I start my UPSC preparation?")
        s.append("Make me a study routine")
    s.append("I'm feeling low — motivate me")
    return s[:5]


class MentorReviseRequest(BaseModel):
    topic: str
    days: Optional[int] = 3

_TOPIC_STOP = ("what should", "what do i", "motivate", "feeling", "on track",
               "make me", "my mistake", "my progress", "help me", "start my",
               "how do i start", "routine", "hello", "hi ", "thanks")
def _mentor_topic(msg):
    """Best-effort topic extraction from a student's chat message (for the learner model)."""
    import re as _re
    t = (msg or "").strip().lower()
    if len(t) < 4 or any(st in t for st in _TOPIC_STOP):
        return None
    t = _re.sub(r"^(what\s+is|what\s+are|what's|whats|explain|tell me about|define|describe|"
                r"who\s+is|who\s+are|meaning of|teach me(\s+about)?|give me(\s+an?)?|tell me)\b[:\s]*",
                "", t)
    t = t.strip(" ?.!:")
    if len(t) < 3 or len(t) > 70 or not _re.search(r"[a-z]", t):
        return None
    return t.title()

def _log_topic(db, user, topic):
    if not topic:
        return
    try:
        row = (db.query(DBMentorTopic)
               .filter(DBMentorTopic.user_id == user.id, DBMentorTopic.topic.ilike(topic)).first())
        now = datetime.datetime.utcnow()
        if row:
            row.times_seen = (row.times_seen or 1) + 1
            row.last_seen = now
            if row.times_seen >= 5:
                row.level = "confident"
            elif row.times_seen >= 3:
                row.level = "reviewing"
        else:
            db.add(DBMentorTopic(user_id=user.id, topic=topic, level="introduced",
                                 times_seen=1, first_seen=now, last_seen=now))
        db.commit()
    except Exception:
        db.rollback()

@app.post("/me/mentor/revise", tags=["AI Features"])
def mentor_revise(req: MentorReviseRequest, current_user: DBUser = Depends(get_current_user),
                  db: Session = Depends(get_db)):
    topic = (req.topic or "").strip()
    if not topic:
        raise HTTPException(status_code=400, detail="No topic to revise.")
    days = max(1, min(30, int(req.days or 3)))
    due = datetime.datetime.utcnow() + datetime.timedelta(days=days)
    row = (db.query(DBMentorTopic)
           .filter(DBMentorTopic.user_id == current_user.id, DBMentorTopic.topic.ilike(topic)).first())
    if row:
        row.revise_due = due
        row.revised = False
    else:
        db.add(DBMentorTopic(user_id=current_user.id, topic=topic, level="introduced",
                             times_seen=1, revise_due=due, revised=False))
    db.commit()
    return {"status": "success", "topic": topic, "days": days, "due_at": due.isoformat()}

@app.post("/me/mentor/revise/done", tags=["AI Features"])
def mentor_revise_done(req: MentorReviseRequest, current_user: DBUser = Depends(get_current_user),
                       db: Session = Depends(get_db)):
    row = (db.query(DBMentorTopic)
           .filter(DBMentorTopic.user_id == current_user.id,
                   DBMentorTopic.topic.ilike((req.topic or "").strip())).first())
    if row:
        row.revised = True
        db.commit()
    return {"status": "success"}

@app.get("/me/mentor/history", tags=["AI Features"])
def mentor_history(current_user: DBUser = Depends(get_current_user), db: Session = Depends(get_db)):
    rows = (db.query(DBChatMessage).filter(DBChatMessage.user_id == current_user.id)
            .order_by(DBChatMessage.id.desc()).limit(40).all())
    rows = list(reversed(rows))
    _ctx, name, km = _student_context(db, current_user)
    _due = (db.query(DBMentorTopic).filter(
                DBMentorTopic.user_id == current_user.id, DBMentorTopic.revised == False,
                DBMentorTopic.revise_due != None,
                DBMentorTopic.revise_due <= datetime.datetime.utcnow())
            .order_by(DBMentorTopic.revise_due.asc()).all())
    return {"status": "success",
            "due_revisions": [t.topic for t in _due],
            "messages": [{"role": r.role, "content": r.content} for r in rows],
            "suggestions": _mentor_suggestions(km, name),
            "greeting": (f"Namaste {name}! 🙏 I'm AIVORA — your guide, teacher, mentor and companion. "
                         f"I know your progress and I'm here for the whole journey. What's on your mind today?")}


@app.post("/ai/chat/", tags=["AI Features"])
def ai_chat(request: AIChatRequest, current_user: DBUser = Depends(get_current_user),
            db: Session = Depends(get_db)):
    msg = (request.message or "").strip()
    if not msg:
        raise HTTPException(status_code=400, detail="Say something to your mentor.")
    try:
        context, _name, _km = _student_context(db, current_user)
        passages = _retrieve_for_query(db, msg)
        if passages:
            context = (context +
                       "\n\nRelevant material from the student's AIVORA library (ground your answer in "
                       "this, prefer it over general knowledge, and don't invent facts beyond it):\n" +
                       passages)
        history = [{"role": r.role, "content": r.content} for r in
                   reversed(db.query(DBChatMessage).filter(DBChatMessage.user_id == current_user.id)
                            .order_by(DBChatMessage.id.desc()).limit(10).all())]
        reply = chat_with_mentor(msg, context, history)
        db.add(DBChatMessage(user_id=current_user.id, role="user", content=msg[:4000]))
        db.add(DBChatMessage(user_id=current_user.id, role="assistant", content=(reply or "")[:8000]))
        db.commit()
        _topic = _mentor_topic(msg)
        if _topic:
            _log_topic(db, current_user, _topic)
        return {"status": "success", "response": reply, "topic": _topic or ""}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI Error: {str(e)}")

@app.post("/questions/{question_id}/explain/", tags=["AI Features"])
def ai_explain_question(question_id: int, db: Session = Depends(get_db), current_user: DBUser = Depends(get_current_user)):
    """On-demand AI explanation for a specific question (used on the review screen)."""
    q = db.query(DBQuestion).filter(DBQuestion.id == question_id).first()
    if not q:
        raise HTTPException(status_code=404, detail="Question not found")
    correct_text = {"A": q.option_a, "B": q.option_b, "C": q.option_c, "D": q.option_d}.get(q.correct_answer, "")
    message = (
        "Explain this UPSC Prelims question for an aspirant. State briefly why the correct "
        "option is right and why the others are wrong. Be concise and exam-focused.\n\n"
        f"Question: {q.text}\n"
        f"A) {q.option_a}\nB) {q.option_b}\nC) {q.option_c}\nD) {q.option_d}\n"
        f"Correct answer: {q.correct_answer}) {correct_text}"
    )
    try:
        result = chat_with_mentor(message)
        return {"status": "success", "explanation": result}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI Error: {str(e)}")


# ══════════════════════════════════════════════════════════════════════════════
#  AIMENTORA Guided Success Programme (GSP) — Phase-1 spine (Prelims)
#  Macro "Guidance Engine": placement -> module state -> daily mission -> promotion.
#  Modules are data (gsp.py); gating reuses the live AML student model
#  (concept_mastery moved by /me/attempt). Pilot content: Fundamental Rights.
# ══════════════════════════════════════════════════════════════════════════════
import gsp as _gsp
from models import GspEnrollment as DBGspEnrollment, GspModuleProgress as DBGspModuleProgress


def _gsp_slim(states):
    """Compact per-module view for the state/dashboard."""
    return [{"module_id": s["module_id"], "order": s["order"], "title": s["title"],
             "stage": s["stage"], "stage_name": s["stage_name"], "objective": s["objective"],
             "concepts": s["progress"]["total"], "state": s["state"],
             "mastery_pct": round(100 * s["progress"]["mastery"]),
             "coverage_pct": round(100 * s["progress"]["coverage"]),
             "gate_ready": s["gate_ready"], "prereqs": s["prereqs"],
             "exit_mastery_pct": round(100 * s["exit_mastery"])}
            for s in states]


class GspPlacementIn(BaseModel):
    reset: Optional[bool] = False


@app.post("/me/programme/placement", tags=["GSP"])
def gsp_placement(payload: Optional[GspPlacementIn] = None,
                  db: Session = Depends(get_db), current_user: DBUser = Depends(get_current_user)):
    """Measure -> place the student at the entry stage/module and open enrollment."""
    uid = current_user.id
    prof = db.query(DBStudentProfile).filter(DBStudentProfile.user_id == uid).first()
    dg = getattr(prof, "diagnostic_gs", None) if prof else None
    placed_level = "Foundation"
    if dg is not None:
        placed_level = "Foundation" if dg < 50 else ("Standard Books" if dg < 75 else "Concept Integration")
    e = db.query(DBGspEnrollment).filter(DBGspEnrollment.user_id == uid).first()
    if not e:
        e = DBGspEnrollment(user_id=uid, track="prelims", current_stage=1,
                            current_module=_gsp.FIRST_MODULE, placed_level=placed_level, intensity="standard")
        db.add(e)
    elif payload and payload.reset:
        e.current_stage = 1
        e.current_module = _gsp.FIRST_MODULE
        e.placed_level = placed_level
    states = _gsp.modules_state(db, uid)
    e.readiness_pct = _gsp.readiness(states)
    db.commit()
    first = _gsp.MODULE_BY_ID[_gsp.FIRST_MODULE]
    return {"status": "success", "enrolled": True, "track": e.track, "placed_level": placed_level,
            "current_module": e.current_module, "readiness_pct": e.readiness_pct,
            "message": f"Placed at {placed_level}. Begin with '{first['title']}'."}


@app.get("/me/programme/state", tags=["GSP"])
def gsp_state(db: Session = Depends(get_db), current_user: DBUser = Depends(get_current_user)):
    """The whole journey: stages, every module with state + progress, current module."""
    uid = current_user.id
    e = db.query(DBGspEnrollment).filter(DBGspEnrollment.user_id == uid).first()
    if not e:
        return {"status": "success", "enrolled": False,
                "message": "Not placed yet — POST /me/programme/placement to begin."}
    states = _gsp.modules_state(db, uid)
    cur = _gsp.current_module(states)
    read = _gsp.readiness(states)
    if e.readiness_pct != read:
        e.readiness_pct = read
        if cur:
            e.current_module = cur["module_id"]
            e.current_stage = cur["stage"]
        db.commit()
    return {"status": "success", "enrolled": True, "track": e.track,
            "placed_level": e.placed_level, "intensity": e.intensity,
            "readiness_pct": read, "topic": "Fundamental Rights",
            "stages": _gsp.STAGES, "current_module": cur["module_id"] if cur else None,
            "mastered": sum(1 for s in states if s["state"] == "mastered"),
            "total_modules": len(states), "modules": _gsp_slim(states)}


@app.get("/me/programme/mission", tags=["GSP"])
def gsp_mission(db: Session = Depends(get_db), current_user: DBUser = Depends(get_current_user)):
    """The 3-part daily mission (mission + why + expected outcome) for the current module."""
    uid = current_user.id
    e = db.query(DBGspEnrollment).filter(DBGspEnrollment.user_id == uid).first()
    if not e:
        return {"status": "success", "enrolled": False,
                "message": "Not placed yet — POST /me/programme/placement to begin."}
    states = _gsp.modules_state(db, uid)
    cur = _gsp.current_module(states)
    read = _gsp.readiness(states)
    if cur is None:
        return {"status": "success", "enrolled": True, "done": True,
                "headline": "All Fundamental Rights modules mastered — topic complete.",
                "readiness_pct": read}
    p = cur["progress"]
    why = []
    if cur["progress"]["total"]:
        why.append(f"{cur['title']} is your current step ({cur['stage_name']} stage).")
    why.append(f"Mastery here is {round(100*p['mastery'])}% over {p['practiced']}/{p['total']} concepts "
               f"— the gate needs {round(100*cur['exit_mastery'])}%.")
    if cur["prereqs"]:
        why.append("Prerequisites cleared: " + ", ".join(cur["prereqs"]) + ".")
    gap = max(0, round(100 * (cur["exit_mastery"] - p["mastery"])))
    return {"status": "success", "enrolled": True, "readiness_pct": read,
            "headline": f"Work on {cur['title']} — {round(100*p['mastery'])}% mastered",
            "mission": {
                "module_id": cur["module_id"], "title": cur["title"],
                "objective": cur["objective"], "stage": cur["stage_name"],
                "read": "Study notes for this module (Fundamental Rights pack)",
                "practice": f"Solve {min(15, max(8, cur['checkpoint_q']))} questions on these concepts",
            },
            "why": why,
            "expected_outcome": (f"Complete today's practice → mastery {round(100*p['mastery'])}% "
                                 f"→ ~{min(100, round(100*p['mastery'])+8)}%; "
                                 f"clear this module's gate in a few sessions." if gap
                                 else "You're at the gate — take the checkpoint via /me/programme/promote.")}


@app.get("/me/programme/module/{module_id}", tags=["GSP"])
def gsp_module_detail(module_id: str, db: Session = Depends(get_db),
                      current_user: DBUser = Depends(get_current_user)):
    """Module detail: objective, concepts (with verified key-facts), gate, current mastery."""
    m = _gsp.MODULE_BY_ID.get(module_id)
    if not m:
        raise HTTPException(status_code=404, detail="Module not found")
    prog = _gsp.module_progress(db, current_user.id, m)
    facts = _gsp.concept_facts(db, m["concept_keys"]) if m["concept_keys"] else []
    return {"status": "success", "module_id": m["module_id"], "title": m["title"],
            "stage": m["stage"], "stage_name": m["stage_name"],
            "objective": m["objective"], "bloom": m["bloom"], "prerequisites": m["prereqs"],
            "exit_gate": {"mastery_pct": round(100 * m["exit_mastery"]),
                          "checkpoint_questions": m["checkpoint_q"],
                          "min_coverage_pct": round(100 * _gsp.MIN_COVERAGE)},
            "progress": {"mastery_pct": round(100 * prog["mastery"]),
                         "coverage_pct": round(100 * prog["coverage"]),
                         "practiced": prog["practiced"], "total": prog["total"],
                         "attempts": prog["attempts"]},
            "concepts": facts, "concept_keys": m["concept_keys"]}


@app.get("/me/programme/module/{module_id}/practice", tags=["GSP"])
def gsp_module_practice(module_id: str, n: int = 12, db: Session = Depends(get_db),
                        current_user: DBUser = Depends(get_current_user)):
    """Serve questions scoped to a module's concepts (answers go via POST /me/attempt)."""
    m = _gsp.MODULE_BY_ID.get(module_id)
    if not m:
        raise HTTPException(status_code=404, detail="Module not found")
    n = max(1, min(int(n or 12), 30))
    keys = m["concept_keys"]
    qs = []
    if keys:
        qs = (db.query(DBQuestion).filter(DBQuestion.concept_key.in_(keys)).limit(n).all())
    if len(qs) < n:                      # fallback: topic/subject match keeps the loop runnable
        from sqlalchemy import or_ as _or
        got = {q.id for q in qs}
        extra = (db.query(DBQuestion)
                 .filter(DBQuestion.subject.ilike("%polity%"),
                         _or(DBQuestion.topic.ilike("%fundamental right%"),
                             DBQuestion.chapter.ilike("%fundamental right%"),
                             DBQuestion.topic.ilike("%right to%")))
                 .limit(n * 2).all())
        for q in extra:
            if q.id not in got:
                qs.append(q); got.add(q.id)
            if len(qs) >= n:
                break
    out = [{"id": q.id, "text": q.text, "option_a": q.option_a, "option_b": q.option_b,
            "option_c": q.option_c, "option_d": q.option_d, "difficulty": q.difficulty,
            "concept_key": q.concept_key, "pattern": q.pattern} for q in qs[:n]]
    return {"status": "success", "module_id": module_id, "served": len(out),
            "note": "Post each answer to /me/attempt (attempt_context='programme') to move mastery.",
            "questions": out}


@app.post("/me/programme/promote", tags=["GSP"])
def gsp_promote(db: Session = Depends(get_db), current_user: DBUser = Depends(get_current_user)):
    """Evaluate the current module's exit gate; if passed, mark mastered and advance."""
    uid = current_user.id
    e = db.query(DBGspEnrollment).filter(DBGspEnrollment.user_id == uid).first()
    if not e:
        raise HTTPException(status_code=400, detail="Not placed yet — call /me/programme/placement first.")
    states = _gsp.modules_state(db, uid)
    cur = _gsp.current_module(states)
    if cur is None:
        return {"status": "success", "promoted": False, "done": True,
                "message": "All modules already mastered."}
    if not cur["gate_ready"]:
        p = cur["progress"]
        return {"status": "success", "promoted": False, "module_id": cur["module_id"],
                "message": (f"Not yet. {cur['title']}: mastery {round(100*p['mastery'])}% "
                            f"(need {round(100*cur['exit_mastery'])}%), coverage {round(100*p['coverage'])}% "
                            f"(need {round(100*_gsp.MIN_COVERAGE)}%), {p['attempts']} attempts "
                            f"(need {_gsp.MIN_ATTEMPTS}). Keep practising this module.")}
    row = (db.query(DBGspModuleProgress)
           .filter(DBGspModuleProgress.user_id == uid,
                   DBGspModuleProgress.module_id == cur["module_id"]).first())
    if not row:
        row = DBGspModuleProgress(user_id=uid, module_id=cur["module_id"])
        db.add(row)
    row.state = "mastered"
    row.mastery = cur["progress"]["mastery"]
    row.mastered_at = datetime.datetime.utcnow()
    db.commit()
    states = _gsp.modules_state(db, uid)
    nxt = _gsp.current_module(states)
    read = _gsp.readiness(states)
    e.readiness_pct = read
    if nxt:
        e.current_module = nxt["module_id"]
        e.current_stage = nxt["stage"]
    db.commit()
    return {"status": "success", "promoted": True, "mastered_module": cur["module_id"],
            "mastered_title": cur["title"], "readiness_pct": read,
            "next_module": (nxt["module_id"] if nxt else None),
            "next_title": (nxt["title"] if nxt else None),
            "message": (f"🏆 Mastered {cur['title']}! Next: {nxt['title']}." if nxt
                        else f"🏆 Mastered {cur['title']} — Fundamental Rights journey complete!")}


@app.post("/admin/gsp/seed-fr", tags=["Admin"])
def gsp_seed_fr(db: Session = Depends(get_db), admin: DBUser = Depends(require_admin)):
    """Idempotently insert the grounded Fundamental-Rights seed questions (tagged with
    each module's concept_key) so GSP modules are practiceable/promotable."""
    import gsp_seed as _seed
    inserted = skipped = 0
    for q in _seed.SEED_QUESTIONS:
        exists = (db.query(DBQuestion)
                  .filter(DBQuestion.book == _seed.SENTINEL,
                          DBQuestion.concept_key == q["ck"],
                          DBQuestion.text == q["text"]).first())
        if exists:
            skipped += 1
            continue
        db.add(DBQuestion(
            text=q["text"], option_a=q["a"], option_b=q["b"], option_c=q["c"], option_d=q["d"],
            correct_answer=q["correct"], explanation=q["exp"],
            subject=_seed.SUBJECT, topic=_seed.TOPIC, chapter=_seed.TOPIC, book=_seed.SENTINEL,
            difficulty=q["difficulty"], concept_key=q["ck"],
            pattern=q["pattern"], question_type=q["pattern"]))
        inserted += 1
    db.commit()
    total = db.query(DBQuestion).filter(DBQuestion.book == _seed.SENTINEL).count()
    return {"status": "success", "inserted": inserted, "skipped": skipped, "seed_questions_live": total}
